from io import BytesIO
from pathlib import Path

from docx import Document
import pytest

from calendar_pedagoga.matching import MatchStatus, match_utp_to_program
from calendar_pedagoga.parsing import Hours, parse_utp
from calendar_pedagoga.program_parsing import parse_program
from calendar_pedagoga.resolve_utp import (
    AUTO_WORKLOAD_WARNING,
    UtpResolutionError,
    resolve_utp,
)
from calendar_pedagoga.scheduling import build_schedule
from calendar_pedagoga.upload_validation import UploadPurpose, validate_upload


REFERENCES = Path(__file__).resolve().parents[1] / "references"


def _program_without_utp() -> bytes:
    document = Document()
    document.add_paragraph("Дополнительная программа «ТЕСТ»")
    document.add_paragraph("Цель программы: проверка извлечения УТП.")
    document.add_paragraph("Задачи программы:")
    document.add_paragraph("научиться находить таблицу часов.")
    heading = document.add_paragraph()
    heading.add_run("Содержание программы 2-го года обучения").bold = True
    topic = document.add_paragraph()
    topic.add_run("1. Введение").bold = True
    document.add_paragraph("Знакомство с программой.")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_embedded_utp_from_tour_guides_program() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    program = validate_upload(
        UploadPurpose.PROGRAM, program_path.name, program_path.read_bytes()
    )
    result = resolve_utp(None, program)
    assert result.table_totals == Hours(72, 27, 45)
    assert len(result.topics) >= 20
    assert result.metadata.study_weeks == 36
    assert result.metadata.hours_per_week == 2
    assert result.metadata.workload_provenance == "derived_36x2"
    assert AUTO_WORKLOAD_WARNING in result.warnings
    schedule = build_schedule(result)
    assert len(schedule.weeks) == 36


def test_separate_utp_has_priority_over_embedded_table() -> None:
    key_utp = REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx"
    tp_program = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    validated_utp = validate_upload(UploadPurpose.UTP, key_utp.name, key_utp.read_bytes())
    validated_program = validate_upload(
        UploadPurpose.PROGRAM, tp_program.name, tp_program.read_bytes()
    )
    result = resolve_utp(validated_utp, validated_program)
    assert result.table_totals == Hours(72, 22, 50)
    assert len(result.topics) == 13
    assert result.metadata.workload_provenance == "document"
    assert AUTO_WORKLOAD_WARNING not in result.warnings


def test_key_regression_separate_files() -> None:
    utp_path = REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx"
    program_path = REFERENCES / "Программа КЛЮЧ.DOC"
    validated_utp = validate_upload(UploadPurpose.UTP, utp_path.name, utp_path.read_bytes())
    validated_program = validate_upload(
        UploadPurpose.PROGRAM, program_path.name, program_path.read_bytes()
    )
    utp = resolve_utp(validated_utp, validated_program)
    program = parse_program(
        validated_program.content,
        validated_program.filename,
        study_year=2,
    )
    matches = match_utp_to_program(utp.topics, program.content_items)
    assert len(utp.topics) == 13
    assert utp.table_totals == Hours(72, 22, 50)
    assert utp.metadata.study_weeks == 36
    assert sum(match.status is not MatchStatus.NOT_MATCHED for match in matches) == 13
    assert len(build_schedule(utp).weeks) == 36


def test_program_without_embedded_or_separate_utp_fails() -> None:
    program = validate_upload(
        UploadPurpose.PROGRAM,
        "program.docx",
        _program_without_utp(),
    )
    with pytest.raises(UtpResolutionError, match="не найден учебно-тематический план"):
        resolve_utp(None, program)
