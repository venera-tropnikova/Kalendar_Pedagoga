from io import BytesIO
from pathlib import Path

from docx import Document
import pytest

from calendar_pedagoga.matching import MatchStatus, match_utp_to_program
from calendar_pedagoga.parsing import Hours, Topic
from calendar_pedagoga.program_parsing import (
    LegacyDocUnsupportedError,
    ProgramContentItem,
    convert_legacy_doc,
    extract_attestation_statements,
    parse_age_range,
    parse_duration_years,
    parse_program_docx,
    parse_program,
    study_year_label,
)
from calendar_pedagoga.parsing import parse_utp


REFERENCES = Path(__file__).resolve().parents[1] / "references"


def _program_docx() -> bytes:
    document = Document()
    headings = {
        "Содержание программы 2-го года обучения", "1. Введение",
        "2. Краеведение", "Моя семья", "3.1 Пеший туризм",
    }
    for text in (
        "Дополнительная общеобразовательная программа «КЛЮЧ»",
        "Срок реализации программы: 2 года",
        "Возраст обучающихся: 11-12 лет",
        "Цель программы: развитие интереса к родному краю.",
        "Задачи программы:",
        "изучить историю города;",
        "развить навыки наблюдения.",
        "Формы организации занятий: групповая работа.",
        "Методы обучения: практический метод.",
        "Ожидаемые результаты:",
        "обучающийся знает историю города.",
        "Содержание программы 2-го года обучения",
        "1. Введение",
        "Знакомство с программой и правилами работы.",
        "2. Краеведение",
        "Изучение родного края.",
        "Моя семья",
        "История семьи и семейные традиции.",
        "3.1 Пеший туризм",
        "Основы организации пешего похода.",
        "По окончанию второго года обучения",
    ):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        if text in headings:
            run.bold = True
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_parse_program_docx_preserves_source_wording() -> None:
    result = parse_program_docx(_program_docx())

    assert result.title == "КЛЮЧ"
    assert result.duration == "2 года"
    assert result.duration_years == 2
    assert result.student_age == "11-12 лет"
    assert result.age_min == 11
    assert result.age_max == 12
    assert result.goal == "развитие интереса к родному краю."
    assert result.tasks == (
        "изучить историю города;",
        "развить навыки наблюдения.",
    )
    assert result.lesson_forms == ("групповая работа.",)
    assert result.teaching_methods == ("практический метод.",)
    assert result.expected_results == (
        "обучающийся знает историю города.",
    )
    assert result.knowledge_outcomes == ()
    assert result.skill_outcomes == ()
    assert result.content_items[0].title == "Введение"
    assert "Знакомство с программой" in result.content_items[0].content
    assert result.attestation_statements == ()


def test_attestation_extracted_only_from_explicit_phrase() -> None:
    assert extract_attestation_statements(
        "Ожидаемые результаты:\nобучающийся знает правила."
    ) == ()
    assert extract_attestation_statements(
        "Формы контроля: текущий опрос и практическое задание."
    ) == ()
    found = extract_attestation_statements(
        "Итоговая аттестация проводится в конце года обучения."
    )
    assert found == ("Итоговая аттестация проводится в конце года обучения.",)


def test_legacy_doc_without_libreoffice_is_reported(monkeypatch) -> None:
    monkeypatch.setattr("calendar_pedagoga.program_parsing.find_libreoffice", lambda: None)
    with pytest.raises(LegacyDocUnsupportedError):
        convert_legacy_doc(b"legacy")


def test_matching_uses_deterministic_priority() -> None:
    topics = (
        Topic("1", "Введение", Hours(2, 2, 0), "Введение", True),
        Topic(None, "Моя семья", Hours(2, 1, 1), "Краеведение"),
        Topic("3.1", "Пеший туризм", Hours(6, 2, 4), "Туризм"),
        Topic(None, "Несуществующая тема", Hours(1, 1, 0), "Раздел"),
    )
    items = (
        ProgramContentItem("1", "Введение", "Текст 1"),
        ProgramContentItem(None, "МОЯ   СЕМЬЯ!", "Текст 2", "Краеведение"),
        ProgramContentItem("3.1", "Походы пешком", "Текст 3", "Туризм"),
    )

    matches = match_utp_to_program(topics, items)

    assert [match.status for match in matches] == [
        MatchStatus.EXACT,
        MatchStatus.NORMALIZED,
        MatchStatus.NUMBER_MATCH,
        MatchStatus.NOT_MATCHED,
    ]


def test_ambiguous_normalized_match_is_not_guessed() -> None:
    topic = Topic(None, "Моя семья", Hours(2, 1, 1), "Краеведение")
    items = (
        ProgramContentItem(None, "Моя семья!", "Первый текст"),
        ProgramContentItem(None, "МОЯ СЕМЬЯ", "Второй текст"),
    )

    match = match_utp_to_program((topic,), items)[0]

    assert match.status is MatchStatus.NOT_MATCHED
    assert match.program_item is None
    assert len(match.ambiguous_candidates) == 2


def test_unique_contained_title_is_text_match() -> None:
    topic = Topic(None, "Аптечка", Hours(2, 1, 1), "Раздел")
    items = (ProgramContentItem(None, "Медицинская аптечка.", "Текст"),)
    match = match_utp_to_program((topic,), items)[0]
    assert match.status is MatchStatus.TEXT_MATCH


def test_real_key_program_matches_all_13_positions() -> None:
    program_path = REFERENCES / "Программа КЛЮЧ.DOC"
    utp_path = REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx"
    program = parse_program(program_path.read_bytes(), program_path.name, study_year=2)
    utp = parse_utp(utp_path)
    matches = match_utp_to_program(utp.topics, program.content_items)

    assert program.title == "КЛЮЧ"
    assert program.duration == "3 года"
    assert program.duration_years == 3
    assert program.student_age == "8 – 11 лет"
    assert program.age_min == 8
    assert program.age_max == 11
    assert parse_duration_years("несколько лет") is None
    assert parse_age_range("11 лет") == (None, None)
    assert len(program.tasks) == 3
    assert len(program.content_items) == 16
    assert len(matches) == 13
    assert all(match.status is not MatchStatus.NOT_MATCHED for match in matches)
def test_tour_guides_year1_program_finds_content_items() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    program = parse_program(program_path.read_bytes(), program_path.name, study_year=1)
    assert program.title == "Туристы-проводники"
    assert len(program.content_items) >= 20
    titles = {item.title for item in program.content_items}
    assert "Основы туристской подготовки" in titles
    assert any("Туристские путешествия" in title for title in titles)
    assert any(item.content for item in program.content_items)
    assert any("палатк" in item.casefold() for item in program.skill_outcomes)
    assert any("ориентир" in item.casefold() for item in program.knowledge_outcomes)
    assert all("тест" not in item.casefold() for item in program.expected_results)


def test_study_year_label_uses_same_hints_as_pipeline() -> None:
    assert study_year_label(None) is None
    assert study_year_label("второй") == "2 год обучения"
    assert (
        study_year_label(
            None,
            "УТП из файла «Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx»",
            "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx",
        )
        == "1 год обучения"
    )
    assert study_year_label("специальный набор") == "специальный набор год обучения"


def test_key_year2_program_parses_know_and_able_outcomes() -> None:
    program_path = REFERENCES / "Программа КЛЮЧ.DOC"
    program = parse_program(program_path.read_bytes(), program_path.name, study_year=2)
    assert any("истори" in item.casefold() for item in program.knowledge_outcomes)
    assert any("палатк" in item.casefold() for item in program.skill_outcomes)
    # Не смешивать исходы другого года.
    assert not any("адрес школы" in item.casefold() for item in program.knowledge_outcomes)
