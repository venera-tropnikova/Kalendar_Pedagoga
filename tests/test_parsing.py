from io import BytesIO
from pathlib import Path

from docx import Document
import pytest

from calendar_pedagoga.parsing import Hours, UtpMetadata, UtpParseResult, parse_utp
from calendar_pedagoga.validation import validate_utp


REFERENCES = Path(__file__).resolve().parents[1] / "references"


def test_parse_key_utp() -> None:
    result = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    assert result.metadata.program_name == "КЛЮЧ"
    assert result.metadata.academic_year == "2026–2027"
    assert result.metadata.study_year == "второй"
    assert result.metadata.student_age == "11 -12 лет"
    assert result.metadata.hours_per_week == 2
    assert result.metadata.hours_per_year == 72
    assert result.metadata.study_weeks == 36
    assert result.metadata.teacher_name == "И.М.Саранцева"
    assert len(result.sections) == 7
    assert [section.number for section in result.sections] == [
        "1", "2", "3", "4", "5", "6", "7"
    ]
    assert [section.is_standalone_position for section in result.sections] == [
        True, False, False, False, True, True, True
    ]
    assert len(result.topics) == 13
    family = next(topic for topic in result.topics if topic.title == "Моя семья")
    assert family.number is None
    assert family.parent_section == "Краеведение"
    assert result.table_totals == Hours(72, 22, 50)
    assert result.warnings == ()


def test_parse_tour_guides_utp_and_report_conflict() -> None:
    result = parse_utp(REFERENCES / "УТП ТП 3г. 2ч.docx")
    assert result.metadata.program_name == "Туристы проводники"
    assert result.metadata.academic_year == "2026–2027"
    assert result.metadata.study_year == "третий"
    assert result.metadata.student_age == "15-16 лет"
    assert result.metadata.hours_per_week == 2
    assert result.metadata.hours_per_year == 140
    assert result.metadata.study_weeks == 36
    assert result.metadata.stated_schedule_hours == 72
    assert result.metadata.teacher_name == "И.М.Саранцева"
    assert len(result.sections) == 5
    assert len(result.topics) == 22
    assert result.table_totals == Hours(72, 24, 48)
    assert result.topics[-2].hours == Hours(10, 0, 10)
    assert result.topics[-1].hours == Hours(10, 0, 10)
    assert len(result.warnings) == 1
    assert "информационная справка: 140 ч." in result.warnings[0]
    assert "36 недель × 2 часа: 72 ч." in result.warnings[0]
    assert "итоговая строка УТП: 72 ч." in result.warnings[0]


def _five_column_utp_docx(*, extra_calendar: bool = False) -> bytes:
    document = Document()
    document.add_paragraph("Учебно-тематический план")
    if extra_calendar:
        calendar = document.add_table(rows=3, cols=8)
        headers = (
            "Месяц",
            "Неделя",
            "Тема",
            "Теоретические занятия",
            "Практические занятия",
            "Планируемый результат",
            "Вид контроля",
            "Примечание",
        )
        for cell, value in zip(calendar.rows[0].cells, headers, strict=True):
            cell.text = value
        calendar.rows[1].cells[0].text = "Сентябрь"
        calendar.rows[1].cells[1].text = "1"
        calendar.rows[1].cells[2].text = "Вводное занятие"
    table = document.add_table(rows=8, cols=5)
    headers = ("№", "Тема", "всего", "теория/лекции", "практика")
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    rows = (
        ("1", "Раздел первый", "4", "2", "2"),
        ("1.1", "Введение", "2", "2", "0"),
        ("1.2", "Практика", "2", "0", "2"),
        ("2", "Раздел второй", "6", "2", "4"),
        ("2.1", "Теория раздела", "2", "2", "0"),
        ("2.2", "Практика раздела", "4", "0", "4"),
        ("Итого", "", "10", "4", "6"),
    )
    for index, values in enumerate(rows, start=1):
        for cell, value in zip(table.rows[index].cells, values, strict=True):
            cell.text = value
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_five_column_parser_reads_theory_practice() -> None:
    result = parse_utp(_five_column_utp_docx())
    assert result.table_totals == Hours(10, 4, 6)
    assert [section.title for section in result.sections] == [
        "Раздел первый",
        "Раздел второй",
    ]
    assert [topic.title for topic in result.topics] == [
        "Введение",
        "Практика",
        "Теория раздела",
        "Практика раздела",
    ]
    assert result.topics[0].hours == Hours(2, 2, 0)
    assert result.topics[-1].hours == Hours(4, 0, 4)


def test_selects_utp_table_among_calendar_and_other_tables() -> None:
    result = parse_utp(_five_column_utp_docx(extra_calendar=True))
    assert result.table_totals == Hours(10, 4, 6)
    assert len(result.sections) == 2
    assert all(topic.parent_section for topic in result.topics)


def test_five_column_embedded_utp_in_tour_guides_program() -> None:
    result = parse_utp(REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx")
    assert result.table_totals == Hours(72, 27, 45)
    assert len(result.sections) == 5
    assert len(result.topics) >= 20
    assert any(topic.number and "." in topic.number for topic in result.topics)
    assert all(section.title for section in result.sections)


def test_parse_utp_rejects_calendar_form() -> None:
    with pytest.raises(ValueError, match="не найдена таблица УТП"):
        parse_utp(REFERENCES / "Календарный план.docx")


def test_validation_reports_arithmetic_and_table_mismatch() -> None:
    result = UtpParseResult(
        metadata=UtpMetadata(hours_per_week=2, hours_per_year=72, study_weeks=36),
        sections=(),
        topics=(),
        table_totals=Hours(72, 20, 40),
    )
    warnings = validate_utp(result)
    assert any("теория + практика" in warning for warning in warnings)
    assert any("Сумма тем" in warning for warning in warnings)
