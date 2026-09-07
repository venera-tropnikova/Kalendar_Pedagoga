from io import BytesIO
from types import SimpleNamespace

from docx import Document
import pytest

from calendar_pedagoga import docx_qa as qa
from calendar_pedagoga.docx_generation import (
    _allow_row_split, _prevent_row_split, _repeat_table_header_rows,
    _merge_month_cells_by_page_segments, _columns_for_table,
)


def test_vertical_minimum_tracks_text_font_and_margins_without_changing_layout():
    from calendar_pedagoga.docx_generation import _protect_vertical_cell_height
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    doc = Document()
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    for row, text, size in zip(table.rows[2:], ('7\n01–07.03', '28\n15–21.03', '28\n15–21.03'), (12, 12, 18)):
        cell = row.cells[1]
        cell.text = text
        cell.paragraphs[0].runs[0].font.size = Pt(size)
        direction = OxmlElement('w:textDirection')
        direction.set(qn('w:val'), 'btLr')
        cell._tc.get_or_add_tcPr().append(direction)
    before = [[cell._tc.xml for cell in row.cells] for row in table.rows]
    _protect_vertical_cell_height(table)
    heights = [int(row._tr.trPr.find(qn('w:trHeight')).get(qn('w:val'))) for row in table.rows[2:]]
    assert heights[0] == heights[1]
    assert heights[2] > heights[1] > 45.89 * 20
    assert all(row._tr.trPr.find(qn('w:trHeight')).get(qn('w:hRule')) == 'atLeast' for row in table.rows[2:])
    assert before == [[cell._tc.xml for cell in row.cells] for row in table.rows]
    assert not table.rows[0]._tr.xpath('./w:trPr/w:trHeight')
    xml = table._tbl.xml
    _protect_vertical_cell_height(table)
    assert table._tbl.xml == xml


def test_horizontal_cells_do_not_receive_vertical_minimum():
    from calendar_pedagoga.docx_generation import _protect_vertical_cell_height
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(2, 1).text = 'Long horizontal content ' * 10
    _protect_vertical_cell_height(table)
    assert not table._tbl.xpath('.//w:trHeight')


@pytest.mark.parametrize('direction', ['btLr', 'tbRl'])
def test_vertical_minimum_uses_actual_text_and_inherited_padding(direction):
    from calendar_pedagoga.docx_generation import _protect_vertical_cell_height
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    cell = table.cell(2, 0)
    cell.text = 'Май'
    cell.paragraphs[0].runs[0].font.size = Pt(12)
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    cell.paragraphs[0].paragraph_format.line_spacing = 1.2
    marker = OxmlElement('w:textDirection')
    marker.set(qn('w:val'), direction)
    cell._tc.get_or_add_tcPr().append(marker)
    def height():
        row = table.rows[2]
        for old in row._tr.xpath('./w:trPr/w:trHeight'):
            old.getparent().remove(old)
        _protect_vertical_cell_height(table)
        return int(row._tr.trPr.find(qn('w:trHeight')).get(qn('w:val')))
    short = height()
    cell.paragraphs[0].runs[0].text = 'Сентябрь / Октябрь'
    long = height()
    assert long > short
    margins = OxmlElement('w:tblCellMar')
    left = OxmlElement('w:left')
    left.set(qn('w:w'), '200')
    margins.append(left)
    table._tbl.tblPr.append(margins)
    inherited_left = int(table.style.element.xpath('./w:tblPr/w:tblCellMar/w:left')[0].get(qn('w:w')))
    assert height() == long + 200 - inherited_left
    assert cell.paragraphs[0].paragraph_format.space_after == Pt(3)
    assert cell.paragraphs[0].paragraph_format.line_spacing == 1.2


def test_allow_split_removes_only_pagination_constraints():
    doc = Document()
    table = doc.add_table(rows=3, cols=8)
    _repeat_table_header_rows(table)
    row = table.rows[2]
    row.cells[2].text = 'Unchanged content'
    _prevent_row_split(row)
    row.cells[2].paragraphs[0].paragraph_format.page_break_before = True
    _allow_row_split(row)
    assert not row._tr.xpath('.//w:cantSplit | .//w:pageBreakBefore')
    assert row.cells[2].text == 'Unchanged content'
    assert len(table._tbl.xpath('.//w:tblHeader')) == 2


def test_merge_does_not_bridge_excluded_split_row():
    doc = Document()
    table = doc.add_table(rows=5, cols=8)
    for row in table.rows[2:]:
        row.cells[0].text = 'Сентябрь'
    _merge_month_cells_by_page_segments(
        table, _columns_for_table(table), ('Сентябрь',) * 3, ((0, 2),),
    )
    assert not table._tbl.xpath('.//w:vMerge')


def _source():
    doc = Document()
    table = doc.add_table(rows=4, cols=3)
    for row, values in zip(table.rows[2:], [('Month', '19', 'abcdef'), ('Month', '20', 'gh')]):
        for cell, value in zip(row.cells, values):
            cell.text = value
    result = BytesIO()
    doc.save(result)
    return result.getvalue()


def _pdf(monkeypatch, fragments):
    import pymupdf
    class Pdf:
        def __enter__(self): return self
        def __exit__(self, *args): pass
        def __iter__(self):
            for rows in fragments:
                table = SimpleNamespace(col_count=3, extract=lambda rows=rows: [[], []] + rows)
                yield SimpleNamespace(find_tables=lambda table=table: SimpleNamespace(tables=[table]))
    monkeypatch.setattr(pymupdf, 'open', lambda **kwargs: Pdf())


def test_pdf_measures_continuation_and_monthly_week_numbers(monkeypatch):
    _pdf(monkeypatch, [[['Month', '19', 'abc']], [['', '', 'def'], ['Month', '20', 'gh']]])
    assert qa._data_row_page_spans_pdf(_source(), b'pdf', 2) == (
        qa.DataRowPageSpan(1, 2, True), qa.DataRowPageSpan(2, 2, True),
    )
    layouts = qa._data_row_page_layout_pdf(_source(), b'pdf', 2)
    assert layouts is not None
    first = layouts[0].segments
    assert [segment.page_number for segment in first] == [1, 2]
    assert [segment.cells[:2] for segment in first] == [('Month', '19')] * 2
    assert ''.join(segment.cells[2] for segment in first) == 'abcdef'


@pytest.mark.parametrize('month,week', [('Mon', '19'), ('Month', '1'), ('', '')])
def test_split_with_incomplete_identifier_is_unsafe(monkeypatch, month, week):
    _pdf(monkeypatch, [[[month, week, 'abc']], [['', '', 'def'], ['Month', '20', 'gh']]])
    spans = qa._data_row_page_spans_pdf(_source(), b'pdf', 2)
    assert spans[0] == qa.DataRowPageSpan(1, 2, False)
    assert spans[1] == qa.DataRowPageSpan(2, 2, True)


def test_identifiers_may_be_intact_on_continuation_page(monkeypatch):
    _pdf(monkeypatch, [[['', '', 'abc']], [['Month', '19', 'def'], ['Month', '20', 'gh']]])
    assert qa._data_row_page_spans_pdf(_source(), b'pdf', 2)[0].split_safe


def test_generation_materializes_exact_non_splitting_page_segments(monkeypatch):
    from calendar_pedagoga import docx_generation as gen
    def populate(document, *args, **kwargs):
        table = document.add_table(rows=4, cols=8)
        _repeat_table_header_rows(table)
        for row, week, body in zip(table.rows[2:], ('19\nDate 19', '20\nDate 20'), ('abcdef', 'gh')):
            row.cells[0].text = 'Month'
            row.cells[1].text = week
            row.cells[2].text = body
        return table, _columns_for_table(table), ('Month', 'Month')
    monkeypatch.setattr(gen, '_load_template', lambda template: Document())
    monkeypatch.setattr(gen, '_populate_calendar_table', populate)
    layouts = (
        qa.DataRowPageLayout(
            qa.DataRowPageSpan(1, 2, True),
            (
                qa.DataRowPageSegment(1, ('Month', '19\nDate 19', 'abc', '', '', '', '', '')),
                qa.DataRowPageSegment(2, ('Month', '19\nDate 19', 'def', '', '', '', '', '')),
            ),
        ),
        qa.DataRowPageLayout(
            qa.DataRowPageSpan(2, 2, True),
            (qa.DataRowPageSegment(2, ('Month', '20\nDate 20', 'gh', '', '', '', '', '')),),
        ),
    )
    monkeypatch.setattr(qa, 'detect_data_row_page_layout', lambda *args, **kwargs: layouts)
    monkeypatch.setattr(
        qa,
        'detect_data_row_page_spans',
        lambda *args, **kwargs: tuple(
            qa.DataRowPageSpan(page, page, True) for page in (1, 2, 2)
        ),
    )
    doc = Document(BytesIO(gen.generate_calendar_docx(
        None, (None, None), SimpleNamespace(uses_organization_template=False), '2026–2027')))
    rows = doc.tables[0].rows[2:]
    assert len(rows) == 3
    assert [row.cells[0].text for row in rows] == ['Month'] * 3
    assert [row.cells[1].text for row in rows] == ['19\nDate 19', '19\nDate 19', '20\nDate 20']
    assert rows[0].cells[2].text + rows[1].cells[2].text == 'abcdef'
    assert rows[2].cells[2].text == 'gh'
    assert all(row._tr.xpath('./w:trPr/w:cantSplit') for row in rows)
    assert not rows[0]._tr.xpath('.//w:vMerge')
    assert not rows[1]._tr.xpath('.//w:vMerge')
    assert not doc._element.xpath('.//w:pageBreakBefore')


@pytest.mark.parametrize('fragments', [
    [[['Month', '19', 'wrong']]],
    [[['Month', '19', 'abc']]],
    [[['Month', '19', 'abcdef'], ['Month', '20', 'gh'], ['Month', '20', 'gh']]],
])
def test_pdf_refuses_missing_changed_or_duplicate_content(monkeypatch, fragments):
    _pdf(monkeypatch, fragments)
    assert qa._data_row_page_spans_pdf(_source(), b'pdf', 2) is None


def test_month_label_verification_uses_one_consistent_render(monkeypatch):
    word_calls = []
    libreoffice_calls = []

    def word_pdf(_content):
        word_calls.append(True)
        # A second export would reproduce the former Word-COM failure.
        return b'word-pdf' if len(word_calls) == 1 else None

    def libreoffice_pdf(_content):
        libreoffice_calls.append(True)
        return b'libreoffice-pdf-with-different-pagination'

    monkeypatch.setattr(qa, '_docx_to_pdf_bytes_word', word_pdf)
    monkeypatch.setattr(qa, '_docx_to_pdf_bytes_libreoffice', libreoffice_pdf)
    monkeypatch.setattr(
        qa,
        'detect_data_row_indices_by_page',
        lambda *args, **kwargs: pytest.fail('independent page mapping must not be used'),
    )
    monkeypatch.setattr(
        qa,
        '_data_row_page_spans_pdf',
        lambda content, pdf, total_rows: (
            qa.DataRowPageSpan(1, 2, True),
        ) if pdf == b'word-pdf' and total_rows == 1 else None,
    )

    pages = tuple(
        SimpleNamespace(
            rect=SimpleNamespace(height=800),
            # Rotated week identifiers may be absent from PDF text extraction.
            get_text=lambda _kind: [],
            find_tables=lambda: SimpleNamespace(tables=[]),
        )
        for _ in range(2)
    )

    class Pdf:
        def __len__(self): return len(pages)
        def __getitem__(self, index): return pages[index]
        def close(self): pass

    import pymupdf
    monkeypatch.setattr(pymupdf, 'open', lambda **kwargs: Pdf())
    monkeypatch.setattr(
        qa,
        '_month_words_on_page',
        lambda page, month: [(20, 100, month)] if page is pages[0] else [],
    )

    source = Document()
    table = source.add_table(rows=3, cols=3)
    table.rows[2].cells[0].text = 'Сентябрь'
    table.rows[2].cells[1].text = '1'
    table.rows[2].cells[2].text = 'Content'
    source_buffer = BytesIO()
    source.save(source_buffer)
    assert qa.verify_month_labels_by_page(
        source_buffer.getvalue(), months=('Сентябрь',)
    ) == ()
    assert len(word_calls) == 1
    assert not libreoffice_calls


def test_word_measurement_uses_all_cells_and_both_ends(monkeypatch):
    monkeypatch.setattr(qa, '_docx_to_pdf_bytes_word', lambda content: None)
    monkeypatch.setattr(qa, '_docx_to_pdf_bytes_libreoffice', lambda content: None)
    monkeypatch.setattr(qa, '_pagination_measurement_copy', lambda content: content)
    class Range:
        Start = 0
        End = 10
        def __init__(self, start, end): self.pages = (start, end)
        @property
        def Duplicate(self): return Range(*self.pages)
        def Collapse(self, direction): self.page = self.pages[0 if direction == 1 else 1]
        def Information(self, code): return self.page
    cells = [SimpleNamespace(RowIndex=3, ColumnIndex=column, Range=Range(1, end))
             for column, end in [(1, 8), (2, 1), (3, 2)]]
    document = SimpleNamespace(Repaginate=lambda: None,
        Tables=lambda index: SimpleNamespace(Range=SimpleNamespace(Cells=cells)))
    monkeypatch.setattr(qa, '_run_with_word_document', lambda content, callback: callback(None, document))
    assert qa.detect_data_row_page_spans(b'docx', total_rows=1) == (qa.DataRowPageSpan(1, 2),)


def test_libreoffice_measurement_removes_vertical_direction_only_from_copy(monkeypatch):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    table = doc.add_table(rows=3, cols=8)
    _repeat_table_header_rows(table)
    table.cell(0, 0).text = 'Header'
    table.cell(2, 0).text = 'Month'
    table.cell(2, 1).text = '19'
    table.cell(2, 2).text = '01–07.03'
    for cell in (*table.rows[0].cells[:1], *table.rows[2].cells[:3]):
        marker = OxmlElement('w:textDirection')
        marker.set(qn('w:val'), 'btLr')
        cell._tc.get_or_add_tcPr().append(marker)
    source = BytesIO()
    doc.save(source)
    original = source.getvalue()

    captured = {}
    monkeypatch.setattr(qa, '_docx_to_pdf_bytes_word', lambda content: None)
    def libreoffice(content):
        captured['copy'] = content
        return b'pdf'
    monkeypatch.setattr(qa, '_docx_to_pdf_bytes_libreoffice', libreoffice)
    monkeypatch.setattr(
        qa,
        '_data_row_page_spans_pdf',
        lambda content, pdf, total_rows: (
            qa.DataRowPageSpan(1, 2, True),
        ),
    )

    assert qa.detect_data_row_page_spans(original, total_rows=1) == (
        qa.DataRowPageSpan(1, 2, True),
    )
    assert Document(BytesIO(original)).tables[0]._tbl.xpath('.//w:textDirection')
    measured = Document(BytesIO(captured['copy'])).tables[0]
    assert measured.rows[0]._tr.xpath('.//w:textDirection')
    assert not measured.rows[2]._tr.xpath('.//w:textDirection')
    assert [cell.text for cell in measured.rows[2].cells[:3]] == [
        'Month', '19', '01–07.03',
    ]


def test_single_page_segment_is_kept_together(monkeypatch):
    from calendar_pedagoga import docx_generation as gen

    def populate(document, *args, **kwargs):
        table = document.add_table(rows=3, cols=8)
        _repeat_table_header_rows(table)
        table.rows[2].cells[0].text = 'Month'
        return table, _columns_for_table(table), ('Month',)

    monkeypatch.setattr(gen, '_load_template', lambda template: Document())
    monkeypatch.setattr(gen, '_populate_calendar_table', populate)
    layout = (
        qa.DataRowPageLayout(
            qa.DataRowPageSpan(1, 1, True),
            (qa.DataRowPageSegment(1, ('Month', '', '', '', '', '', '', '')),),
        ),
    )
    monkeypatch.setattr(qa, 'detect_data_row_page_layout', lambda *args, **kwargs: layout)
    monkeypatch.setattr(
        qa, 'detect_data_row_page_spans',
        lambda *args, **kwargs: (qa.DataRowPageSpan(1, 1, True),),
    )
    result = gen.generate_calendar_docx(
        None, (None,), SimpleNamespace(uses_organization_template=False), '2026–2027',
    )
    doc = Document(BytesIO(result))
    assert doc.tables[0].rows[2]._tr.xpath('./w:trPr/w:cantSplit')
    assert not doc._element.xpath('.//w:pageBreakBefore')


def test_last_row_of_measured_page_is_not_in_vertical_merge(monkeypatch):
    from calendar_pedagoga import docx_generation as gen

    def populate(document, *args, **kwargs):
        table = document.add_table(rows=5, cols=8)
        _repeat_table_header_rows(table)
        for row in table.rows[2:]:
            row.cells[0].text = 'Month'
        return table, _columns_for_table(table), ('Month',) * 3

    monkeypatch.setattr(gen, '_load_template', lambda template: Document())
    monkeypatch.setattr(gen, '_populate_calendar_table', populate)
    layouts = tuple(
        qa.DataRowPageLayout(
            qa.DataRowPageSpan(1, 1, True),
            (qa.DataRowPageSegment(1, ('Month', '', '', '', '', '', '', '')),),
        ) for _ in range(3)
    )
    spans = tuple(qa.DataRowPageSpan(1, 1, True) for _ in range(3))
    monkeypatch.setattr(qa, 'detect_data_row_page_layout', lambda *args, **kwargs: layouts)
    monkeypatch.setattr(qa, 'detect_data_row_page_spans', lambda *args, **kwargs: spans)
    result = gen.generate_calendar_docx(
        None, (None,) * 3, SimpleNamespace(uses_organization_template=False), '2026–2027',
    )
    rows = Document(BytesIO(result)).tables[0].rows[2:]
    assert rows[0]._tr.xpath('.//w:vMerge[@w:val="restart"]')
    assert rows[1]._tr.xpath('.//w:vMerge[not(@w:val)]')
    assert not rows[2]._tr.xpath('.//w:vMerge')


def test_generation_fails_closed_when_physical_segment_still_splits(monkeypatch):
    from calendar_pedagoga import docx_generation as gen

    def populate(document, *args, **kwargs):
        table = document.add_table(rows=4, cols=8)
        _repeat_table_header_rows(table)
        for row in table.rows[2:]:
            row.cells[0].text = 'Month'
        return table, _columns_for_table(table), ('Month', 'Month')

    monkeypatch.setattr(gen, '_load_template', lambda template: Document())
    monkeypatch.setattr(gen, '_populate_calendar_table', populate)
    layouts = tuple(
        qa.DataRowPageLayout(
            qa.DataRowPageSpan(1, 1, True),
            (qa.DataRowPageSegment(1, ('Month', '', '', '', '', '', '', '')),),
        ) for _ in range(2)
    )
    monkeypatch.setattr(qa, 'detect_data_row_page_layout', lambda *args, **kwargs: layouts)
    monkeypatch.setattr(
        qa, 'detect_data_row_page_spans',
        lambda *args, **kwargs: (
            qa.DataRowPageSpan(1, 2, False), qa.DataRowPageSpan(2, 2, True),
        ),
    )
    with pytest.raises(ValueError, match='page-segment'):
        gen.generate_calendar_docx(
            None, (None, None), SimpleNamespace(uses_organization_template=False), '2026–2027',
        )


def test_generation_fails_closed_when_final_merge_reflows_segment(monkeypatch):
    from calendar_pedagoga import docx_generation as gen

    def populate(document, *args, **kwargs):
        table = document.add_table(rows=3, cols=8)
        _repeat_table_header_rows(table)
        table.rows[2].cells[0].text = 'Month'
        table.rows[2].cells[1].text = '1\nDate 1'
        table.rows[2].cells[2].text = 'Content'
        return table, _columns_for_table(table), ('Month',)

    monkeypatch.setattr(gen, '_load_template', lambda template: Document())
    monkeypatch.setattr(gen, '_populate_calendar_table', populate)
    layout = (
        qa.DataRowPageLayout(
            qa.DataRowPageSpan(1, 1, True),
            (qa.DataRowPageSegment(
                1, ('Month', '1\nDate 1', 'Content', '', '', '', '', '')
            ),),
        ),
    )
    monkeypatch.setattr(qa, 'detect_data_row_page_layout', lambda *args, **kwargs: layout)
    measurements = iter((
        (qa.DataRowPageSpan(1, 1, True),),
        (qa.DataRowPageSpan(1, 2, False),),
    ))
    monkeypatch.setattr(
        qa, 'detect_data_row_page_spans', lambda *args, **kwargs: next(measurements)
    )
    with pytest.raises(ValueError, match='финальная merge-структура'):
        gen.generate_calendar_docx(
            None, (None,), SimpleNamespace(uses_organization_template=False), '2026–2027',
        )


@pytest.mark.parametrize('measurement_available', [True, False])
def test_monthly_restores_deleted_merge_anchor_and_respects_new_pages(monkeypatch, measurement_available):
    from calendar_pedagoga import ui
    from calendar_pedagoga.docx_generation import _merge_month_cell_group, _save_document
    doc = Document()
    table = doc.add_table(rows=7, cols=8)
    _repeat_table_header_rows(table)
    for row, week in zip(table.rows[2:], range(5, 10)):
        row.cells[0].text = 'Октябрь'
        row.cells[1].text = f'{week}\nДата {week}'
        row.cells[2].text = f'Topic {week}'
        row.cells[4].text = f'Practice {week}'
    _merge_month_cell_group(table.rows[2:], _columns_for_table(table), 'Октябрь')
    source_xml = doc._element.xml
    annual = _save_document(doc)
    header = [row._tr.xml for row in table.rows[:2]]
    spans = tuple(qa.DataRowPageSpan(page, page, True) for page in (1, 1, 2, 2))
    monkeypatch.setattr(ui, 'detect_data_row_page_spans',
                        lambda *args, **kwargs: spans if measurement_available else None)
    monthly = Document(BytesIO(ui._monthly_plan_docx(annual, '2026–2027', 2026, 10)))
    rows = monthly.tables[0].rows[2:]
    assert [row.cells[1].text for row in rows] == [f'{week}\nДата {week}' for week in range(6, 10)]
    assert [row.cells[0].text for row in rows] == ['Октябрь'] * 4
    assert [row.cells[2].text for row in rows] == [f'Topic {week}' for week in range(6, 10)]
    assert [row._tr.xml for row in monthly.tables[0].rows[:2]] == header
    assert not monthly._element.xpath('.//w:pageBreakBefore')
    # Row 8 starts page 2: its physical month cell must contain its own label.
    assert rows[2]._tr.tc_lst[0].xpath('.//w:t/text()') == ['Октябрь']
    assert doc._element.xml == source_xml


def test_monthly_restores_week_merge_without_losing_multiple_parts(monkeypatch):
    from calendar_pedagoga import ui
    from calendar_pedagoga.docx_generation import _save_document
    doc = Document()
    table = doc.add_table(rows=5, cols=8)
    for row, week, part in zip(table.rows[2:], (7, 7, 8), ('A', 'B', 'C')):
        row.cells[0].text = 'Октябрь'
        row.cells[1].text = str(week)
        row.cells[2].text = part
    table.cell(2, 1).merge(table.cell(3, 1)).text = '7'
    monkeypatch.setattr(ui, 'detect_data_row_page_spans', lambda *args, **kwargs: None)
    monthly = Document(BytesIO(ui._monthly_plan_docx(_save_document(doc), '2026–2027', 2026, 10)))
    rows = monthly.tables[0].rows[2:]
    assert [row.cells[1].text for row in rows] == ['7', '7', '8']
    assert [row.cells[2].text for row in rows] == ['A', 'B', 'C']
    assert not monthly._element.xpath('.//w:vMerge')
