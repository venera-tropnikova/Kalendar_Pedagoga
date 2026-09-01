from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

from docx import Document
import pytest

from calendar_pedagoga.program_parsing import convert_legacy_doc
from calendar_pedagoga.upload_validation import (
    MAX_UPLOAD_BYTES,
    UploadPurpose,
    UploadValidationError,
    validate_upload,
)


REFERENCES = Path(__file__).resolve().parents[1] / "references"


def _docx(paragraphs: tuple[str, ...], table_headers: tuple[str, ...] = ()) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_headers:
        table = document.add_table(rows=1, cols=len(table_headers))
        for cell, value in zip(table.rows[0].cells, table_headers, strict=True):
            cell.text = value
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _normative_docx() -> bytes:
    return _docx(("ПРИКАЗ МИНИСТЕРСТВА", "Нормативный правовой акт"))


def test_valid_utp_and_calendar_template_are_recognized() -> None:
    utp = REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx"
    template = REFERENCES / "Календарный план Образец.docx"
    assert validate_upload(UploadPurpose.UTP, utp.name, utp.read_bytes()).parsed
    assert (
        validate_upload(
            UploadPurpose.CALENDAR_TEMPLATE, template.name, template.read_bytes()
        ).parsed
        is None
    )


@pytest.mark.parametrize(
    ("purpose", "reference_name"),
    [
        (UploadPurpose.UTP, "Календарный план Образец.docx"),
        (UploadPurpose.CALENDAR_TEMPLATE, "УТП КЛЮЧ 2 г. 2ч.docx"),
        (UploadPurpose.PROGRAM, "УТП КЛЮЧ 2 г. 2ч.docx"),
    ],
)
def test_wrong_document_type_is_rejected(
    purpose: UploadPurpose, reference_name: str
) -> None:
    source = REFERENCES / reference_name
    with pytest.raises(UploadValidationError, match="Неправильный тип документа"):
        validate_upload(purpose, source.name, source.read_bytes())


def test_extension_substitution_is_rejected() -> None:
    source = REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx"
    with pytest.raises(UploadValidationError, match="Неправильный формат"):
        validate_upload(UploadPurpose.PROGRAM, "program.doc", source.read_bytes())


def test_empty_file_is_rejected() -> None:
    with pytest.raises(UploadValidationError, match="Файл пустой"):
        validate_upload(UploadPurpose.UTP, "plan.docx", b"")


def test_corrupted_docx_is_rejected() -> None:
    with pytest.raises(UploadValidationError, match="Файл повреждён"):
        validate_upload(UploadPurpose.UTP, "plan.docx", b"PK\x03\x04broken")


def test_oversized_file_is_rejected_before_parsing() -> None:
    oversized = b"PK\x03\x04" + b"0" * (MAX_UPLOAD_BYTES + 1)
    with pytest.raises(UploadValidationError, match="Файл слишком большой"):
        validate_upload(UploadPurpose.UTP, "plan.docx", oversized)


@pytest.mark.parametrize(
    "purpose",
    (UploadPurpose.UTP, UploadPurpose.PROGRAM, UploadPurpose.CALENDAR_TEMPLATE),
)
def test_normative_documents_are_rejected_by_user_uploaders(
    purpose: UploadPurpose,
) -> None:
    with pytest.raises(UploadValidationError, match="Нормативный документ загружать не нужно"):
        validate_upload(purpose, "document.docx", _normative_docx())


def test_legacy_conversion_always_removes_temporary_directory(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output_bytes = _docx(("Образовательная программа",))
    observed_temp_dir: Path | None = None

    def fake_run(command, **_kwargs):
        nonlocal observed_temp_dir
        output_dir = Path(command[command.index("--outdir") + 1])
        observed_temp_dir = output_dir
        (output_dir / "program.docx").write_bytes(output_bytes)
        return SimpleNamespace(returncode=0, stderr="", stdout="")

    monkeypatch.setattr("calendar_pedagoga.program_parsing.subprocess.run", fake_run)
    result = convert_legacy_doc(b"legacy", soffice=Path("soffice.exe"))

    assert result == output_bytes
    assert observed_temp_dir is not None
    assert not observed_temp_dir.exists()
