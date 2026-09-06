from io import BytesIO
from pathlib import Path
from functools import lru_cache

import pytest
from calendar_pedagoga.docx_generation import (
    PRINT_TOP_MARGIN_CM,
    STANDARD_GROUP_SPACE_AFTER_PT,
    STANDARD_TABLE_FONT_FAMILY,
    STANDARD_TEMPLATE_PATH,
    _fill_organization_header_paragraph,
    _group_class_line,
    _program_header_line,
    _resolve_header_from_rows,
    _write_document_header,
    build_output_filename,
    generate_calendar_docx,
)
from calendar_pedagoga.docx_qa import (
    QASeverity,
    _month_cell_is_continuation,
    detect_data_row_indices_by_page,
    has_blocking_qa_issues,
    validate_calendar_docx,
    verify_month_labels_by_page,
)
from calendar_pedagoga.lesson_resolution import resolve_lesson_content
from calendar_pedagoga.organization_template import select_calendar_template
from calendar_pedagoga.parsing import parse_utp
from calendar_pedagoga.program_parsing import parse_program
from calendar_pedagoga.resolve_utp import resolve_utp
from calendar_pedagoga.upload_validation import UploadPurpose, validate_upload
from calendar_pedagoga.content_generation import build_content_model
from calendar_pedagoga.lesson_content import build_lesson_content
from calendar_pedagoga.scheduling import build_schedule
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REFERENCES = Path(__file__).resolve().parents[1] / "references"


@lru_cache(maxsize=1)
def _key_docx() -> bytes:
    utp_path = REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx"
    program_path = REFERENCES / "Программа КЛЮЧ.DOC"
    utp = parse_utp(utp_path)
    program = parse_program(program_path.read_bytes(), program_path.name, study_year=2)
    content = build_content_model(build_schedule(utp), utp, program, utp_path.name)
    resolved = resolve_lesson_content(build_lesson_content(content))
    return generate_calendar_docx(
        utp,
        resolved,
        select_calendar_template(),
        "2026–2027",
    )


def test_standard_template_exists() -> None:
    assert STANDARD_TEMPLATE_PATH.is_file()


def test_generated_table_marks_header_rows_for_word_repeat() -> None:
    document = Document(BytesIO(_key_docx()))
    table = document.tables[0]
    header_tag = qn("w:tblHeader")
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.find(qn("w:trPr"))
        marker = tr_pr.find(header_tag) if tr_pr is not None else None
        if index < 2:
            assert marker is not None, f"header row {index} without tblHeader"
            assert not marker.attrib, f"header row {index} must be <w:tblHeader/>"
        else:
            assert marker is None, f"data row {index} must not repeat as header"


def _explicit_run_fonts(run) -> tuple[str | None, ...]:
    run_properties = run._r.find(qn("w:rPr"))
    fonts = run_properties.find(qn("w:rFonts")) if run_properties is not None else None
    if fonts is None:
        return (None, None, None, None)
    return tuple(
        fonts.get(qn(f"w:{attribute}"))
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs")
    )


def test_standard_table_uses_one_explicit_font_family() -> None:
    table = Document(BytesIO(_key_docx())).tables[0]
    header_runs = [
        run
        for row in table.rows[:2]
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    ]
    body_runs = [
        run
        for row in table.rows[2:]
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    ]
    expected = (STANDARD_TABLE_FONT_FAMILY,) * 4

    assert header_runs
    assert body_runs
    assert all(_explicit_run_fonts(run) == expected for run in header_runs)
    assert all(_explicit_run_fonts(run) == expected for run in body_runs)
    assert all(run.bold is True for run in header_runs)
    assert all(run.bold is False for run in body_runs)
    assert {run.font.size.pt for run in header_runs if run.font.size} == {12.0}
    assert {run.font.size.pt for run in body_runs if run.font.size} == {12.0}


def _key_docx_for_year(academic_year: str) -> bytes:
    utp_path = REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx"
    program_path = REFERENCES / "Программа КЛЮЧ.DOC"
    utp = parse_utp(utp_path)
    program = parse_program(program_path.read_bytes(), program_path.name, study_year=2)
    content = build_content_model(build_schedule(utp, academic_year), utp, program, utp_path.name)
    resolved = resolve_lesson_content(build_lesson_content(content))
    return generate_calendar_docx(
        utp,
        resolved,
        select_calendar_template(),
        academic_year,
    )


def _week_cells(content: bytes) -> list[str]:
    table = Document(BytesIO(content)).tables[0]
    return [row.cells[1].text for row in table.rows[2:]]


def test_key_generation_produces_36_data_rows() -> None:
    content = _key_docx()
    document = Document(BytesIO(content))
    assert document.paragraphs[0].text == "Календарный план"
    assert document.paragraphs[1].text == "«КЛЮЧ» — 2 год обучения (2 часа в неделю)"
    assert document.paragraphs[2].text == "2026–2027 учебный год"
    assert document.paragraphs[3].text == "Группа № ___________ (Класс _________)"
    table = document.tables[0]
    assert len(table.rows) == 38
    assert len(table.columns) >= 8


def test_standard_docx_keeps_2026_dates_and_builds_2027_without_gap() -> None:
    year_2026 = _key_docx_for_year("2026–2027")
    year_2027 = _key_docx_for_year("2027–2028")
    doc_2026 = Document(BytesIO(year_2026))
    doc_2027 = Document(BytesIO(year_2027))
    assert doc_2026.paragraphs[2].text == "2026–2027 учебный год"
    assert doc_2027.paragraphs[2].text == "2027–2028 учебный год"
    weeks_2026 = _week_cells(year_2026)
    weeks_2027 = _week_cells(year_2027)
    assert any("01–06.09" in cell for cell in weeks_2026)
    assert any("28–30.12" in cell for cell in weeks_2026)
    assert any("11–17.01" in cell for cell in weeks_2026)
    assert any("01–05.09" in cell for cell in weeks_2027)
    assert all("28–30.12" not in cell for cell in weeks_2027)
    assert all("11–17.01" not in cell for cell in weeks_2027)
    assert not has_blocking_qa_issues(validate_calendar_docx(year_2026, expected_weeks=36))
    assert not has_blocking_qa_issues(validate_calendar_docx(year_2027, expected_weeks=36))


def test_key_generation_passes_structural_qa() -> None:
    issues = validate_calendar_docx(_key_docx(), expected_weeks=36)
    assert not has_blocking_qa_issues(issues)


def _structural_qa_heading_fixture(title: str = "Календарный план") -> Document:
    document = Document()
    document.add_paragraph(title)
    document.add_paragraph("Учебная программа")
    document.add_paragraph("Группа № 1")
    table = document.add_table(rows=38, cols=8)
    table.rows[0].cells[0].text = (
        "Месяц Неделя Теоретические занятия Практические занятия "
        "Тип занятия Планируемый результат Вид контроля"
    )
    for number, row in enumerate(table.rows[2:], start=1):
        row.cells[0].text = "Сентябрь"
        row.cells[1].text = f"{number}\n01–07.09"
        row.cells[2].text = "Тема"
    return document


def test_structural_qa_accepts_semantic_heading_after_blank_template_paragraph() -> None:
    document = _structural_qa_heading_fixture()
    document.paragraphs[0]._p.addprevious(OxmlElement("w:p"))
    output = BytesIO()
    document.save(output)

    issues = validate_calendar_docx(output.getvalue(), expected_weeks=36)
    assert not has_blocking_qa_issues(issues)


def test_structural_qa_does_not_accept_calendar_heading_after_table() -> None:
    document = _structural_qa_heading_fixture("Рабочий документ")
    document.add_paragraph("Календарный план")
    output = BytesIO()
    document.save(output)

    issues = validate_calendar_docx(output.getvalue(), expected_weeks=36)
    assert any(
        issue.severity is QASeverity.ERROR
        and "заголовок календарного плана" in issue.message.casefold()
        for issue in issues
    )


def _assert_print_safe_margins(document, source) -> None:
    generated = document.sections[0]
    template = source.sections[0]
    assert abs(generated.top_margin.cm - PRINT_TOP_MARGIN_CM) < 0.02
    assert abs(generated.bottom_margin.cm - template.bottom_margin.cm) < 0.02
    assert abs(generated.left_margin.cm - template.left_margin.cm) < 0.02
    assert abs(generated.right_margin.cm - template.right_margin.cm) < 0.02


def test_key_generation_sets_print_safe_top_margin() -> None:
    content = _key_docx()
    document = Document(BytesIO(content))
    source = Document(str(STANDARD_TEMPLATE_PATH))
    _assert_print_safe_margins(document, source)
    assert "Группа Нет" not in " ".join(paragraph.text for paragraph in document.paragraphs[:6])


def test_header_line_uses_program_title_and_filename_year() -> None:
    utp = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    assert _program_header_line(utp) == "«КЛЮЧ» — 2 год обучения (2 часа в неделю)"
    assert _group_class_line() == "Группа № ___________ (Класс _________)"
    assert _group_class_line("12", "5Б") == "Группа № 12 (Класс 5Б)"


def test_tour_guides_header_uses_program_and_filename_year() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    validated_program = validate_upload(
        UploadPurpose.PROGRAM,
        program_path.name,
        program_path.read_bytes(),
    )
    utp = resolve_utp(None, validated_program)
    program = parse_program(program_path.read_bytes(), program_path.name, study_year=1)
    line = _program_header_line(
        utp,
        program_title=program.title,
        study_year_hints=(f"УТП из файла «{program_path.name}»", program_path.name),
    )
    assert line == "«Туристы-проводники» — 1 год обучения (2 часа в неделю)"
    title, hints = _resolve_header_from_rows(
        resolve_lesson_content(
            build_lesson_content(
                build_content_model(
                    build_schedule(utp),
                    utp,
                    program,
                    f"УТП из файла «{program_path.name}»",
                )
            )
        ),
        program_title=None,
        study_year_hints=(),
    )
    assert _program_header_line(
        utp,
        program_title=title,
        study_year_hints=hints,
    ) == "«Туристы-проводники» — 1 год обучения (2 часа в неделю)"


def test_document_header_writes_year_and_optional_group() -> None:
    document = Document(str(STANDARD_TEMPLATE_PATH))
    utp = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    _write_document_header(
        document,
        utp,
        academic_year="2026–2027",
        program_title="КЛЮЧ",
        group_number="5",
        class_name="7А",
    )
    assert document.paragraphs[0].text == "Календарный план"
    assert document.paragraphs[1].text == "«КЛЮЧ» — 2 год обучения (2 часа в неделю)"
    assert document.paragraphs[2].text == "2026–2027 учебный год"
    assert document.paragraphs[3].text == "Группа № 5 (Класс 7А)"


def test_output_filename_uses_program_and_year() -> None:
    utp = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    filename = build_output_filename(utp, "2026–2027")
    assert filename.endswith(".docx")
    assert "2026-2027" in filename


def test_tour_guides_without_program_generates_valid_empty_content_docx() -> None:
    utp_path = REFERENCES / "УТП ТП 3г. 2ч.docx"
    utp = parse_utp(utp_path)
    content = build_content_model(build_schedule(utp), utp, None, utp_path.name)
    resolved = resolve_lesson_content(build_lesson_content(content))
    docx_bytes = generate_calendar_docx(
        utp,
        resolved,
        select_calendar_template(),
        "2026–2027",
    )
    issues = validate_calendar_docx(docx_bytes, expected_weeks=36)
    assert not has_blocking_qa_issues(issues)
    document = Document(BytesIO(docx_bytes))
    first_page_rows = detect_data_row_indices_by_page(docx_bytes, total_rows=36)
    assert first_page_rows is not None and first_page_rows[0]
    first_page_labels = [
        document.tables[0].rows[2 + index].cells[0].text.strip()
        for index in first_page_rows[0]
    ]
    assert "Сентябрь" in first_page_labels


def test_key_reference_calendar_matches_generated_structure() -> None:
    """Structural QA: сгенерированный календарь совпадает с эталоном по размерности таблицы."""
    generated = Document(BytesIO(_key_docx()))
    reference = Document(REFERENCES / "Календарный_план_КЛЮЧ_2026-2027_с_датами.docx")
    gen_table = generated.tables[0]
    ref_table = reference.tables[0]
    assert len(gen_table.rows) == len(ref_table.rows)
    assert len(gen_table.columns) >= 8
    assert generated.paragraphs[0].text == reference.paragraphs[0].text


def test_key_generation_passes_visual_qa_all_pages() -> None:
    from calendar_pedagoga.docx_qa import validate_calendar_docx_visual, has_blocking_qa_issues
    from calendar_pedagoga.program_parsing import find_libreoffice

    if find_libreoffice() is None:
        pytest.skip("LibreOffice недоступен для visual QA")

    issues = validate_calendar_docx_visual(_key_docx())
    assert not has_blocking_qa_issues(issues)


def test_qa_detects_missing_weeks() -> None:
    issues = validate_calendar_docx(_key_docx(), expected_weeks=40)
    assert any(issue.severity is QASeverity.ERROR for issue in issues)


def test_visual_qa_checks_all_data_rows_have_week_and_month() -> None:
    content = _key_docx()
    document = Document(BytesIO(content))
    for index, row in enumerate(document.tables[0].rows[2:], start=1):
        month_cell = row.cells[0]
        if not month_cell.text.strip():
            assert _month_cell_is_continuation(month_cell), f"month missing row {index}"
        assert row.cells[1].text.strip(), f"week missing row {index}"
        assert str(index) in row.cells[1].text.splitlines()[0]


def test_key_docx_does_not_truncate_source_with_ellipsis() -> None:
    document = Document(BytesIO(_key_docx()))
    table = document.tables[0]
    week3 = table.rows[4].cells[2].text  # week 3 theory
    week7 = table.rows[8].cells[4].text  # week 7 practice
    week16 = table.rows[17].cells[4].text  # week 16 practice
    for text in (week3, week7, week16):
        assert "…" not in text
        assert "..." not in text
    assert "строительства города" in week3
    assert "Праздник курая" in week7
    assert "Найди середину" in week16
    assert "Мой город" in week3
    assert "История родного края" in week7
    assert "Ориентирование" in week16


def test_key_docx_fills_type_result_control_and_keeps_mark_empty() -> None:
    document = Document(BytesIO(_key_docx()))
    data_rows = document.tables[0].rows[2:]
    assert len(data_rows) == 36
    for index, row in enumerate(data_rows, start=1):
        cells = [cell.text.strip() for cell in row.cells]
        assert cells[5], f"lesson type empty week {index}"
        assert cells[6], f"planned result empty week {index}"
        assert cells[7], f"assessment empty week {index}"
        assert not cells[6].casefold().startswith("учащийся сможет изучить тему")
        assert cells[3] == ""


def test_organization_template_preserves_vertical_columns_and_merges_months() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    template_path = REFERENCES / "Календарный план.docx"
    program_upload = parse_program(program_path.read_bytes(), program_path.name, study_year=1)
    validated_program = validate_upload(
        UploadPurpose.PROGRAM,
        program_path.name,
        program_path.read_bytes(),
    )
    utp = resolve_utp(None, validated_program)
    schedule = build_schedule(utp)
    content = build_content_model(schedule, utp, program_upload, program_path.name)
    resolved = resolve_lesson_content(build_lesson_content(content))
    generated = generate_calendar_docx(
        utp,
        resolved,
        select_calendar_template(template_path.name, template_path.read_bytes()),
        "2026–2027",
    )

    table = Document(BytesIO(generated)).tables[0]
    data_rows = table.rows[2:]
    expected_months = tuple(week.month for week in schedule.weeks)
    active_month = None
    saw_restart = False
    for index, row in enumerate(data_rows):
        raw_cells = row._tr.tc_lst
        for column in (0, 1):
            direction = raw_cells[column].tcPr.find(qn("w:textDirection"))
            assert direction is not None
            assert direction.get(qn("w:val")) == "btLr"
            vertical_alignment = raw_cells[column].tcPr.find(qn("w:vAlign"))
            assert vertical_alignment is not None
            assert vertical_alignment.get(qn("w:val")) == "center"

        month_merge = raw_cells[0].tcPr.find(qn("w:vMerge"))
        month_xml_text = "".join(raw_cells[0].xpath(".//w:t/text()")).strip()
        if month_merge is None:
            assert month_xml_text == expected_months[index]
            active_month = None
        elif month_merge.get(qn("w:val")) == "restart":
            assert month_xml_text == expected_months[index]
            active_month = expected_months[index]
            saw_restart = True
        else:
            assert month_xml_text == ""
            assert active_month == expected_months[index]
    assert saw_restart, "at least one page-safe month segment must be merged"


def test_tour_guides_month_labels_visible_on_each_page_segment() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    template_path = REFERENCES / "Календарный план.docx"
    program_upload = parse_program(program_path.read_bytes(), program_path.name, study_year=1)
    validated_program = validate_upload(
        UploadPurpose.PROGRAM,
        program_path.name,
        program_path.read_bytes(),
    )
    utp = resolve_utp(None, validated_program)
    schedule = build_schedule(utp)
    content = build_content_model(schedule, utp, program_upload, program_path.name)
    resolved = resolve_lesson_content(build_lesson_content(content))
    generated = generate_calendar_docx(
        utp,
        resolved,
        select_calendar_template(template_path.name, template_path.read_bytes()),
        "2026–2027",
    )

    expected_months = tuple(week.month for week in schedule.weeks)
    rows_by_page = detect_data_row_indices_by_page(generated, total_rows=len(expected_months))
    assert rows_by_page is not None

    issues = verify_month_labels_by_page(generated, months=expected_months)
    assert not issues, "; ".join(issues)


def _run_font(paragraph):
    run = paragraph.runs[0]
    size = run.font.size.pt if run.font.size else None
    name = run.font.name
    if name is None:
        rpr = run._r.find(qn("w:rPr"))
        rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
        if rfonts is not None:
            name = rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))
    return name, size


def test_organization_header_replaces_values_without_academic_year_line() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    tourists = resolve_utp(
        None,
        validate_upload(
            UploadPurpose.PROGRAM,
            program_path.name,
            program_path.read_bytes(),
        ),
    )
    filled = _fill_organization_header_paragraph(
        "« Туристы проводники »  1 год обучения  ( 72 чса )",
        tourists,
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        group_number=None,
        class_name=None,
    )
    assert "Туристы-проводники" in filled
    assert "1 год обучения" in filled
    assert "72" in filled
    assert "учебный год" not in filled

    key = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    sample = _fill_organization_header_paragraph(
        "« название программы »  ____  г.об. ( количество часов в неделю)",
        key,
        program_title="КЛЮЧ",
        study_year_hints=("УТП КЛЮЧ 2 г. 2ч.docx",),
        group_number=None,
        class_name=None,
    )
    assert sample.startswith("«КЛЮЧ»")
    assert "2 г.об." in sample
    assert "2 часа в неделю" in sample

    filled_group = _fill_organization_header_paragraph(
        "Группа № ___________ (Класс _________)",
        tourists,
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        group_number="3",
        class_name="5А",
    )
    assert filled_group == "\tГруппа № 3 (Класс 5А)"
    empty_group = _fill_organization_header_paragraph(
        "Группа № ___________ (Класс _________)",
        tourists,
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        group_number="",
        class_name="",
    )
    assert empty_group == "\tГруппа № ___________ (Класс _________)"
    assert "Нет" not in empty_group

    empty_teacher = _fill_organization_header_paragraph(
        "Группа № ___________ (Класс _________)",
        tourists,
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        group_number="",
        class_name="",
        teacher_name="",
    )
    assert empty_teacher == "\tГруппа № ___________ (Класс _________)"

    filled_teacher = _fill_organization_header_paragraph(
        "Группа № ___________ (Класс _________)",
        tourists,
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        group_number="",
        class_name="",
        teacher_name="Иванов И.И.",
    )
    assert filled_teacher == "\tГруппа № ___________ (Класс _________)\tИванов И.И."

    preserved_tail = _fill_organization_header_paragraph(
        "Группа № ___________ (Класс _________)                                            Саранцева И.М.",
        tourists,
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        group_number="3",
        class_name="5А",
        teacher_name="",
    )
    assert preserved_tail == "\tГруппа № 3 (Класс 5А)"
    assert "Саранцева И.М." not in preserved_tail


def test_standard_header_adds_teacher_name_to_group_line() -> None:
    document = Document(str(STANDARD_TEMPLATE_PATH))
    source_count = len(document.paragraphs)
    utp = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    _write_document_header(
        document,
        utp,
        academic_year="2026–2027",
        program_title="КЛЮЧ",
        group_number="5",
        class_name="7А",
        teacher_name="Иванов И.И.",
    )
    assert len(document.paragraphs) == source_count
    assert document.paragraphs[2].text == "2026–2027 учебный год"
    group = document.paragraphs[3]
    assert group.text == "\tГруппа № 5 (Класс 7А)\tИванов И.И."
    assert group.alignment == 0  # LEFT: позицию задают center/right tabs
    assert group.paragraph_format.space_after.pt == STANDARD_GROUP_SPACE_AFTER_PT
    tabs = group._p.find(qn("w:pPr")).find(qn("w:tabs"))
    assert tabs is not None
    assert [tab.get(qn("w:val")) for tab in tabs.findall(qn("w:tab"))] == [
        "center",
        "right",
    ]
    assert all(paragraph.text.strip() != "Иванов И.И." for paragraph in document.paragraphs)


def test_standard_header_keeps_empty_teacher_behavior() -> None:
    document = Document(str(STANDARD_TEMPLATE_PATH))
    source_count = len(document.paragraphs)
    source_alignment = document.paragraphs[3].alignment
    source_properties = document.paragraphs[3]._p.find(qn("w:pPr"))
    source_tabs = source_properties.find(qn("w:tabs")) if source_properties is not None else None
    source_tab_values = (
        [tab.get(qn("w:val")) for tab in source_tabs.findall(qn("w:tab"))]
        if source_tabs is not None
        else []
    )
    utp = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    _write_document_header(
        document,
        utp,
        academic_year="2026–2027",
        program_title="КЛЮЧ",
        group_number="5",
        class_name="7А",
        teacher_name="",
    )
    assert len(document.paragraphs) == source_count
    assert document.paragraphs[3].text == "Группа № 5 (Класс 7А)"
    assert document.paragraphs[3].alignment == source_alignment
    assert (
        document.paragraphs[3].paragraph_format.space_after.pt
        == STANDARD_GROUP_SPACE_AFTER_PT
    )
    paragraph_properties = document.paragraphs[3]._p.find(qn("w:pPr"))
    tabs = paragraph_properties.find(qn("w:tabs")) if paragraph_properties is not None else None
    tab_values = (
        [tab.get(qn("w:val")) for tab in tabs.findall(qn("w:tab"))]
        if tabs is not None
        else []
    )
    assert tab_values == source_tab_values


def test_standard_header_uses_explicit_times_new_roman() -> None:
    document = Document(str(STANDARD_TEMPLATE_PATH))
    utp = parse_utp(REFERENCES / "УТП КЛЮЧ 2 г. 2ч.docx")
    _write_document_header(
        document,
        utp,
        academic_year="2026–2027",
        program_title="КЛЮЧ",
        group_number="5",
        class_name="7А",
        teacher_name="Иванов И.И.",
    )
    expected = (STANDARD_TABLE_FONT_FAMILY,) * 4
    runs = [
        run
        for paragraph in document.paragraphs[:4]
        for run in paragraph.runs
        if run.text.strip()
    ]
    assert runs
    assert all(_explicit_run_fonts(run) == expected for run in runs)
    assert {run.font.size.pt for run in runs if run.font.size} == {12.0}


def test_organization_template_keeps_visual_header_and_times_new_roman() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    template_path = REFERENCES / "Календарный план.docx"
    program_upload = parse_program(program_path.read_bytes(), program_path.name, study_year=1)
    validated_program = validate_upload(
        UploadPurpose.PROGRAM,
        program_path.name,
        program_path.read_bytes(),
    )
    utp = resolve_utp(None, validated_program)
    content = build_content_model(build_schedule(utp), utp, program_upload, program_path.name)
    resolved = resolve_lesson_content(build_lesson_content(content))
    generated = generate_calendar_docx(
        utp,
        resolved,
        select_calendar_template(template_path.name, template_path.read_bytes()),
        "2026–2027",
        program_title=program_upload.title,
        study_year_hints=(program_path.name,),
    )

    source = Document(str(template_path))
    document = Document(BytesIO(generated))
    _assert_print_safe_margins(document, source)
    header_texts = [paragraph.text for paragraph in document.paragraphs[:4]]
    assert header_texts[0] == "Календарный план"
    assert "Туристы-проводники" in header_texts[1]
    assert "1 год обучения" in header_texts[1]
    assert "72" in header_texts[1]
    assert header_texts[2] == "\tГруппа № ___________ (Класс _________)"
    assert all("учебный год" not in text for text in header_texts)

    for paragraph in document.paragraphs[:3]:
        name, size = _run_font(paragraph)
        assert name == "Times New Roman"
        assert size == 12.0

    source_headers = [
        "".join(cell.xpath(".//w:t/text()")).strip()
        for cell in source.tables[0].rows[1]._tr.tc_lst
    ]
    generated_headers = [
        "".join(cell.xpath(".//w:t/text()")).strip()
        for cell in document.tables[0].rows[1]._tr.tc_lst
    ]
    assert generated_headers == source_headers

    source_grid = [
        width.get(qn("w:w"))
        for width in source.tables[0]._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    generated_grid = [
        width.get(qn("w:w"))
        for width in document.tables[0]._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    assert generated_grid == source_grid

    data_run = document.tables[0].rows[2].cells[2].paragraphs[0].runs[0]
    data_size = data_run.font.size.pt if data_run.font.size else None
    assert data_size == 12.0

    header_height = source.tables[0].rows[1]._tr.find(qn("w:trPr")).find(qn("w:trHeight"))
    generated_header_height = document.tables[0].rows[1]._tr.find(qn("w:trPr")).find(
        qn("w:trHeight")
    )
    assert header_height is not None
    assert generated_header_height is not None
    assert generated_header_height.get(qn("w:val")) == header_height.get(qn("w:val"))

    for row in document.tables[0].rows[2:]:
        row_properties = row._tr.find(qn("w:trPr"))
        height = row_properties.find(qn("w:trHeight")) if row_properties is not None else None
        assert height is not None, "vertical identifiers require a safe minimum height"
        assert height.get(qn("w:hRule")) == "atLeast"
        assert height.get(qn("w:val")) != header_height.get(qn("w:val"))
        for paragraph in row.cells[2].paragraphs:
            paragraph_properties = paragraph._p.find(qn("w:pPr"))
            assert paragraph_properties is not None
            indent = paragraph_properties.find(qn("w:ind"))
            tabs = paragraph_properties.find(qn("w:tabs"))
            spacing = paragraph_properties.find(qn("w:spacing"))
            assert indent is None, "data cells must not inherit header firstLine indent"
            assert tabs is None, "data cells must not inherit header tabs"
            assert spacing is not None, "data cells must override docDefaults spacing"
            assert spacing.get(qn("w:before")) == "0"
            assert spacing.get(qn("w:after")) == "0"
            assert spacing.get(qn("w:line")) == "240"
            assert spacing.get(qn("w:lineRule")) == "auto"
            assert paragraph_properties.find(qn("w:jc")) is None

    theory_cell = document.tables[0].rows[2].cells[2]
    assert len(theory_cell.paragraphs) == 1


def test_organization_docx_appends_teacher_name_without_new_paragraph() -> None:
    program_path = REFERENCES / "Программа ТУРИСТЫ-ПРОВОДНИКИ 1 г.docx"
    template_path = REFERENCES / "Календарный план.docx"
    tourists = resolve_utp(
        None,
        validate_upload(
            UploadPurpose.PROGRAM,
            program_path.name,
            program_path.read_bytes(),
        ),
    )
    source = Document(str(template_path))
    empty = Document(str(template_path))
    _write_document_header(
        empty,
        tourists,
        academic_year="2026–2027",
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        teacher_name="",
        uses_organization_template=True,
    )
    assert len(empty.paragraphs) == len(source.paragraphs)
    assert empty.paragraphs[2].text == "\tГруппа № ___________ (Класс _________)"
    assert empty.paragraphs[2].alignment == 0  # LEFT
    assert all("учебный год" not in paragraph.text for paragraph in empty.paragraphs[:4])
    empty_tabs = empty.paragraphs[2]._p.find(qn("w:pPr")).find(qn("w:tabs"))
    empty_vals = [tab.get(qn("w:val")) for tab in empty_tabs.findall(qn("w:tab"))]
    assert empty_vals == ["center", "right"]

    filled = Document(str(template_path))
    _write_document_header(
        filled,
        tourists,
        academic_year="2026–2027",
        program_title="Туристы-проводники",
        study_year_hints=(program_path.name,),
        teacher_name="Иванов И.И.",
        uses_organization_template=True,
    )
    assert len(filled.paragraphs) == len(source.paragraphs)
    group = filled.paragraphs[2]
    assert group.text == "\tГруппа № ___________ (Класс _________)\tИванов И.И."
    assert group.alignment == 0  # LEFT
    tabs = group._p.find(qn("w:pPr")).find(qn("w:tabs"))
    assert tabs is not None
    tab_vals = [tab.get(qn("w:val")) for tab in tabs.findall(qn("w:tab"))]
    assert tab_vals == ["center", "right"]
    assert all(paragraph.text.strip() != "Иванов И.И." for paragraph in filled.paragraphs)
    assert all("учебный год" not in paragraph.text for paragraph in filled.paragraphs[:4])
