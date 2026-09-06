from datetime import date
import hashlib
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from streamlit.testing.v1 import AppTest

from calendar_pedagoga import ui
from calendar_pedagoga.academic_year import default_academic_year_start, format_academic_year
from calendar_pedagoga.practice_slots import SLOT_CONTINUE_WARNING, SLOT_PACK_WARNING
from calendar_pedagoga.ui import _teacher_generation_warnings


APP_PATH = Path(__file__).resolve().parents[1] / "app.py"
REFERENCES = Path(__file__).resolve().parents[1] / "references"
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _reference_named(*parts: str) -> Path:
    lowered = tuple(part.casefold() for part in parts)
    for path in REFERENCES.iterdir():
        name = path.name.casefold()
        if all(part in name for part in lowered):
            return path
    raise FileNotFoundError(" ".join(parts))


def _program_file() -> Path:
    return _reference_named("туристы-проводники", "1")


def _utp_file() -> Path:
    return _reference_named("утп", "ключ")


def _default_year() -> str:
    return format_academic_year(default_academic_year_start())


def _template_file() -> Path:
    exact = REFERENCES / "Календарный план.docx"
    if exact.exists():
        return exact
    return _reference_named("календарный план.docx")


def _non_utp_docx() -> bytes:
    document = Document()
    document.add_paragraph("Это не учебно-тематический план.")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _upload(app: AppTest, index: int, path: Path) -> None:
    app.get("file_uploader")[index].set_value((path.name, path.read_bytes(), DOCX_MIME))


def _clear_buttons(app: AppTest):
    return [button for button in app.button if button.label == "×"]


def _page_text(app: AppTest) -> str:
    return " ".join(item.value or "" for item in app.markdown)


def _analysis_ready(app: AppTest) -> bool:
    return "analysis_ready" in app.session_state and bool(app.session_state["analysis_ready"])


def _check_button(app: AppTest):
    return next(button for button in app.button if button.label == "Проверить документы")


def test_initial_screen_contains_required_controls() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=10).run()

    assert not app.exception
    assert app.title[0].value == "Календарь педагога"
    uploaders = app.get("file_uploader")
    assert len(uploaders) == 3
    assert [uploader.label for uploader in uploaders] == [
        "Загрузите образовательную программу",
        "Загрузите УТП",
        "Шаблон календарного плана вашей организации",
    ]
    assert [uploader.help for uploader in uploaders] == [
        "Программа — образовательная программа, DOC/DOCX, до 10 МБ.",
        "УТП — учебно-тематический план, DOCX, до 10 МБ.",
        "Шаблон — только образец календарного плана организации, DOCX, до 10 МБ.",
    ]
    assert app.number_input[0].label == "Начало учебного года"
    assert int(app.number_input[0].value) == default_academic_year_start()
    assert "2026–2027 / 2027–2028 / 2028–2029" not in _page_text(app)
    assert [item.label for item in app.text_input] == ["Группа №", "Класс", "ФИО педагога"]
    assert all(item.value in {"", None} for item in app.text_input)
    assert not any("ИИ" in (item.label or "") for item in getattr(app, "checkbox", []))
    assert "Дополнить содержание с помощью ИИ" not in _page_text(app)
    assert "Группа Нет" not in _page_text(app)
    assert any(button.label == "Проверить документы" for button in app.button)
    assert any(button.label == "Открыть календарь" for button in app.button)
    assert "1. Документы" in _page_text(app)
    assert "2. Сведения для плана" in _page_text(app)
    assert f"Календарь {_default_year()} учебного года" in _page_text(app)
    assert "Недели №1–36 соответствуют строкам календарного плана" in _page_text(app)
    notes = " ".join(item.value or "" for item in app.markdown)
    assert "Загрузите отдельно, только если УТП находится в другом файле" in notes
    assert "Если есть образец вашей организации — загрузите его; иначе используем стандартный" in notes


def test_check_requires_program() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=10).run()
    _check_button(app).click().run()

    assert not app.exception
    assert app.error[0].value == "Загрузите программу обучения."
    assert not _clear_buttons(app)


def test_clear_template_keeps_program() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=20).run()
    program = _program_file()
    template = _template_file()
    _upload(app, 0, program)
    _upload(app, 2, template)
    app.run()

    assert [button.label for button in _clear_buttons(app)] == ["×", "×"]
    _clear_buttons(app)[1].click().run()

    assert not app.exception
    names = " ".join(item.value or "" for item in app.markdown)
    assert program.name in names
    assert template.name not in names
    assert len(_clear_buttons(app)) == 1
    assert int(app.number_input[0].value) == default_academic_year_start()


def test_clear_program_resets_analysis_but_keeps_other_files() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=30).run()
    program = _program_file()
    template = _template_file()
    _upload(app, 0, program)
    _upload(app, 2, template)
    app.run()
    _check_button(app).click().run()

    assert _analysis_ready(app) is True
    assert "Документы проверены" in _page_text(app)

    template_nonce = app.session_state["upload_nonce_template"]
    _clear_buttons(app)[0].click().run()

    assert _analysis_ready(app) is False
    assert "calendar_context" not in app.session_state
    assert "Документы проверены" not in _page_text(app)
    assert app.session_state["upload_nonce_program"] == 1
    assert app.session_state["upload_nonce_template"] == template_nonce
    assert int(app.number_input[0].value) == default_academic_year_start()


def test_clear_wrong_utp_keeps_program_and_template() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=20).run()
    program = _program_file()
    template = _template_file()
    _upload(app, 0, program)
    app.get("file_uploader")[1].set_value(("wrong-utp.docx", _non_utp_docx(), DOCX_MIME))
    _upload(app, 2, template)
    app.run()

    assert len(_clear_buttons(app)) == 3
    program_nonce = app.session_state["upload_nonce_program"]
    template_nonce = app.session_state["upload_nonce_template"]
    _clear_buttons(app)[1].click().run()

    assert app.session_state["upload_nonce_utp"] == 1
    assert app.session_state["upload_nonce_program"] == program_nonce
    assert app.session_state["upload_nonce_template"] == template_nonce
    assert "wrong-utp.docx" not in _page_text(app)


def test_cleared_slot_accepts_new_upload() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=20).run()
    program = _program_file()
    template = _template_file()
    _upload(app, 0, program)
    _upload(app, 2, template)
    app.run()
    _clear_buttons(app)[1].click().run()

    _upload(app, 2, template)
    app.run()

    names = " ".join(item.value or "" for item in app.markdown)
    assert program.name in names
    assert template.name in names
    assert len(_clear_buttons(app)) == 2


def test_utp_year_is_suggested_after_upload() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=20).run()
    _upload(app, 0, _program_file())
    _upload(app, 1, _utp_file())
    app.run()

    assert not app.exception
    assert int(app.number_input[0].value) == 2026
    notices = " ".join(
        [
            _page_text(app),
            " ".join(item.value or "" for item in getattr(app, "info", [])),
            " ".join(item.value or "" for item in getattr(app, "caption", [])),
        ]
    )
    assert "2026–2027" in notices
    assert "УТП" in notices
    assert "2026–2027 / 2027–2028" not in notices


def test_analysis_screen_shows_study_year_from_program_filename() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=30).run()
    _upload(app, 0, _program_file())
    app.run()
    _check_button(app).click().run()

    text = _page_text(app)
    assert not app.exception
    assert "Документы проверены" in text
    assert _default_year() in text
    assert "1 год обучения" in text
    assert "Часы совпадают" in text
    assert "Календарь проверен" in text
    assert "замечан" in text
    assert any(item.label == "Подробнее о проверке" for item in app.expander)
    assert any(button.label == "Изменить данные" for button in app.button)
    assert any(
        button.label == "Сформировать календарный план" for button in app.button
    )
    assert "Нормативная и методическая проверка" in text
    assert "Документы закона" in text
    assert "Календарь учреждения" in text
    assert "Сверка ваших часов" in text
    assert "Код приложения обновлён" not in text
    assert "Перезапустите приложение" not in text

def test_calendar_uses_the_scheduling_grid_and_keeps_short_weeks_in_note() -> None:
    weeks = ui.build_academic_weeks(_default_year())
    study, short, breaks = ui._academic_day_sets(_default_year())
    calendar_html = ui._academic_calendar_html(_default_year())

    assert [week.number for week in weeks] == list(range(1, 37))
    assert ui._calendar_day_class(weeks[0].start, study, short, breaks) == "kp-cal-study"
    assert ui._calendar_day_class(weeks[0].end, study, short, breaks) == "kp-cal-wknd"
    assert (
        ui._calendar_day_class(date(2026, 10, 26), study, short, breaks)
        == "kp-cal-break"
    )
    assert "kp-cal-short" not in calendar_html
    assert "№1" in ui._short_weeks_note(_default_year())
    assert "№18" in ui._short_weeks_note(_default_year())
    assert len(ui._TEACHER_LINKS) >= 2


def test_calendar_week_before_generation_requests_calendar_plan() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=20).run()
    next(
        button for button in app.button if button.label == "Открыть календарь"
    ).click().run()

    week_buttons = [button for button in app.button if button.label.startswith("№")]
    assert len(week_buttons) == 36
    assert [button.label for button in week_buttons] == [f"№{number}" for number in range(1, 37)]

    week_buttons[1].click().run()
    assert not app.exception
    assert any(
        "Сформируйте календарный план" in (item.value or "")
        for item in app.info
    )


def test_calendar_month_before_generation_requests_calendar_plan() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=20).run()
    next(
        button for button in app.button if button.label == "Открыть календарь"
    ).click().run()

    month_button = next(
        button for button in app.button if button.label == "Октябрь 2026"
    )
    assert month_button.help == "Открыть план на месяц"
    month_button.click().run()

    assert not app.exception
    assert "Октябрь 2026" in _page_text(app)
    assert "Недели №6–№9" in _page_text(app)
    assert any(
        "Сначала сформируйте календарный план."
        in (item.value or "")
        for item in app.info
    )
    assert not any(
        item.label.startswith("Скачать план на")
        for item in app.get("download_button")
    )


def test_long_week_topic_caption_keeps_every_source_topic() -> None:
    parts = (
        SimpleNamespace(topic_number="4", topic_title="Подготовка маршрута"),
        SimpleNamespace(topic_number="5", topic_title="Работа с картой"),
    )
    row = SimpleNamespace(
        source=SimpleNamespace(
            source=SimpleNamespace(
                week_parts=parts,
                topic_number="4",
                topic_title="Подготовка маршрута",
            )
        )
    )

    assert ui._week_topic_caption(row) == (
        "4. Подготовка маршрута; 5. Работа с картой"
    )

def test_calendar_marks_confirmed_holidays_and_uses_plain_language_legend() -> None:
    study, short, breaks = ui._academic_day_sets(_default_year())

    assert (
        ui._calendar_day_class(date(2026, 11, 4), study, short, breaks)
        == "kp-cal-holiday"
    )
    assert (
        ui._calendar_day_class(date(2027, 1, 1), study, short, breaks)
        == "kp-cal-holiday"
    )
    source = Path(ui.__file__).read_text(encoding="utf-8")
    assert "Красный — выходной / официальный праздник" in source
    assert "Бежевый — рекомендуемые каникулы / перерыв" in source
    assert "переносы выходных 2027 года не размечены" in source
    html_row = ui._calendar_day_row_html(
        year=2026, month=11, day_numbers=[2, 3, 4, 5, 6, 7, 8],
        study=study, short=short, breaks=breaks,
    )
    assert 'title=\"День народного единства\"' in html_row
    january_row = ui._calendar_day_row_html(
        year=2027, month=1, day_numbers=[4, 5, 6, 7, 8, 9, 10],
        study=study, short=short, breaks=breaks,
    )
    assert 'title=\"Рождество Христово\"' in january_row
    assert january_row.count("kp-cal-holiday") == 5
    assert 'data-calendar-week-header=\"numero\"' in source
    header = ui._calendar_days_header_html()
    assert ui._CALENDAR_WEEK_HEADER == "\N{NUMERO SIGN}"
    assert ord(ui._CALENDAR_WEEK_HEADER) == 0x2116
    assert header.startswith(
        '<div class="kp-cal-days-head" lang="ru" translate="no">'
        '<span class="kp-cal-week-head notranslate" translate="no" '
        'data-calendar-week-header="numero" aria-label="№ нед.">&#8470; нед.</span>'
    )
    assert "Нет" not in header
    assert "None" not in header
    assert "False" not in header
    assert (
        ui._calendar_day_class(date(2026, 10, 5), study, short, breaks)
        == "kp-cal-professional"
    )
    teacher_day = ui._calendar_day_row_html(
        year=2026, month=10, day_numbers=[5, 6, 7, 8, 9, 10, 11],
        study=study, short=short, breaks=breaks,
    )
    assert 'title="5 октября — День учителя"' in teacher_day
    assert 'data-professional="5 октября — День учителя"' in teacher_day
    assert "Профессиональная дата" in source
    recommended = ui._recommended_school_breaks(_default_year())
    assert date(2026, 10, 26) in recommended
    assert date(2026, 11, 3) in recommended
    assert date(2026, 12, 31) in recommended
    assert date(2027, 1, 10) in recommended
    assert date(2027, 3, 27) in recommended
    assert date(2027, 4, 4) in recommended
    assert date(2027, 5, 27) in recommended
    assert date(2027, 8, 31) in recommended
    assert date(2027, 2, 15) not in recommended
    first_class = ui._recommended_school_breaks(_default_year(), "1 класс")
    assert date(2027, 2, 15) in first_class
    assert date(2027, 2, 21) in first_class
    assert (
        ui._calendar_day_class(date(2026, 10, 31), study, short, breaks)
        == "kp-cal-break kp-cal-wknd"
    )


def test_recommended_break_cards_reuse_confirmed_calendar_periods() -> None:
    cards_html = ui._recommended_break_cards_html(_default_year())

    assert "Рекомендуемые школьные каникулы" in cards_html
    assert cards_html.count('class="kp-cal-vacation-card"') == 4
    assert "26 октября — 3 ноября 2026" in cards_html
    assert "31 декабря 2026 — 10 января 2027" in cards_html
    assert "27 марта — 4 апреля 2027" in cards_html
    assert "27 мая — 31 августа 2027" in cards_html
    assert "Дополнительные каникулы для 1 класса" not in cards_html

    first_class_html = ui._recommended_break_cards_html(
        _default_year(), "1 класс"
    )
    assert (
        "Дополнительные каникулы для 1 класса: 15–21 февраля 2027"
        in first_class_html
    )


def test_month_detail_uses_scheduling_months_and_ready_snapshot_rows() -> None:
    september = ui._month_weeks(_default_year(), 2026, 9)
    october = ui._month_weeks(_default_year(), 2026, 10)
    january = ui._month_weeks(_default_year(), 2027, 1)
    may = ui._month_weeks(_default_year(), 2027, 5)

    assert [week.number for week in september] == [1, 2, 3, 4, 5]
    assert [week.number for week in october] == [6, 7, 8, 9]
    assert [week.number for week in january] == [19, 20, 21]
    assert [week.number for week in may] == [35, 36]

    row6 = ("05–11.10", "6. Тема", "Теория", "Практика", "Тип", "Результат", "Контроль", "")
    row9 = ("26.10–01.11", "9. Тема", "Теория", "Практика", "Тип", "Результат", "Контроль", "")
    rows = ui._month_detail_rows({6: (row6,), 9: (row9,)}, october)

    assert rows == (("№6", *row6), ("№9", *row9))


def test_monthly_docx_is_an_unchanged_subset_of_the_annual_docx(monkeypatch) -> None:
    # This is a content-projection test; renderer boundaries have separate tests.
    monkeypatch.setattr(ui, 'detect_data_row_page_spans', lambda *args, **kwargs: None)
    annual_path = (
        REFERENCES
        / "Календарный_план_Туристы_проводники_3_год_2026-2027_Верно.docx"
    )
    annual_content = annual_path.read_bytes()
    annual_hash = hashlib.sha256(annual_content).hexdigest()

    def rows_by_week(content: bytes) -> dict[int, tuple[str, ...]]:
        document = Document(BytesIO(content))
        table = document.tables[0]
        columns = ui._columns_for_table(table)
        rows: dict[int, tuple[str, ...]] = {}
        for row in table.rows[2:]:
            first_line = row.cells[columns.week].text.splitlines()[:1]
            try:
                number = int(first_line[0].strip()) if first_line else None
            except ValueError:
                number = None
            if number is not None:
                rows[number] = tuple(cell.text for cell in row.cells)
        return rows

    annual_rows = rows_by_week(annual_content)
    for year, month, expected in (
        (2026, 9, {1, 2, 3, 4, 5}),
        (2026, 10, {6, 7, 8, 9}),
        (2027, 1, {19, 20, 21}),
        (2027, 5, {35, 36}),
    ):
        monthly_content = ui._monthly_plan_docx(
            annual_content, _default_year(), year, month
        )
        monthly_rows = rows_by_week(monthly_content)
        assert set(monthly_rows) == expected
        assert monthly_rows == {
            number: annual_rows[number] for number in sorted(expected)
        }

    assert hashlib.sha256(annual_content).hexdigest() == annual_hash


def test_week_detail_projects_every_ready_plan_row_without_inventing_mark() -> None:
    def row(topic: str, theory: str, practice: str) -> SimpleNamespace:
        return SimpleNamespace(
            source=SimpleNamespace(
                source=SimpleNamespace(
                    week_number=19,
                    date_range="11–17 января",
                    week_parts=(),
                    topic_number="5.1",
                    topic_title=topic,
                )
            ),
            theory_text=theory,
            practice_text=practice,
            lesson_type="Комбинированное занятие",
            planned_result="Планируемый результат",
            assessment_method="Практическое задание",
        )

    rows = ui._week_detail_rows(
        (
            row("Длинная тема, часть 1", "Теория 1", "Практика 1"),
            row("Длинная тема, часть 2", "Теория 2", "Практика 2"),
        ),
        19,
    )

    assert len(rows) == 2
    assert rows[0][:4] == (
        "11–17 января",
        "5.1. Длинная тема, часть 1",
        "Теория 1",
        "Практика 1",
    )
    assert rows[1][1] == "5.1. Длинная тема, часть 2"
    assert all(item[-1] == "" for item in rows)

def test_generation_click_runs_pipeline_and_exposes_download() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=30).run()
    _upload(app, 0, _program_file())
    _upload(app, 2, _template_file())
    app.run()
    _check_button(app).click().run()

    generated = SimpleNamespace(
        filename="calendar.docx",
        content=b"generated-docx",
        warnings=(
            "Ширина таблицы не задана явно; возможны переносы на новые страницы.",
            SLOT_CONTINUE_WARNING,
            SLOT_PACK_WARNING,
            "Неоднозначное соответствие для «Тема»: вариант А",
        ),
        ai_usage=None,
    )
    assert [item.label for item in app.text_input] == ["Группа №", "Класс", "ФИО педагога"]
    assert not any("ИИ" in (item.label or "") for item in getattr(app, "checkbox", []))
    generate = next(
        button for button in app.button if button.label == "Сформировать календарный план"
    )
    with patch("calendar_pedagoga.ui.run_calendar_pipeline", return_value=generated) as pipeline:
        generate.click().run()

    pipeline.assert_called_once()
    assert pipeline.call_args.kwargs["use_ai"] is False
    assert "ai_provider" not in pipeline.call_args.kwargs
    assert pipeline.call_args.kwargs["group_number"] == ""
    assert pipeline.call_args.kwargs["class_name"] == ""
    assert pipeline.call_args.kwargs["teacher_name"] == ""
    assert pipeline.call_args.kwargs["academic_year"] == _default_year()
    assert "1 г" in (pipeline.call_args.kwargs["program_filename"] or "")
    assert "Дополнить содержание с помощью ИИ" not in _page_text(app)
    assert "Группа Нет" not in _page_text(app)
    assert app.session_state["calendar_generation_pending"] is False
    assert app.session_state["calendar_generation_succeeded"] is True
    assert app.session_state["calendar_download"].content == b"generated-docx"
    assert "Календарный план готов" in _page_text(app)
    assert app.get("download_button")[0].label == (
        f"Скачать план за {_default_year()} учебный год"
    )
    assert "Ширина таблицы" not in _page_text(app)
    assert not any("Ширина таблицы" in (item.value or "") for item in app.warning)
    assert not any("продолжение уже представленного" in (item.value or "") for item in app.warning)
    assert not any("нескольких исходных практических" in (item.value or "") for item in app.warning)
    assert any(
        "Неоднозначное соответствие для «Тема»" in (item.value or "")
        for item in app.warning
    )
    stored = app.session_state["calendar_warnings"]
    assert SLOT_CONTINUE_WARNING in stored
    assert SLOT_PACK_WARNING in stored


def test_generated_plan_survives_calendar_and_week_click_reruns() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=30).run()
    _upload(app, 0, _program_file())
    _upload(app, 2, _template_file())
    app.run()
    _check_button(app).click().run()

    source = SimpleNamespace(
        week_number=19,
        date_range="11–17.01",
        week_parts=(),
        topic_number="4.2",
        topic_title="Аптечка",
    )
    resolved = SimpleNamespace(
        source=SimpleNamespace(source=source),
        theory_text="Фактическая теория",
        practice_text="Фактическая практика",
        lesson_type="Комбинированное занятие",
        planned_result="Фактический результат",
        assessment_method="Практическое задание",
    )
    october_source = SimpleNamespace(
        week_number=6,
        date_range="05–11.10",
        week_parts=(),
        topic_number="2.1",
        topic_title="Октябрьская тема",
    )
    october_resolved = SimpleNamespace(
        source=SimpleNamespace(source=october_source),
        theory_text="Октябрьская теория",
        practice_text="Октябрьская практика",
        lesson_type="Практическое занятие",
        planned_result="Октябрьский результат",
        assessment_method="Октябрьский контроль",
    )
    generated = SimpleNamespace(
        filename="calendar.docx",
        content=b"generated-docx",
        warnings=(),
        ai_usage=None,
        resolved_lessons=(october_resolved, resolved),
    )
    generate = next(
        button for button in app.button
        if button.label == "Сформировать календарный план"
    )
    with patch("calendar_pedagoga.ui.run_calendar_pipeline", return_value=generated):
        generate.click().run()

    open_buttons = [
        button for button in app.button
        if button.label == "Открыть календарь"
    ]
    assert len(open_buttons) == 1
    with patch(
        "calendar_pedagoga.ui._monthly_plan_docx",
        return_value=b"monthly-docx",
    ):
        open_buttons[0].click().run()
        next(
            button for button in app.button if button.label == "Октябрь 2026"
        ).click().run()

        month_text = _page_text(app)
        assert "Сформируйте календарный план" not in month_text
        assert "Недели №6–№9" in month_text
        assert "Октябрьская теория" in month_text
        assert "Октябрьская практика" in month_text
        assert "Октябрьский результат" in month_text
        assert "Октябрьский контроль" in month_text
        assert any(
            button.label == "Скачать план на октябрь 2026"
            for button in app.get("download_button")
        )

        next(
            button for button in app.button if button.label == "← К календарю"
        ).click().run()
        next(button for button in app.button if button.label == "№19").click().run()

    text = _page_text(app)
    assert "Сформируйте календарный план" not in text
    assert "Неделя №19 · 11–17 января" in text
    assert "Фактическая теория" in text
    assert "Фактическая практика" in text
    assert "Фактический результат" in text
    assert "Практическое задание" in text
    assert app.session_state["calendar_resolved_lessons"] == (
        october_resolved,
        resolved,
    )

def test_teacher_name_is_optional_and_invalidates_download() -> None:
    app = AppTest.from_file(str(APP_PATH), default_timeout=30).run()
    _upload(app, 0, _program_file())
    _upload(app, 2, _template_file())
    app.run()
    _check_button(app).click().run()

    generated = SimpleNamespace(
        filename="calendar.docx",
        content=b"generated-docx",
        warnings=(),
        ai_usage=None,
    )
    teacher = next(item for item in app.text_input if item.label == "ФИО педагога")
    assert teacher.value in {"", None}

    generate = next(
        button for button in app.button if button.label == "Сформировать календарный план"
    )
    with patch("calendar_pedagoga.ui.run_calendar_pipeline", return_value=generated) as pipeline:
        generate.click().run()
    assert pipeline.call_args.kwargs["teacher_name"] == ""
    assert app.session_state["calendar_generation_succeeded"] is True
    assert app.session_state["calendar_download"].content == b"generated-docx"

    teacher = next(item for item in app.text_input if item.label == "ФИО педагога")
    teacher.set_value("Иванов И.И.").run()
    assert "calendar_generation_invalidated" in app.session_state
    assert app.session_state["calendar_generation_invalidated"]
    assert "calendar_download" not in app.session_state
    assert app.session_state["analysis_ready"] is True

    generate = next(
        button for button in app.button if button.label == "Сформировать календарный план"
    )
    with patch("calendar_pedagoga.ui.run_calendar_pipeline", return_value=generated) as pipeline:
        generate.click().run()
    assert pipeline.call_args.kwargs["teacher_name"] == "Иванов И.И."
    assert app.session_state["calendar_generation_succeeded"] is True


def test_teacher_generation_warnings_hide_internal_diagnostics_and_collapse_ce2(
    caplog,
) -> None:
    visible = "Неоднозначное соответствие для «Тема»: вариант А"
    with caplog.at_level("INFO", logger="calendar_pedagoga.ui"):
        shown = _teacher_generation_warnings(
            (
                "Ширина таблицы не задана явно; возможны переносы на новые страницы.",
                SLOT_CONTINUE_WARNING,
                "Безопасный шаблон CE2: broken_clause_join.",
                "Безопасный шаблон CE2: unproven_object_case.",
                "Безопасный шаблон CE2: broken_clause_join.",
                SLOT_PACK_WARNING,
                visible,
            )
        )
    assert shown == (
        visible,
        "Некоторые формулировки автоматически приведены "
        "к безопасному нейтральному виду.",
    )
    assert "broken_clause_join, unproven_object_case" in caplog.text
    assert all(
        code not in " ".join(shown)
        for code in ("broken_clause_join", "unproven_object_case")
    )
