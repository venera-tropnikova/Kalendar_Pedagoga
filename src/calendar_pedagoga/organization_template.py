"""Выбор и валидация шаблона календарного плана организации."""

from __future__ import annotations

from dataclasses import dataclass
from enum import StrEnum
from io import BytesIO
from pathlib import Path

from docx import Document


ORG_TEMPLATE_UNSUPPORTED_MESSAGE = (
    "Структура шаблона организации пока не поддерживается. "
    "Можно использовать стандартный шаблон приложения."
)

_SUPPORTED_COLUMN_COUNT = 8
_MIN_HEADER_ROWS = 2
_MAX_TEMPLATE_ROWS = 3

_REQUIRED_HEADER_MARKERS = (
    "месяц",
    "неделя",
    "теоретические",
    "практические",
    "тип занятия",
    "планируемый результат",
    "вид контроля",
)


class CalendarTemplateSource(StrEnum):
    STANDARD = "standard"
    ORGANIZATION = "organization"


class OrganizationTemplateError(ValueError):
    """Загруженный шаблон организации не соответствует поддерживаемой структуре."""


@dataclass(frozen=True)
class CalendarTemplateSelection:
    """Источник структуры будущего календарного плана."""

    source: CalendarTemplateSource
    filename: str | None = None
    content: bytes | None = None

    @property
    def uses_organization_template(self) -> bool:
        return self.source is CalendarTemplateSource.ORGANIZATION


def validate_organization_template(content: bytes) -> None:
    """Проверить минимально поддерживаемую структуру пустого шаблона организации."""

    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise OrganizationTemplateError(ORG_TEMPLATE_UNSUPPORTED_MESSAGE) from error

    if len(document.tables) != 1:
        raise OrganizationTemplateError(ORG_TEMPLATE_UNSUPPORTED_MESSAGE)

    table = document.tables[0]
    if len(table.columns) != _SUPPORTED_COLUMN_COUNT:
        raise OrganizationTemplateError(ORG_TEMPLATE_UNSUPPORTED_MESSAGE)

    if not (_MIN_HEADER_ROWS <= len(table.rows) <= _MAX_TEMPLATE_ROWS):
        raise OrganizationTemplateError(ORG_TEMPLATE_UNSUPPORTED_MESSAGE)

    header_text = " ".join(
        cell.text for row in table.rows[:_MIN_HEADER_ROWS] for cell in row.cells
    ).casefold()
    if any(marker not in header_text for marker in _REQUIRED_HEADER_MARKERS):
        raise OrganizationTemplateError(ORG_TEMPLATE_UNSUPPORTED_MESSAGE)

    for row in table.rows[_MIN_HEADER_ROWS:]:
        if any(cell.text.strip() for cell in row.cells):
            raise OrganizationTemplateError(ORG_TEMPLATE_UNSUPPORTED_MESSAGE)


def select_calendar_template(
    filename: str | None = None,
    content: bytes | None = None,
) -> CalendarTemplateSelection:
    """Выбрать загруженный шаблон или стандартный шаблон приложения."""

    if filename is None and content is None:
        return CalendarTemplateSelection(CalendarTemplateSource.STANDARD)
    if not filename or content is None:
        raise ValueError("Для шаблона организации нужны имя файла и содержимое.")
    if Path(filename).suffix.lower() not in {".doc", ".docx"}:
        raise ValueError("Шаблон организации должен быть в формате DOC или DOCX.")
    if not content:
        raise ValueError("Файл шаблона организации пуст.")
    validate_organization_template(content)
    return CalendarTemplateSelection(
        CalendarTemplateSource.ORGANIZATION,
        filename=Path(filename).name,
        content=bytes(content),
    )
