"""Структурная и визуальная проверка сгенерированного календарного DOCX."""

from __future__ import annotations

from dataclasses import dataclass
from enum import StrEnum
from io import BytesIO
import json
import logging
from pathlib import Path
import re
import shutil
import struct
import subprocess
import tempfile

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from calendar_pedagoga.program_parsing import find_libreoffice

logger = logging.getLogger(__name__)

_WD_ACTIVE_END_PAGE_NUMBER = 3
_WD_EXPORT_FORMAT_PDF = 17
_WD_ALERTS_NONE = 0


def find_microsoft_word() -> bool:
    """Проверить, доступен ли Microsoft Word через COM."""

    try:
        import win32com.client  # noqa: F401
    except ImportError:
        return False
    try:
        word = win32com.client.DispatchEx("Word.Application")
    except Exception:
        return False
    try:
        word.Quit()
    except Exception:
        pass
    return True


def _run_with_word_document(content: bytes, callback):
    """Открыть DOCX в Microsoft Word и выполнить callback(word, doc)."""

    import pythoncom
    import win32com.client

    temp_path = Path(tempfile.mkdtemp(prefix="calendar_pedagoga_word_"))
    docx_path = temp_path / "calendar.docx"
    docx_path.write_bytes(content)
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = _WD_ALERTS_NONE
        document = word.Documents.Open(
            str(docx_path),
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )
        return callback(word, document)
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        shutil.rmtree(temp_path, ignore_errors=True)


def _detect_data_row_indices_by_page_word(
    content: bytes,
    *,
    total_rows: int,
) -> tuple[tuple[int, ...], ...] | None:
    """Определить строки данных по страницам через пагинацию Microsoft Word."""

    try:
        import win32com.client  # noqa: F401
    except ImportError:
        return None

    def _collect(_word, document) -> tuple[tuple[int, ...], ...] | None:
        if document.Tables.Count < 1:
            return None
        table = document.Tables(1)
        pages: dict[int, list[int]] = {}
        for data_index in range(total_rows):
            row_number = data_index + 3  # две строки заголовка
            if row_number > table.Rows.Count:
                break
            # Колонка недели: не зависит от вертикального merge месяца.
            cell = table.Cell(row_number, 2)
            page_number = int(cell.Range.Information(_WD_ACTIVE_END_PAGE_NUMBER))
            pages.setdefault(page_number, []).append(data_index)
        if not pages:
            return None
        return tuple(tuple(pages[page]) for page in sorted(pages))

    try:
        return _run_with_word_document(content, _collect)
    except Exception:
        return None


def _docx_to_pdf_bytes_word(content: bytes) -> bytes | None:
    """Сконвертировать DOCX в PDF через Microsoft Word (пагинация как в Word)."""

    try:
        import win32com.client  # noqa: F401
    except ImportError:
        return None

    temp_path = Path(tempfile.mkdtemp(prefix="calendar_pedagoga_word_pdf_"))
    docx_path = temp_path / "calendar.docx"
    pdf_path = temp_path / "calendar.pdf"
    docx_path.write_bytes(content)

    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = _WD_ALERTS_NONE
        document = word.Documents.Open(
            str(docx_path),
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )
        document.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=_WD_EXPORT_FORMAT_PDF,
            OpenAfterExport=False,
            OptimizeFor=0,
            BitmapMissingFonts=True,
            DocStructureTags=True,
            CreateBookmarks=0,
        )
        if not pdf_path.is_file():
            return None
        return pdf_path.read_bytes()
    except Exception:
        return None
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        shutil.rmtree(temp_path, ignore_errors=True)


class QASeverity(StrEnum):
    ERROR = "error"
    WARNING = "warning"


@dataclass(frozen=True)
class QAIssue:
    severity: QASeverity
    message: str


@dataclass(frozen=True)
class VisualPageReport:
    page_number: int
    path: Path
    width: int
    height: int
    size_bytes: int
    ink_ratio: float


_REQUIRED_HEADER_MARKERS = (
    "месяц",
    "неделя",
    "теоретические",
    "практические",
    "тип занятия",
    "планируемый результат",
    "вид контроля",
)

_MIN_PAGE_FILE_BYTES = 8_000
_MIN_INK_RATIO = 0.01
_MIN_PAGE_DIMENSION = 400


def _month_cell_is_continuation(cell) -> bool:
    """Пустая ячейка месяца допустима как продолжение w:vMerge."""
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return False
    merge = tc_pr.find(qn("w:vMerge"))
    if merge is None:
        return False
    return merge.get(qn("w:val")) != "restart"


def _header_paragraph_texts(document) -> tuple[str, ...]:
    """Return body paragraphs before the calendar table, preserving template order."""

    texts: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:tbl"):
            break
        if child.tag == qn("w:p"):
            text = Paragraph(child, document).text.strip()
            if text:
                texts.append(text)
    return tuple(texts)


def _has_calendar_plan_heading(texts: tuple[str, ...]) -> bool:
    return any(
        "календар" in text.casefold() and "план" in text.casefold()
        for text in texts
    )


def validate_calendar_docx(
    content: bytes,
    *,
    expected_weeks: int,
) -> tuple[QAIssue, ...]:
    """Проверить DOCX после генерации: структура, таблица и постраничная целостность."""

    issues: list[QAIssue] = []
    try:
        document = Document(BytesIO(content))
    except Exception as error:
        return (QAIssue(QASeverity.ERROR, f"DOCX не читается: {error}"),)

    header_paragraphs = _header_paragraph_texts(document)
    if not header_paragraphs:
        issues.append(QAIssue(QASeverity.ERROR, "Отсутствует заголовок документа."))
    elif not _has_calendar_plan_heading(header_paragraphs):
        issues.append(
            QAIssue(
                QASeverity.ERROR,
                "В шапке документа отсутствует заголовок календарного плана.",
            )
        )

    if len(document.paragraphs) < 3:
        issues.append(QAIssue(QASeverity.ERROR, "Недостаточно служебных абзацев шапки."))

    if not document.tables:
        issues.append(QAIssue(QASeverity.ERROR, "В документе нет таблицы календаря."))
        return tuple(issues)

    if len(document.tables) != 1:
        issues.append(QAIssue(QASeverity.WARNING, "Ожидалась одна таблица календаря."))

    table = document.tables[0]
    if len(table.rows) < 2:
        issues.append(QAIssue(QASeverity.ERROR, "Таблица не содержит строк заголовка."))
        return tuple(issues)

    header_text = " ".join(
        cell.text for row in table.rows[:2] for cell in row.cells
    ).casefold()
    missing = [marker for marker in _REQUIRED_HEADER_MARKERS if marker not in header_text]
    if missing:
        issues.append(
            QAIssue(
                QASeverity.ERROR,
                "В заголовке таблицы отсутствуют обязательные колонки: "
                + ", ".join(missing),
            )
        )

    data_rows = table.rows[2:]
    if len(data_rows) != expected_weeks:
        issues.append(
            QAIssue(
                QASeverity.ERROR,
                f"Ожидалось {expected_weeks} строк данных, найдено {len(data_rows)}.",
            )
        )

    column_count = len(table.columns)
    if column_count < 8:
        issues.append(
            QAIssue(
                QASeverity.ERROR,
                f"Недостаточно колонок в таблице: {column_count}.",
            )
        )

    for section in document.sections:
        if section.page_width <= 0 or section.page_height <= 0:
            issues.append(
                QAIssue(QASeverity.ERROR, "Некорректные параметры страницы документа.")
            )
            break
        if section.left_margin <= 0 or section.right_margin <= 0:
            issues.append(
                QAIssue(
                    QASeverity.WARNING,
                    "Нетипичные поля страницы; проверьте визуальное отображение.",
                )
            )

    if document.tables:
        table = document.tables[0]
        tbl_pr = table._tbl.tblPr
        tbl_width = getattr(tbl_pr, "tblW", None) if tbl_pr is not None else None
        if tbl_pr is not None and tbl_width is None:
            issues.append(
                QAIssue(
                    QASeverity.WARNING,
                    "Ширина таблицы не задана явно; возможны переносы на новые страницы.",
                )
            )
        for index, row in enumerate(table.rows[2:], start=1):
            if not any(cell.text.strip() for cell in row.cells[:5]):
                issues.append(
                    QAIssue(
                        QASeverity.ERROR,
                        f"Строка {index + 2}: полностью пустая строка календаря.",
                    )
                )

    week_numbers: list[int] = []
    for index, row in enumerate(data_rows, start=1):
        cells = row.cells
        if len(cells) < 8:
            issues.append(
                QAIssue(
                    QASeverity.ERROR,
                    f"Строка {index + 2}: неполный набор ячеек ({len(cells)}).",
                )
            )
            continue

        month = cells[0].text.strip()
        week_cell = cells[1].text.strip()
        if not month and not _month_cell_is_continuation(cells[0]):
            issues.append(
                QAIssue(QASeverity.ERROR, f"Строка {index + 2}: пустой месяц.")
            )
        if not week_cell:
            issues.append(
                QAIssue(QASeverity.ERROR, f"Строка {index + 2}: пустая неделя/дата.")
            )
        else:
            match = re.search(r"(\d+)", week_cell.splitlines()[0])
            if match:
                week_numbers.append(int(match.group(1)))
            else:
                issues.append(
                    QAIssue(
                        QASeverity.ERROR,
                        f"Строка {index + 2}: не распознан номер недели.",
                    )
                )

        for cell in cells:
            if cell.text is None:
                issues.append(
                    QAIssue(
                        QASeverity.ERROR,
                        f"Строка {index + 2}: недоступная ячейка таблицы.",
                    )
                )

    if week_numbers and week_numbers != list(range(1, len(week_numbers) + 1)):
        issues.append(
            QAIssue(
                QASeverity.ERROR,
                "Нумерация недель в таблице не непрерывна.",
            )
        )

    return tuple(issues)


def has_blocking_qa_issues(issues: tuple[QAIssue, ...]) -> bool:
    return any(issue.severity is QASeverity.ERROR for issue in issues)


def _pdf_page_count(pdf_path: Path) -> int:
    data = pdf_path.read_bytes()
    counts = [int(value) for value in re.findall(rb"/Count\s+(\d+)", data)]
    return max(counts) if counts else 0


def _png_dimensions(path: Path) -> tuple[int, int]:
    with path.open("rb") as handle:
        signature = handle.read(8)
        if signature != b"\x89PNG\r\n\x1a\n":
            raise ValueError(f"Not a PNG: {path}")
        handle.read(4)
        chunk = handle.read(4)
        if chunk != b"IHDR":
            raise ValueError(f"PNG IHDR missing: {path}")
        width, height = struct.unpack(">II", handle.read(8))
        return width, height


def _png_ink_ratio(path: Path) -> float:
    """Доля не-белых пикселей (эвристика читаемости)."""

    try:
        from PIL import Image
    except ImportError:
        return _png_ink_ratio_fallback(path)

    with Image.open(path) as image:
        rgb = image.convert("RGB")
        width, height = rgb.size
        step = max(1, (width * height) // 20_000)
        dark = 0
        seen = 0
        for index in range(0, width * height, step):
            x = index % width
            y = index // width
            red, green, blue = rgb.getpixel((x, y))
            seen += 1
            if red < 245 or green < 245 or blue < 245:
                dark += 1
        return dark / seen if seen else 0.0


def _png_ink_ratio_fallback(path: Path) -> float:
    raw = path.read_bytes()
    if raw[:8] != b"\x89PNG\r\n\x1a\n":
        return 0.0
    return 0.05 if path.stat().st_size >= _MIN_PAGE_FILE_BYTES else 0.0


def _libreoffice_version(soffice: Path) -> str:
    try:
        result = subprocess.run(
            [str(soffice), "--headless", "--version"],
            capture_output=True,
            text=True,
            timeout=15,
        )
    except (OSError, subprocess.TimeoutExpired) as error:
        return f"unavailable: {type(error).__name__}: {error}"
    output = (result.stdout or result.stderr or "").strip()
    return output or f"unavailable: return code {result.returncode}"


def _log_libreoffice_failure(
    soffice: Path,
    command: list[str],
    *,
    return_code: int | None,
    stdout: str | None,
    stderr: str | None,
    pdf_path: Path | None = None,
    pdf_page_count: int | None = None,
    pymupdf_error: BaseException | None = None,
) -> None:
    pdf_exists = pdf_path.is_file() if pdf_path is not None else False
    diagnostics = {
        "soffice_executable": str(soffice),
        "libreoffice_version": _libreoffice_version(soffice),
        "command": command,
        "return_code": return_code,
        "stdout": stdout or "",
        "stderr": stderr or "",
        "pdf_exists": pdf_exists,
        "pdf_size": pdf_path.stat().st_size if pdf_exists else None,
        "pdf_page_count": pdf_page_count,
        "pymupdf_exception_type": (
            type(pymupdf_error).__name__ if pymupdf_error is not None else None
        ),
        "pymupdf_exception_message": (
            str(pymupdf_error) if pymupdf_error is not None else None
        ),
    }
    logger.error(
        "LibreOffice visual QA failure: %s",
        json.dumps(diagnostics, ensure_ascii=False, sort_keys=True),
    )


def _run_soffice(
    soffice: Path,
    arguments: list[str],
    temp_path: Path,
) -> subprocess.CompletedProcess[str]:
    profile = Path(tempfile.mkdtemp(prefix="lo_profile_", dir=temp_path))
    command = [
        str(soffice),
        "--headless",
        "--norestore",
        f"-env:UserInstallation={profile.as_uri()}",
        *arguments,
    ]
    try:
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except subprocess.TimeoutExpired as error:
        _log_libreoffice_failure(
            soffice,
            command,
            return_code=None,
            stdout=error.stdout,
            stderr=error.stderr,
        )
        raise RuntimeError("LibreOffice convert timed out; diagnostics logged.") from error
    except OSError as error:
        _log_libreoffice_failure(
            soffice,
            command,
            return_code=None,
            stdout=None,
            stderr=str(error),
        )
        raise RuntimeError("LibreOffice could not start; diagnostics logged.") from error

    if result.returncode != 0:
        _log_libreoffice_failure(
            soffice,
            command,
            return_code=result.returncode,
            stdout=result.stdout,
            stderr=result.stderr,
        )
        raise RuntimeError("LibreOffice convert failed; diagnostics logged.")
    return result


def _raise_missing_libreoffice_output(
    soffice: Path,
    result: subprocess.CompletedProcess[str],
    message: str,
) -> None:
    _log_libreoffice_failure(
        soffice,
        list(result.args),
        return_code=result.returncode,
        stdout=result.stdout,
        stderr=result.stderr,
    )
    raise RuntimeError(f"{message}; diagnostics logged.")


def _docx_to_pdf_bytes_libreoffice(content: bytes) -> bytes | None:
    """Convert DOCX to PDF with LibreOffice only."""

    soffice = find_libreoffice()
    if soffice is None:
        return None

    temp_path = Path(tempfile.mkdtemp(prefix="calendar_pedagoga_pdf_"))
    try:
        docx_path = temp_path / "calendar.docx"
        docx_path.write_bytes(content)
        pdf_dir = temp_path / "pdf"
        pdf_dir.mkdir()
        _run_soffice(
            soffice,
            [
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_dir),
                str(docx_path),
            ],
            temp_path,
        )
        pdfs = sorted(pdf_dir.glob("*.pdf"))
        if not pdfs:
            return None
        return pdfs[0].read_bytes()
    finally:
        shutil.rmtree(temp_path, ignore_errors=True)


def _docx_to_pdf_bytes(content: bytes) -> bytes | None:
    """Сконвертировать DOCX в PDF: сначала Word (эталон), иначе LibreOffice."""

    word_pdf = _docx_to_pdf_bytes_word(content)
    if word_pdf is not None:
        return word_pdf
    return _docx_to_pdf_bytes_libreoffice(content)


def _pagination_measurement_copy(content: bytes) -> bytes:
    """Make vertical identifiers measurable by LibreOffice without changing output.

    LibreOffice keeps a table row whole when it contains vertical Month/Week/Date
    text even if ``w:cantSplit`` is absent.  The copy is used only to measure the
    narrative cells; the returned production DOCX retains its text directions.
    """

    document = Document(BytesIO(content))
    if not document.tables:
        return content
    for row in document.tables[0].rows[2:]:
        # Calendar schema: the first three physical columns are Month, Week
        # and Date. Narrative cells must retain their production geometry even
        # in the measurement copy.
        for cell in row._tr.tc_lst[:3]:
            for direction in list(cell.xpath("./w:tcPr/w:textDirection")):
                direction.getparent().remove(direction)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@dataclass(frozen=True)
class DataRowPageSpan:
    """Physical span; split_safe requires positive rendered-cell evidence."""
    start_page: int
    end_page: int
    split_safe: bool = False


def _data_row_page_spans_pdf(
    content: bytes, pdf: bytes, total_rows: int,
) -> tuple[DataRowPageSpan, ...] | None:
    """Match complete cell contents across PDF table fragments, never week numbers alone."""
    import pymupdf

    source = Document(BytesIO(content)).tables[0]

    def normalized(text):
        return "".join(char for char in (text or "").casefold() if char.isalnum())

    def identifier(text):
        # Punctuation belongs to dates/identifiers and must not disappear.
        return re.sub(r"\s+", "", text or "")

    expected = []
    identifiers = []
    protected_columns = []
    merged_rows = []
    for row in source.rows[2:]:
        seen = set()
        cells = []
        for cell in row.cells:
            cells.append("" if cell._tc in seen else normalized(cell.text))
            seen.add(cell._tc)
        expected.append(cells)
        identifiers.append([identifier(cell.text) for cell in row.cells])
        protected_columns.append({0, 1} | {
            index for index, cell in enumerate(row.cells)
            if cell._tc.xpath('./w:tcPr/w:textDirection | ./w:tcPr/w:vMerge')
        })
        merged_rows.append(bool(row._tr.xpath('.//w:vMerge')))
    if len(expected) != total_rows:
        return None
    spans = []
    accumulated = [""] * len(source.columns)
    start_page = None
    complete_identifiers = set()
    with pymupdf.open(stream=pdf, filetype="pdf") as document:
        for page_number, page in enumerate(document, start=1):
            tables = [table for table in page.find_tables().tables
                      if table.col_count == len(source.columns)]
            if len(tables) != 1:
                return None
            for fragment in tables[0].extract()[2:]:
                if len(spans) >= total_rows:
                    return None
                target = expected[len(spans)]
                protected = protected_columns[len(spans)]
                start_page = start_page or page_number
                # Rotated cells can be clipped or attached to the next PDF row.
                # Match row identity by the unchanged narrative cells instead.
                for column in protected:
                    if identifier(fragment[column]) == identifiers[len(spans)][column]:
                        complete_identifiers.add(column)
                body_columns = set(range(len(target))) - protected
                if not any(target[column] for column in body_columns):
                    return None
                for column in body_columns:
                    accumulated[column] += normalized(fragment[column])
                    if not target[column].startswith(accumulated[column]):
                        return None
                if all(accumulated[column] == target[column] for column in body_columns):
                    safe = (protected <= complete_identifiers and not merged_rows[len(spans)])
                    spans.append(DataRowPageSpan(start_page, page_number, safe))
                    accumulated = [""] * len(target)
                    start_page = None
                    complete_identifiers = set()
    return tuple(spans) if len(spans) == total_rows and start_page is None else None


def detect_data_row_page_spans(
    content: bytes, *, total_rows: int,
) -> tuple[DataRowPageSpan, ...] | None:
    """Measure every data row, including both ends of rows crossing a page."""
    if total_rows == 0:
        return ()

    # Word paginates the production layout directly. LibreOffice cannot split
    # rows containing vertical identifiers, so on Linux it receives a temporary
    # measurement copy with those directions removed. The original content is
    # still the source of expected cell text and is never modified.
    pdf = _docx_to_pdf_bytes_word(content)
    if pdf is None:
        pdf = _docx_to_pdf_bytes_libreoffice(
            _pagination_measurement_copy(content)
        )
    if pdf is not None:
        try:
            spans = _data_row_page_spans_pdf(content, pdf, total_rows)
            if spans is not None:
                return spans
        except Exception:
            logger.debug("PDF row-span measurement unavailable", exc_info=False)

    def collect(_word, document):
        document.Repaginate()
        pages = [[] for _ in range(total_rows)]
        for cell in document.Tables(1).Range.Cells:
            index = int(cell.RowIndex) - 3
            if not 0 <= index < total_rows or int(cell.ColumnIndex) == 1:
                continue
            first = cell.Range.Duplicate
            last = cell.Range.Duplicate
            first.Collapse(1)  # wdCollapseStart
            last.End = max(last.Start, last.End - 1)  # exclude end-of-cell marker
            last.Collapse(0)  # wdCollapseEnd
            pages[index].extend((int(first.Information(_WD_ACTIVE_END_PAGE_NUMBER)),
                                 int(last.Information(_WD_ACTIVE_END_PAGE_NUMBER))))
        if any(not values for values in pages):
            return None
        return tuple(DataRowPageSpan(min(values), max(values)) for values in pages)

    try:
        spans = _run_with_word_document(content, collect)
        if spans is not None:
            return spans
    except Exception:
        logger.debug("Word row-span measurement unavailable", exc_info=False)
    return None


def detect_data_row_indices_by_page(
    content: bytes,
    *,
    total_rows: int,
) -> tuple[tuple[int, ...], ...] | None:
    """Определить индексы строк данных (0-based) на каждой странице.

    Эталон — пагинация Microsoft Word. LibreOffice/PDF используется только
    как запасной вариант, если Word COM недоступен.
    """

    word_pages = _detect_data_row_indices_by_page_word(content, total_rows=total_rows)
    if word_pages is not None:
        return word_pages

    pdf_bytes = _docx_to_pdf_bytes(content)
    if pdf_bytes is None:
        return None

    try:
        import pymupdf
    except ImportError:
        return None

    document = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        rows_by_page: list[tuple[int, ...]] = []
        for page in document:
            page_height = page.rect.height
            week_positions: list[tuple[int, float]] = []
            for x0, y0, _x1, _y1, text, *_rest in page.get_text("words"):
                if y0 > page_height - 45:
                    continue
                if re.fullmatch(r"\d{1,2}", text):
                    week_number = int(text)
                    if 1 <= week_number <= total_rows:
                        week_positions.append((week_number, y0))
            week_positions.sort(key=lambda item: item[1])
            rows_by_page.append(
                tuple(week_number - 1 for week_number, _y in week_positions)
            )
        return tuple(rows_by_page)
    finally:
        document.close()


def _month_words_on_page(page, month: str) -> list[tuple[float, float, str]]:
    import pymupdf

    assert isinstance(page, pymupdf.Page)
    tokens: list[tuple[float, float, str]] = []
    for x0, y0, _x1, _y1, text, *_rest in page.get_text("words"):
        cleaned = text.strip()
        if cleaned:
            tokens.append((x0, y0, cleaned))
    month_lower = month.casefold()
    return [
        token
        for token in tokens
        if month_lower in token[2].casefold() or token[2].casefold() in month_lower
    ]


def verify_month_labels_by_page(
    content: bytes,
    *,
    months: tuple[str, ...],
) -> tuple[str, ...]:
    """Проверить, что месяц виден у начала каждого месячного блока на каждой странице PDF."""

    rows_by_page = detect_data_row_indices_by_page(content, total_rows=len(months))
    pdf_bytes = _docx_to_pdf_bytes(content)
    if rows_by_page is None or pdf_bytes is None:
        return ("Не удалось проверить подписи месяцев: PDF/pymupdf недоступны.",)

    try:
        import pymupdf
    except ImportError:
        return ("Не удалось проверить подписи месяцев: pymupdf недоступен.",)

    issues: list[str] = []
    document = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        for page_number, page_rows in enumerate(rows_by_page, start=1):
            if not page_rows:
                continue
            page = document[page_number - 1]
            page_height = page.rect.height
            week_positions: dict[int, float] = {}
            for x0, y0, _x1, _y1, text, *_rest in page.get_text("words"):
                if y0 > page_height - 45:
                    continue
                if re.fullmatch(r"\d{1,2}", text):
                    week_number = int(text)
                    if week_number - 1 in page_rows:
                        week_positions.setdefault(week_number - 1, y0)

            group_start = page_rows[0]
            for index in range(1, len(page_rows)):
                row_index = page_rows[index]
                if months[row_index] != months[group_start]:
                    issues.extend(
                        _month_label_issues_for_segment(
                            page_number=page_number,
                            page=page,
                            months=months,
                            group_start=group_start,
                            group_end=page_rows[index - 1],
                            week_positions=week_positions,
                        )
                    )
                    group_start = row_index
            issues.extend(
                _month_label_issues_for_segment(
                    page_number=page_number,
                    page=page,
                    months=months,
                    group_start=group_start,
                    group_end=page_rows[-1],
                    week_positions=week_positions,
                )
            )
    finally:
        document.close()
    return tuple(issues)


def _month_label_issues_for_segment(
    *,
    page_number: int,
    page,
    months: tuple[str, ...],
    group_start: int,
    group_end: int,
    week_positions: dict[int, float],
) -> list[str]:
    month = months[group_start]
    segment_week_positions = [
        week_positions[row_index]
        for row_index in range(group_start, group_end + 1)
        if row_index in week_positions
    ]
    if not segment_week_positions:
        return [
            f"Страница {page_number}: не найдены недели "
            f"{group_start + 1}–{group_end + 1} для проверки месяца «{month}»."
        ]

    month_hits = _month_words_on_page(page, month)
    min_y = min(segment_week_positions)
    max_y = max(segment_week_positions)
    nearby_hits = [
        hit for hit in month_hits if hit[0] <= 120 and min_y - 20 <= hit[1] <= max_y + 20
    ]
    if not nearby_hits:
        return [
            f"Страница {page_number}: у недели {group_start + 1} не видна подпись "
            f"месяца «{month}» (блок до недели {group_end + 1})."
        ]

    # Подпись должна быть внутри сегмента и ближе к его вертикальному центру,
    # а не у первой/последней строки блока.
    segment_height = max(max_y - min_y, 1.0)
    mid_y = (min_y + max_y) / 2
    label_y = sum(hit[1] for hit in nearby_hits) / len(nearby_hits)
    if abs(label_y - mid_y) > segment_height * 0.4 and segment_height > 40:
        return [
            f"Страница {page_number}: подпись «{month}» (y={label_y:.1f}) "
            f"не по центру блока недель {group_start + 1}–{group_end + 1} "
            f"(центр y={mid_y:.1f})."
        ]
    return []


def render_docx_pages(content: bytes, output_dir: Path) -> tuple[Path, ...]:
    """Отрендерить все страницы DOCX в PNG.

    Приоритет: Microsoft Word → PDF → PNG (пагинация как в Word).
    Запасной путь: LibreOffice → PDF, затем PyMuPDF → PNG.
    """

    output_dir.mkdir(parents=True, exist_ok=True)

    word_pdf = _docx_to_pdf_bytes_word(content)
    if word_pdf is not None:
        try:
            import pymupdf
        except ImportError as error:
            raise RuntimeError("pymupdf недоступен для рендера страниц Word.") from error

        document = pymupdf.open(stream=word_pdf, filetype="pdf")
        try:
            copied: list[Path] = []
            for page_number, page in enumerate(document, start=1):
                pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
                target = output_dir / f"page_{page_number:02d}.png"
                pixmap.save(str(target))
                copied.append(target)
            return tuple(copied)
        finally:
            document.close()

    soffice = find_libreoffice()
    if soffice is None:
        raise RuntimeError("Ни Microsoft Word, ни LibreOffice недоступны для visual QA.")

    temp_path = Path(tempfile.mkdtemp(prefix="calendar_pedagoga_qa_"))
    try:
        docx_path = temp_path / "calendar.docx"
        docx_path.write_bytes(content)
        pdf_dir = temp_path / "pdf_full"
        pdf_dir.mkdir()

        pdf_result = _run_soffice(
            soffice,
            [
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_dir),
                str(docx_path),
            ],
            temp_path,
        )

        pdfs = sorted(pdf_dir.glob("*.pdf"))
        if not pdfs:
            _raise_missing_libreoffice_output(
                soffice,
                pdf_result,
                "LibreOffice не создал PDF для visual QA",
            )
        pdf_path = pdfs[0]
        document = None
        page_count: int | None = None
        try:
            import pymupdf

            document = pymupdf.open(stream=pdf_path.read_bytes(), filetype="pdf")
            page_count = document.page_count
            if page_count <= 0:
                raise RuntimeError("PDF не содержит страниц.")

            copied: list[Path] = []
            for page_number, page in enumerate(document, start=1):
                pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
                target = output_dir / f"page_{page_number:02d}.png"
                pixmap.save(str(target))
                copied.append(target)
            return tuple(copied)
        except Exception as error:
            _log_libreoffice_failure(
                soffice,
                list(pdf_result.args),
                return_code=pdf_result.returncode,
                stdout=pdf_result.stdout,
                stderr=pdf_result.stderr,
                pdf_path=pdf_path,
                pdf_page_count=page_count,
                pymupdf_error=error,
            )
            raise RuntimeError("PyMuPDF не отрендерил PDF; diagnostics logged.") from error
        finally:
            if document is not None:
                document.close()
    finally:
        shutil.rmtree(temp_path, ignore_errors=True)


def analyze_visual_pages(png_paths: tuple[Path, ...]) -> tuple[VisualPageReport, ...]:
    reports: list[VisualPageReport] = []
    for index, path in enumerate(png_paths, start=1):
        width, height = _png_dimensions(path)
        reports.append(
            VisualPageReport(
                page_number=index,
                path=path,
                width=width,
                height=height,
                size_bytes=path.stat().st_size,
                ink_ratio=_png_ink_ratio(path),
            )
        )
    return tuple(reports)


def validate_calendar_docx_visual(content: bytes) -> tuple[QAIssue, ...]:
    """Visual QA: отрендерить все страницы и проверить читаемость/целостность."""

    issues: list[QAIssue] = []
    if not find_microsoft_word() and find_libreoffice() is None:
        return (
            QAIssue(
                QASeverity.ERROR,
                "Ни Microsoft Word, ни LibreOffice недоступны; visual QA невозможен.",
            ),
        )

    with tempfile.TemporaryDirectory(prefix="calendar_pedagoga_visual_") as temp_name:
        try:
            png_paths = render_docx_pages(content, Path(temp_name))
            reports = analyze_visual_pages(png_paths)
        except RuntimeError as error:
            return (QAIssue(QASeverity.ERROR, str(error)),)

        if not reports:
            return (QAIssue(QASeverity.ERROR, "Не удалось отрендерить страницы DOCX."),)

        for report in reports:
            if report.size_bytes < _MIN_PAGE_FILE_BYTES:
                issues.append(
                    QAIssue(
                        QASeverity.ERROR,
                        f"Страница {report.page_number}: подозрение на пустой/обрезанный рендер.",
                    )
                )
            if report.width < _MIN_PAGE_DIMENSION or report.height < _MIN_PAGE_DIMENSION:
                issues.append(
                    QAIssue(
                        QASeverity.ERROR,
                        f"Страница {report.page_number}: некорректный размер изображения.",
                    )
                )
            if report.ink_ratio < _MIN_INK_RATIO:
                issues.append(
                    QAIssue(
                        QASeverity.ERROR,
                        f"Страница {report.page_number}: недостаточно видимого содержания.",
                    )
                )

        first = reports[0]
        if first.ink_ratio < 0.02:
            issues.append(
                QAIssue(
                    QASeverity.ERROR,
                    "Страница 1: шапка/таблица не видны на рендере.",
                )
            )
        last = reports[-1]
        if last.ink_ratio < _MIN_INK_RATIO:
            issues.append(
                QAIssue(
                    QASeverity.ERROR,
                    "Последняя страница: таблица не читается или обрезана.",
                )
            )

    return tuple(issues)
