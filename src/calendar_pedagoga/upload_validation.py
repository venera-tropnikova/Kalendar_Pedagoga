"""Безопасная валидация пользовательских загрузок до дальнейшей обработки."""

from __future__ import annotations

from dataclasses import dataclass
from enum import StrEnum
from io import BytesIO
from pathlib import Path
import zipfile

from docx import Document

from calendar_pedagoga.parsing import UtpParseResult, parse_utp
from calendar_pedagoga.program_parsing import ProgramData, parse_program


MAX_UPLOAD_BYTES = 10 * 1024 * 1024
_DOCX_MAGIC = b"PK\x03\x04"
_DOC_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")


class UploadPurpose(StrEnum):
    UTP = "utp"
    PROGRAM = "program"
    CALENDAR_TEMPLATE = "calendar_template"


class UploadValidationError(ValueError):
    """Загрузка отклонена до сохранения и передачи последующим обработчикам."""


@dataclass(frozen=True)
class ValidatedUpload:
    purpose: UploadPurpose
    filename: str
    content: bytes
    parsed: UtpParseResult | ProgramData | None = None


_ALLOWED_EXTENSIONS = {
    UploadPurpose.UTP: frozenset({".docx"}),
    UploadPurpose.PROGRAM: frozenset({".doc", ".docx"}),
    UploadPurpose.CALENDAR_TEMPLATE: frozenset({".docx"}),
}


def _validate_container(filename: str, data: bytes, purpose: UploadPurpose) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in _ALLOWED_EXTENSIONS[purpose]:
        allowed = ", ".join(sorted(_ALLOWED_EXTENSIONS[purpose]))
        raise UploadValidationError(
            f"Неправильный формат: для файла «{filename}» разрешены только {allowed}."
        )
    if not data:
        raise UploadValidationError(f"Файл пустой: «{filename}». Выберите непустой документ.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise UploadValidationError(
            f"Файл слишком большой: «{filename}». Максимальный размер — 10 МБ."
        )
    expected_magic = _DOC_MAGIC if suffix == ".doc" else _DOCX_MAGIC
    if not data.startswith(expected_magic):
        raise UploadValidationError(
            f"Неправильный формат: содержимое файла «{filename}» не соответствует расширению {suffix}."
        )
    if suffix == ".docx":
        try:
            with zipfile.ZipFile(BytesIO(data)) as archive:
                names = set(archive.namelist())
                if not {"[Content_Types].xml", "word/document.xml"} <= names:
                    raise UploadValidationError(
                        f"Файл повреждён: «{filename}» не является корректным документом Word."
                    )
                if archive.testzip() is not None:
                    raise UploadValidationError(
                        f"Файл повреждён: «{filename}» содержит повреждённые данные."
                    )
        except (zipfile.BadZipFile, RuntimeError, OSError) as error:
            raise UploadValidationError(
                f"Файл повреждён: «{filename}» не является корректным DOCX."
            ) from error
    return suffix


def _docx_text(data: bytes) -> str:
    try:
        document = Document(BytesIO(data))
    except Exception as error:
        raise UploadValidationError("Файл повреждён: документ Word не может быть прочитан.") from error
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return " ".join(values).casefold()


def _looks_normative(text: str) -> bool:
    markers = (
        "федеральный закон",
        "постановление правительства",
        "приказ министерства",
        "зарегистрировано в министерстве юстиции",
        "нормативный правовой акт",
    )
    return any(marker in text for marker in markers)


def _reject_normative(text: str, filename: str) -> None:
    if _looks_normative(text):
        raise UploadValidationError(
            f"Нормативный документ загружать не нужно: файл «{filename}» "
            "не предназначен для пользовательских загрузчиков."
        )


def _validate_utp(filename: str, data: bytes) -> UtpParseResult:
    text = _docx_text(data)
    _reject_normative(text, filename)
    try:
        result = parse_utp(data)
    except Exception as error:
        raise UploadValidationError(
            f"Неправильный тип документа: файл «{filename}» не распознан как УТП — не найдена корректная таблица тем и часов."
        ) from error
    if not result.topics or not result.sections:
        raise UploadValidationError(
            f"Неправильный тип документа: файл «{filename}» не распознан как УТП — отсутствуют разделы или учебные темы."
        )
    return result


def _validate_program(filename: str, data: bytes) -> ProgramData:
    if Path(filename).suffix.lower() == ".docx":
        _reject_normative(_docx_text(data), filename)
    try:
        result = parse_program(data, filename)
    except UploadValidationError:
        raise
    except Exception as error:
        raise UploadValidationError(
            f"Неправильный тип документа: файл «{filename}» не удалось прочитать как образовательную программу."
        ) from error
    if not result.content_items or not (result.title or result.goal or result.tasks):
        raise UploadValidationError(
            f"Неправильный тип документа: файл «{filename}» не распознан как образовательная программа — "
            "не найдены программные разделы и содержательная часть."
        )
    return result


def _validate_calendar_template(filename: str, data: bytes) -> None:
    text = _docx_text(data)
    _reject_normative(text, filename)
    markers = (
        "календарный план",
        "теоретические занятия",
        "практические занятия",
        "тип занятия",
        "планируемый результат",
        "вид контроля",
    )
    if sum(marker in text for marker in markers) < 3:
        raise UploadValidationError(
            f"Неправильный тип документа: файл «{filename}» не распознан как форма календарного плана — "
            "не найдены ожидаемые колонки календаря."
        )


def validate_upload(
    purpose: UploadPurpose,
    filename: str,
    data: bytes,
) -> ValidatedUpload:
    """Проверить файл полностью в памяти до его передачи приложению или AI."""

    safe_name = Path(filename).name
    _validate_container(safe_name, data, purpose)
    parsed: UtpParseResult | ProgramData | None = None
    if purpose is UploadPurpose.UTP:
        parsed = _validate_utp(safe_name, data)
    elif purpose is UploadPurpose.PROGRAM:
        parsed = _validate_program(safe_name, data)
    else:
        _validate_calendar_template(safe_name, data)
    return ValidatedUpload(purpose, safe_name, bytes(data), parsed)
