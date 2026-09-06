"""Пользовательский интерфейс формирования календарного плана."""

from __future__ import annotations

import calendar
import html
import hashlib
import json
import logging
import re
from collections.abc import Callable
from copy import deepcopy
from io import BytesIO
from zipfile import BadZipFile
from datetime import date, timedelta
from pathlib import Path

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from calendar_pedagoga.academic_year import (
    APPROVED_ACADEMIC_YEAR,
    AcademicYearResolution,
    AcademicYearStatus,
    academic_year_period,
    academic_year_start,
    default_academic_year_start,
    format_academic_year,
    mentions_from_program,
    mentions_from_utp,
    resolve_academic_year,
)
from calendar_pedagoga.content_engine_v2 import build_lesson_content_v2
from calendar_pedagoga.content_generation import CalendarContentRow, build_content_model
from calendar_pedagoga.lesson_content import LessonContentRow, build_lesson_content
from calendar_pedagoga.normative_engine import (
    NormativeCheck,
    NormativeLayer,
    NormativeLessonView,
    NormativeReport,
    NormativeVerdict,
    evaluate_normative_mvp,
)
from calendar_pedagoga.normative_registry import (
    CalendarRegistryReference,
    NormativeUpdateChoice,
    get_builtin_normative_registry,
    get_update_notice,
    resolve_registry_reference,
)
from calendar_pedagoga.organization_template import (
    ORG_TEMPLATE_UNSUPPORTED_MESSAGE,
    CalendarTemplateSelection,
    OrganizationTemplateError,
    select_calendar_template,
)
from calendar_pedagoga.pipeline import PipelineError, run_calendar_pipeline
from calendar_pedagoga.docx_generation import (
    _allow_row_split,
    _columns_for_table,
    _merge_month_cells_by_page_segments,
    _prevent_row_split,
    _remove_vmerge,
    _save_document,
)
from calendar_pedagoga.docx_qa import detect_data_row_page_spans
from calendar_pedagoga.practice_slots import SLOT_CONTINUE_WARNING, SLOT_PACK_WARNING
from calendar_pedagoga.resolve_utp import UtpResolutionError, resolve_utp
from calendar_pedagoga.transient_documents import TransientDocumentSession
from calendar_pedagoga.upload_validation import (
    UploadPurpose,
    UploadValidationError,
    ValidatedUpload,
    validate_upload,
)
from calendar_pedagoga.parsing import UtpParseResult, parse_utp
from calendar_pedagoga.matching import ContentMatch, match_utp_to_program
from calendar_pedagoga.program_parsing import (
    ProgramData,
    infer_study_year_number,
    parse_program,
    study_year_label,
)
from calendar_pedagoga.scheduling import (
    ScheduleResult,
    ScheduleValidationError,
    build_academic_weeks,
    build_schedule,
)


def _generator_revision() -> str:
    """Read the working-tree generator identity, not a cached Git revision."""
    root = Path(__file__).resolve().parents[2]
    paths = [root / "app.py", *sorted((root / "src" / "calendar_pedagoga").glob("*.py"))]
    paths.append(root / "references" / "Календарный план Образец.docx")
    digest = hashlib.sha256()
    for path in paths:
        digest.update(str(path.relative_to(root)).encode("utf-8"))
        digest.update(hashlib.sha256(path.read_bytes()).digest())
    return digest.hexdigest()


_LOADED_GENERATOR_REVISION = _generator_revision()


def _inputs_fingerprint(*values: object) -> str:
    parts = []
    for value in values:
        if value is not None and hasattr(value, "getvalue"):
            parts.append([value.name, hashlib.sha256(value.getvalue()).hexdigest()])
        else:
            parts.append(value)
    return hashlib.sha256(json.dumps(parts, ensure_ascii=False).encode("utf-8")).hexdigest()


def _sync_generation_fingerprint(fingerprint: tuple[str, str]) -> bool:
    """Invalidate only; never generate on a rerun or accept an unversioned result."""
    previous = st.session_state.get("calendar_generation_fingerprint")
    if previous == fingerprint:
        return False
    keys = (
        "calendar_download", "calendar_warnings", "calendar_ai_usage",
        "calendar_generation_pending", "calendar_generation_error",
        "calendar_generation_succeeded", "calendar_resolved_lessons",
        "calendar_plan_snapshot",
    )
    if any(st.session_state.get(key) for key in keys):
        st.session_state["calendar_generation_invalidated"] = True
    for key in keys:
        st.session_state.pop(key, None)
    st.session_state["calendar_generation_fingerprint"] = fingerprint
    return True


def _refresh_generation_inputs(
    utp_file, program_file, template_file, academic_year, group_number, class_name,
    teacher_name,
) -> None:
    revision = _generator_revision()
    analysis = _inputs_fingerprint(utp_file, program_file, template_file, academic_year)
    inputs = _inputs_fingerprint(analysis, group_number, class_name, teacher_name)
    _sync_generation_fingerprint((inputs, revision))
    analysis_fingerprint = (analysis, revision)
    if st.session_state.get("calendar_analysis_fingerprint") != analysis_fingerprint:
        _reset_analysis_state()
    st.session_state["calendar_analysis_fingerprint"] = analysis_fingerprint
    st.session_state["calendar_generation_inputs"] = inputs


def _reset_analysis_state() -> None:
    if st.session_state.get("calendar_download") or st.session_state.get("calendar_generation_succeeded"):
        st.session_state["calendar_generation_invalidated"] = True
    st.session_state["analysis_ready"] = False
    for key in (
        "analysis_warnings",
        "calendar_download",
        "calendar_warnings",
        "calendar_ai_usage",
        "calendar_generation_pending",
        "calendar_generation_error",
        "calendar_generation_succeeded",
        "calendar_resolved_lessons",
        "calendar_plan_snapshot",
        "calendar_context",
    ):
        st.session_state.pop(key, None)


def _clear_upload_slot(slot: str) -> None:
    nonce_key = f"upload_nonce_{slot}"
    st.session_state[nonce_key] = int(st.session_state.get(nonce_key, 0)) + 1
    st.session_state.pop(f"upload_name_{slot}", None)
    _reset_analysis_state()
    st.rerun()


def _remember_upload(slot: str, uploaded: object | None) -> None:
    name_key = f"upload_name_{slot}"
    previous = st.session_state.get(name_key)
    current = getattr(uploaded, "name", None)
    if previous is not None and current is None:
        st.session_state.pop(name_key, None)
        if st.session_state.get("analysis_ready"):
            _reset_analysis_state()
    elif current is not None:
        st.session_state[name_key] = current


def _file_uploader_with_clear(
    slot: str,
    *,
    label: str,
    type: tuple[str, ...],
    help: str,
) -> object | None:
    nonce = int(st.session_state.setdefault(f"upload_nonce_{slot}", 0))
    uploaded = st.file_uploader(
        label,
        type=type,
        help=help,
        label_visibility="collapsed",
        key=f"upload_{slot}_{nonce}",
    )
    if uploaded is not None:
        name_col, clear_col = st.columns([0.88, 0.12])
        with name_col:
            st.markdown(
                f'<p class="kp-uploaded-name">{html.escape(uploaded.name)}</p>',
                unsafe_allow_html=True,
            )
        with clear_col:
            if st.button(
                "×",
                key=f"clear_{slot}",
                help="Удалить файл",
                use_container_width=True,
            ):
                _clear_upload_slot(slot)
    _remember_upload(slot, uploaded)
    return uploaded


_MONTH_TITLES = {
    1: "Январь",
    2: "Февраль",
    3: "Март",
    4: "Апрель",
    5: "Май",
    6: "Июнь",
    7: "Июль",
    8: "Август",
    9: "Сентябрь",
    10: "Октябрь",
    11: "Ноябрь",
    12: "Декабрь",
}
_WEEKDAYS_RU = ("Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс")
_MONTHS_GENITIVE = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}
_RECOMMENDED_BREAK_PERIODS = {
    APPROVED_ACADEMIC_YEAR: (
        ("Осенние каникулы", date(2026, 10, 26), date(2026, 11, 3), False),
        ("Зимние каникулы", date(2026, 12, 31), date(2027, 1, 10), False),
        ("Дополнительные каникулы для 1 класса", date(2027, 2, 15), date(2027, 2, 21), True),
        ("Весенние каникулы", date(2027, 3, 27), date(2027, 4, 4), False),
        ("Летние каникулы", date(2027, 5, 27), date(2027, 8, 31), False),
    ),
}


def _is_first_class(value: str) -> bool:
    normalized = " ".join(value.casefold().replace("класс", "").split())
    return normalized in {"1", "1а", "1б", "1в", "1г", "1д"}


def _recommended_school_breaks(
    academic_year: str,
    class_name: str = "",
) -> dict[date, str]:
    dates: dict[date, str] = {}
    for title, start, end, first_class_only in _RECOMMENDED_BREAK_PERIODS.get(
        academic_year, ()
    ):
        if first_class_only and not _is_first_class(class_name):
            continue
        cursor = start
        while cursor <= end:
            dates[cursor] = title
            cursor += timedelta(days=1)
    return dates


def _recommended_break_title(day: date, class_name: str = "") -> str | None:
    return _recommended_school_breaks(
        APPROVED_ACADEMIC_YEAR,
        class_name,
    ).get(day)


def _format_recommended_break_period(start: date, end: date) -> str:
    if start.year != end.year:
        return (
            f"{start.day} {_MONTHS_GENITIVE[start.month]} {start.year} — "
            f"{end.day} {_MONTHS_GENITIVE[end.month]} {end.year}"
        )
    if start.month == end.month:
        return f"{start.day}–{end.day} {_MONTHS_GENITIVE[start.month]} {end.year}"
    return (
        f"{start.day} {_MONTHS_GENITIVE[start.month]} — "
        f"{end.day} {_MONTHS_GENITIVE[end.month]} {end.year}"
    )


def _recommended_break_cards_html(
    academic_year: str,
    class_name: str = "",
) -> str:
    periods = _RECOMMENDED_BREAK_PERIODS.get(academic_year, ())
    cards = "".join(
        '<article class="kp-cal-vacation-card">'
        f'<strong>{html.escape(title)}</strong>'
        f'<span>{html.escape(_format_recommended_break_period(start, end))}</span>'
        '</article>'
        for title, start, end, first_class_only in periods
        if not first_class_only
    )
    first_class_note = ""
    if _is_first_class(class_name):
        for title, start, end, first_class_only in periods:
            if first_class_only:
                first_class_note = (
                    '<p class="kp-cal-vacation-extra">'
                    f'{html.escape(title)}: '
                    f'{html.escape(_format_recommended_break_period(start, end))}'
                    '</p>'
                )
                break
    return (
        '<section class="kp-cal-vacations">'
        '<h3>Рекомендуемые школьные каникулы</h3>'
        f'<div class="kp-cal-vacation-grid">{cards}</div>'
        f'{first_class_note}'
        '</section>'
    )

_OFFICIAL_HOLIDAYS = {
    # Постановление Правительства РФ от 24.09.2025 № 1466.
    date(2026, 12, 31): "Перенесённый выходной день",
    # Статья 112 ТК РФ: День народного единства.
    date(2026, 11, 4): "День народного единства",
    # Статья 112 ТК РФ: фиксированные нерабочие праздничные дни 2027 года.
    date(2027, 1, 1): "Новогодние каникулы",
    date(2027, 1, 2): "Новогодние каникулы",
    date(2027, 1, 3): "Новогодние каникулы",
    date(2027, 1, 4): "Новогодние каникулы",
    date(2027, 1, 5): "Новогодние каникулы",
    date(2027, 1, 6): "Новогодние каникулы",
    date(2027, 1, 7): "Рождество Христово",
    date(2027, 1, 8): "Новогодние каникулы",
    date(2027, 2, 23): "День защитника Отечества",
    date(2027, 3, 8): "Международный женский день",
    date(2027, 5, 1): "Праздник Весны и Труда",
    date(2027, 5, 9): "День Победы",
    date(2027, 6, 12): "День России",
}
_PROFESSIONAL_DATES = {
    date(2026, 10, 5): "5 октября — День учителя",
}
_OFFICIAL_HOLIDAY_SOURCES = (
    (
        "Статья 112 ТК РФ: нерабочие праздничные дни",
        "https://www.consultant.ru/document/cons_doc_LAW_34683/98ef2900507766e70ff29c0b9d8e2353ea80a1cf/",
    ),
    (
        "Постановление Правительства РФ № 1466: переносы 2026",
        "https://government.ru/docs/all/161028/",
    ),
)
_TEACHER_LINKS = (
    (
        "Рекомендации Минпросвещения по каникулам",
        "https://edu.gov.ru/press/11820/minprosvescheniya-rossii-napravilo-v-regiony-rekomendacii-po-grafiku-kanikul-v-2026-27-uchebnomu-godu/",
    ),
    (
        "Расписание школьных каникул 2026–2027",
        "https://t-j.ru/guide/ura-kanikuli/",
    ),
)


def _academic_period_caption(academic_year: str) -> str:
    period = academic_year_period(academic_year)
    if period is None:
        return ""
    start, end = period
    return (
        f"{_MONTH_TITLES[start.month]} {start.year} — "
        f"{_MONTH_TITLES[end.month].lower()} {end.year}"
    )


def _academic_day_sets(
    academic_year: str,
    class_name: str = "",
) -> tuple[set[date], set[date], set[date]]:
    weeks = build_academic_weeks(academic_year)
    study: set[date] = set()
    short: set[date] = set()
    approved = academic_year == APPROVED_ACADEMIC_YEAR
    for week in weeks:
        cursor = week.start
        week_is_short = approved and (week.end - week.start).days < 6
        while cursor <= week.end:
            study.add(cursor)
            if week_is_short:
                short.add(cursor)
            cursor += timedelta(days=1)
    breaks: set[date] = set()
    if approved:
        for previous, nxt in zip(weeks, weeks[1:]):
            cursor = previous.end + timedelta(days=1)
            last = nxt.start - timedelta(days=1)
            while cursor <= last:
                breaks.add(cursor)
                cursor += timedelta(days=1)
    breaks.update(_recommended_school_breaks(academic_year, class_name))
    return study, short, breaks


def _calendar_day_class(
    day: date,
    study: set[date],
    _short: set[date],
    breaks: set[date],
) -> str:
    if day in _OFFICIAL_HOLIDAYS:
        return "kp-cal-holiday"
    if day in _PROFESSIONAL_DATES:
        return "kp-cal-professional"
    if day in breaks and day.weekday() >= 5:
        return "kp-cal-break kp-cal-wknd"
    if day in breaks:
        return "kp-cal-break"
    if day.weekday() >= 5:
        return "kp-cal-wknd"
    if day in study:
        return "kp-cal-study"
    return ""


def _academic_calendar_html(academic_year: str) -> str:
    period = academic_year_period(academic_year)
    if period is None:
        return ""
    study, short, breaks = _academic_day_sets(academic_year)
    month_calendar = calendar.Calendar(firstweekday=0)
    blocks: list[str] = []
    year, month = period[0].year, period[0].month
    last = date(period[1].year, period[1].month, 1)
    while date(year, month, 1) <= last:
        rows: list[str] = []
        for week in month_calendar.monthdayscalendar(year, month):
            cells: list[str] = []
            for day_number in week:
                if day_number == 0:
                    cells.append("<td></td>")
                    continue
                day = date(year, month, day_number)
                kind = _calendar_day_class(day, study, short, breaks)
                attr = f' class="{kind}"' if kind else ""
                cells.append(f"<td{attr}>{day_number}</td>")
            rows.append("<tr>" + "".join(cells) + "</tr>")
        head = "".join(f"<th>{name}</th>" for name in _WEEKDAYS_RU)
        blocks.append(
            '<div class="kp-cal-month">'
            f'<p class="kp-cal-month-title">{_MONTH_TITLES[month]} {year}</p>'
            f"<table><thead><tr>{head}</tr></thead>"
            f"<tbody>{''.join(rows)}</tbody></table></div>"
        )
        month += 1
        if month > 12:
            month = 1
            year += 1
    legend = (
        '<div class="kp-cal-legend">'
        '<span><i class="kp-cal-study"></i>Учебные дни</span>'
        '<span><i class="kp-cal-wknd"></i>Выходные</span>'
        '<span><i class="kp-cal-break"></i>Перерыв</span>'
        '<span><i class="kp-cal-week"></i>№ — номер учебной недели</span>'
        "</div>"
    )
    return '<div class="kp-cal-grid">' + "".join(blocks) + "</div>" + legend


def _week_date_caption(start: date, end: date) -> str:
    if start.month == end.month:
        return f"{start.day}–{end.day} {_MONTHS_GENITIVE[start.month]}"
    return (
        f"{start.day} {_MONTHS_GENITIVE[start.month]} — "
        f"{end.day} {_MONTHS_GENITIVE[end.month]}"
    )


def _short_weeks_note(academic_year: str) -> str:
    short = tuple(
        week for week in build_academic_weeks(academic_year)
        if (week.end - week.start).days < 6
    )
    if not short:
        return "Коротких недель в текущей сетке нет."
    details = "; ".join(
        f"№{week.number} — {_week_date_caption(week.start, week.end)}"
        for week in short
    )
    return f"Короткие недели: {details}."


def _week_topic_caption(row: object) -> str:
    source = row.source.source
    if source.week_parts:
        return "; ".join(
            f"{part.topic_number}. {part.topic_title}"
            if part.topic_number else part.topic_title
            for part in source.week_parts
        )
    return (
        f"{source.topic_number}. {source.topic_title}"
        if source.topic_number else source.topic_title
    )


def _week_detail_cell(value: object) -> str:
    return html.escape(str(value or "")).replace("\n", "<br>")


def _week_detail_rows(
    rows: tuple[object, ...],
    week_number: int,
) -> tuple[tuple[str, ...], ...]:
    """Проекция уже сформированных строк плана без повторной генерации."""

    return tuple(
        (
            row.source.source.date_range,
            _week_topic_caption(row),
            row.theory_text,
            row.practice_text,
            row.lesson_type,
            row.planned_result,
            row.assessment_method,
            "",
        )
        for row in rows
        if row.source.source.week_number == week_number
    )


def _calendar_plan_snapshot(
    rows: tuple[object, ...],
    docx_content: bytes,
) -> dict[int, tuple[tuple[str, ...], ...]]:
    """Снимок фактических ячеек того же результата, который отдан на скачивание."""

    fallback = {
        number: _week_detail_rows(rows, number)
        for number in {row.source.source.week_number for row in rows}
    }
    if not rows:
        return fallback
    try:
        table = Document(BytesIO(docx_content)).tables[0]
        columns = _columns_for_table(table)
        snapshot: dict[int, list[tuple[str, ...]]] = {}
        for index, lesson in enumerate(rows):
            cells = table.rows[index + 2].cells
            week_cell_lines = cells[columns.week].text.splitlines()
            date_value = "\n".join(week_cell_lines[1:]).strip()
            source = lesson.source.source
            snapshot.setdefault(source.week_number, []).append(
                (
                    date_value or source.date_range,
                    _week_topic_caption(lesson),
                    cells[columns.theory].text.strip(),
                    cells[columns.practice].text.strip(),
                    cells[columns.lesson_type].text.strip(),
                    cells[columns.planned_result].text.strip(),
                    cells[columns.assessment].text.strip(),
                    cells[columns.theory_mark].text.strip(),
                )
            )
        return {number: tuple(items) for number, items in snapshot.items()}
    except (BadZipFile, PackageNotFoundError, IndexError):
        return fallback

def _render_week_detail(academic_year: str, week_number: int) -> None:
    week = next(
        (item for item in build_academic_weeks(academic_year) if item.number == week_number),
        None,
    )
    if week is None:
        return
    snapshot = st.session_state.get("calendar_plan_snapshot") or {}
    week_rows = tuple(snapshot.get(week_number) or ())
    st.markdown(
        f'<p class="kp-week-title">Неделя №{week.number} · '
        f'{html.escape(_week_date_caption(week.start, week.end))}</p>',
        unsafe_allow_html=True,
    )
    if not week_rows:
        st.info("Сформируйте календарный план, чтобы посмотреть содержание недели.")
    else:
        headers = (
            "Дата",
            "Номер и название темы",
            "Теоретические занятия",
            "Практические занятия",
            "Тип занятия",
            "Планируемый результат",
            "Вид контроля усвоения содержания темы",
            "Отметка о проведении (дата, кол-во выраб. часов)",
        )
        head = "".join(f"<th>{html.escape(label)}</th>" for label in headers)
        body = "".join(
            "<tr>"
            + "".join(
                f"<td>{_week_detail_cell(value)}</td>"
                for value in row
            )
            + "</tr>"
            for row in week_rows
        )
        st.markdown(
            '<p class="kp-week-source">'
            "Фактические строки сформированного календарного плана"
            "</p>"
            '<div class="kp-week-table-scroll">'
            '<table class="kp-week-table">'
            f"<thead><tr>{head}</tr></thead>"
            f"<tbody>{body}</tbody></table></div>",
            unsafe_allow_html=True,
        )
    if st.button("← К календарю", key="kp_week_back"):
        st.session_state.pop("kp_selected_week", None)
        st.rerun()


def _month_weeks(
    academic_year: str,
    year: int,
    month: int,
) -> tuple[object, ...]:
    """Недели месяца по той же принадлежности, что в календарной сетке."""

    return tuple(
        week
        for week in build_academic_weeks(academic_year)
        if week.start.year == year and week.start.month == month
    )


def _month_detail_rows(
    snapshot: dict[int, tuple[tuple[str, ...], ...]],
    weeks: tuple[object, ...],
) -> tuple[tuple[str, ...], ...]:
    """Проекция готовых DOCX-строк месяца без повторной генерации."""

    return tuple(
        (f"№{week.number}", *row)
        for week in weeks
        for row in tuple(snapshot.get(week.number) or ())
    )


def _monthly_plan_docx(
    annual_content: bytes,
    academic_year: str,
    year: int,
    month: int,
) -> bytes:
    """Выделить строки месяца из готового годового DOCX без пересчёта."""

    from docx.table import _Cell

    selected_weeks = {
        week.number for week in _month_weeks(academic_year, year, month)
    }
    if not selected_weeks:
        raise ValueError("В выбранном месяце нет строк календарного плана.")

    document = Document(BytesIO(annual_content))
    if not document.tables:
        raise ValueError("В готовом календарном плане не найдена таблица.")
    table = document.tables[0]
    columns = _columns_for_table(table)
    kept_rows = 0
    # Resolve inherited identifiers before deleting a possible merge anchor.
    original_rows = [
        (row, deepcopy(row.cells[columns.month]._tc), deepcopy(row.cells[columns.week]._tc))
        for row in table.rows[2:]
    ]
    months = []
    for row, month_cell, week_cell in original_rows:
        first_line = _Cell(week_cell, row).text.splitlines()[:1]
        try:
            week_number = int(first_line[0].strip()) if first_line else None
        except ValueError:
            week_number = None
        if week_number in selected_weeks:
            for column, restored in ((columns.month, month_cell), (columns.week, week_cell)):
                _remove_vmerge(_Cell(restored, row))
                row._tr.replace(row._tr.tc_lst[column], restored)
            months.append(row.cells[columns.month].text)
            _allow_row_split(row)
            kept_rows += 1
            continue
        table._tbl.remove(row._tr)

    if kept_rows == 0:
        raise ValueError("В готовом календарном плане нет строк выбранного месяца.")
    unmerged = _save_document(document)

    def fallback():
        # Missing/unstable pagination is not permission to reuse annual merges.
        for row in table.rows[2:]:
            _prevent_row_split(row)
        return _save_document(document)

    spans = detect_data_row_page_spans(unmerged, total_rows=kept_rows)
    keep_together = set()
    for _ in range(kept_rows + 2):
        if not spans:
            return fallback()
        pages = {}
        for index, span in enumerate(spans):
            if span.start_page == span.end_page:
                pages.setdefault(span.start_page, []).append(index)
            elif not span.split_safe:
                keep_together.add(index)
        candidate = Document(BytesIO(unmerged))
        candidate_table = candidate.tables[0]
        for index in keep_together | {index for indices in pages.values() for index in indices}:
            _prevent_row_split(candidate_table.rows[index + 2])
        # Week/date cells stay independent; never inherit a cross-row merge.
        _merge_month_cells_by_page_segments(
            candidate_table, columns, tuple(months),
            tuple(tuple(indices) for indices in pages.values()),
        )
        result = _save_document(candidate)
        detected = detect_data_row_page_spans(result, total_rows=kept_rows)
        if detected == spans:
            if any(span.start_page != span.end_page and not span.split_safe for span in detected):
                return fallback()
            return result
        spans = detected
    return fallback()


def _adjacent_academic_month(
    academic_year: str,
    year: int,
    month: int,
    delta: int,
) -> tuple[int, int] | None:
    period = academic_year_period(academic_year)
    if period is None:
        return None
    index = year * 12 + month - 1 + delta
    target = (index // 12, index % 12 + 1)
    first = (period[0].year, period[0].month)
    last = (period[1].year, period[1].month)
    return target if first <= target <= last else None


def _render_month_detail(academic_year: str, year: int, month: int) -> None:
    month_title = f"{_MONTH_TITLES[month]} {year}"
    weeks = _month_weeks(academic_year, year, month)
    st.markdown(
        f'<p class="kp-week-title">{html.escape(month_title)}</p>',
        unsafe_allow_html=True,
    )
    if weeks:
        st.markdown(
            f'<p class="kp-month-weeks">Недели №{weeks[0].number}–№{weeks[-1].number}</p>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<p class="kp-month-weeks">Нет недель календарного плана</p>',
            unsafe_allow_html=True,
        )

    snapshot = st.session_state.get("calendar_plan_snapshot") or {}
    annual_download = st.session_state.get("calendar_download")
    month_rows = _month_detail_rows(snapshot, weeks)
    if not snapshot or annual_download is None:
        st.info("Сначала сформируйте календарный план.")
    elif not month_rows:
        st.info("В этом месяце нет строк сформированного календарного плана.")
    else:
        monthly_content = _monthly_plan_docx(
            annual_download.content,
            academic_year,
            year,
            month,
        )
        st.download_button(
            f"Скачать план на {_MONTH_TITLES[month].lower()} {year}",
            data=monthly_content,
            file_name=f"Календарный_план_{_MONTH_TITLES[month]}_{year}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"kp_month_download_{academic_year}_{year}_{month}",
        )
        headers = (
            "Неделя",
            "Дата",
            "Номер и название темы",
            "Теоретические занятия",
            "Практические занятия",
            "Тип занятия",
            "Планируемый результат",
            "Вид контроля усвоения содержания темы",
            "Отметка о проведении",
        )
        head = "".join(f"<th>{html.escape(label)}</th>" for label in headers)
        body = "".join(
            "<tr>"
            + "".join(f"<td>{_week_detail_cell(value)}</td>" for value in row)
            + "</tr>"
            for row in month_rows
        )
        st.markdown(
            '<p class="kp-week-source">'
            "Фактические строки сформированного календарного плана"
            "</p>"
            '<div class="kp-week-table-scroll kp-month-table-scroll">'
            '<table class="kp-week-table kp-month-table">'
            f"<thead><tr>{head}</tr></thead>"
            f"<tbody>{body}</tbody></table></div>",
            unsafe_allow_html=True,
        )

    previous_month = _adjacent_academic_month(academic_year, year, month, -1)
    next_month = _adjacent_academic_month(academic_year, year, month, 1)
    back_column, previous_column, next_column = st.columns((1.25, 1, 1))
    with back_column:
        if st.button("← К календарю", key="kp_month_back", use_container_width=True):
            st.session_state.pop("kp_selected_month", None)
            st.rerun()
    with previous_column:
        if st.button(
            "Предыдущий месяц",
            key="kp_month_previous",
            disabled=previous_month is None,
            use_container_width=True,
        ):
            st.session_state["kp_selected_month"] = previous_month
            st.rerun()
    with next_column:
        if st.button(
            "Следующий месяц",
            key="kp_month_next",
            disabled=next_month is None,
            use_container_width=True,
        ):
            st.session_state["kp_selected_month"] = next_month
            st.rerun()

def _calendar_week_for_month_row(
    *,
    year: int,
    month: int,
    day_numbers: list[int],
    weeks: tuple[object, ...],
) -> object | None:
    """Вернуть учебную неделю, которая начинается в этой строке месяца."""

    dates = {
        date(year, month, day_number)
        for day_number in day_numbers
        if day_number
    }
    return next(
        (
            week
            for week in weeks
            if week.start.year == year
            and week.start.month == month
            and any(week.start <= day <= week.end for day in dates)
        ),
        None,
    )


def _calendar_day_row_html(
    *,
    year: int,
    month: int,
    day_numbers: list[int],
    study: set[date],
    short: set[date],
    breaks: set[date],
) -> str:
    cells: list[str] = []
    for day_number in day_numbers:
        if not day_number:
            cells.append('<span class="kp-cal-day kp-cal-day-empty"></span>')
            continue
        day = date(year, month, day_number)
        kind = _calendar_day_class(day, study, short, breaks)
        official_title = _OFFICIAL_HOLIDAYS.get(day)
        professional_title = _PROFESSIONAL_DATES.get(day)
        title = official_title or professional_title or _recommended_break_title(
            day, str(st.session_state.get("class_name") or "")
        )
        title_attr = f' title="{html.escape(title)}"' if title else ""
        if official_title:
            title_attr += f' data-holiday="{html.escape(official_title)}"'
        elif professional_title:
            title_attr += (
                f' data-professional="{html.escape(professional_title)}"'
            )
        cells.append(
            f'<span class="kp-cal-day {kind}"{title_attr}>{day_number}</span>'
        )
    return '<div class="kp-cal-days-row">' + "".join(cells) + "</div>"


_CALENDAR_WEEK_HEADER = "\N{NUMERO SIGN}"


def _calendar_days_header_html() -> str:
    """Стабильная пользовательская шапка без преобразования bool/None в текст."""

    return (
        '<div class="kp-cal-days-head" lang="ru" translate="no">'
        '<span class="kp-cal-week-head notranslate" translate="no" '
        'data-calendar-week-header="numero" aria-label="№ нед.">&#8470; нед.</span>'
        + "".join(
            f'<span class="{"kp-cal-weekend-head" if index >= 5 else ""}">{name}</span>'
            for index, name in enumerate(_WEEKDAYS_RU)
        )
        + "</div>"
    )


def _install_calendar_translate_guard() -> None:
    """Запретить автоперевод календаря и обернуть подписи недель."""

    components.html(
        """
        <script>
        (() => {
          const parentDocument = window.parent.document;
          const protect = () => {
            const dialog = parentDocument.querySelector('[role="dialog"]');
            if (!dialog) return;
            dialog.setAttribute('lang', 'ru');
            dialog.setAttribute('translate', 'no');
            dialog.classList.add('notranslate');
            dialog.querySelectorAll('[class*="st-key-kp_week_"] button p').forEach((label) => {
              if (label.querySelector('span.notranslate')) return;
              const span = parentDocument.createElement('span');
              span.className = 'notranslate';
              span.setAttribute('translate', 'no');
              span.textContent = label.textContent;
              label.replaceChildren(span);
            });
          };
          protect();
          new MutationObserver(protect).observe(parentDocument.body, {
            childList: true,
            subtree: true,
          });
        })();
        </script>
        """,
        height=0,
        scrolling=False,
    )


def _render_calendar_month(
    *,
    academic_year: str,
    year: int,
    month: int,
    weeks: tuple[object, ...],
    study: set[date],
    short: set[date],
    breaks: set[date],
) -> None:
    month_title = f"{_MONTH_TITLES[month]} {year}"
    if st.button(
        month_title,
        key=f"kp_calendar_month_title_{academic_year}_{year}_{month}",
        help="Открыть план на месяц",
        use_container_width=True,
    ):
        st.session_state["kp_selected_month"] = (year, month)
        st.rerun()
    st.markdown(
        _calendar_days_header_html(),
        unsafe_allow_html=True,
    )
    for row_index, day_numbers in enumerate(
        calendar.Calendar(firstweekday=0).monthdayscalendar(year, month)
    ):
        week = _calendar_week_for_month_row(
            year=year,
            month=month,
            day_numbers=day_numbers,
            weeks=weeks,
        )
        week_column, days_column = st.columns((0.12, 0.88), gap="small")
        with week_column:
            if week is None:
                st.markdown('<div class="kp-cal-week-empty"></div>', unsafe_allow_html=True)
            elif st.button(
                f"№{week.number}",
                key=f"kp_week_{academic_year}_{week.number}",
                help=f"Открыть содержание недели №{week.number}",
                use_container_width=True,
            ):
                st.session_state["kp_selected_week"] = week.number
                st.rerun()
        with days_column:
            st.markdown(
                _calendar_day_row_html(
                    year=year,
                    month=month,
                    day_numbers=day_numbers,
                    study=study,
                    short=short,
                    breaks=breaks,
                ),
                unsafe_allow_html=True,
            )


def _render_calendar_months(academic_year: str) -> None:
    period = academic_year_period(academic_year)
    if period is None:
        return
    study, short, breaks = _academic_day_sets(
        academic_year, str(st.session_state.get("class_name") or "")
    )
    weeks = tuple(build_academic_weeks(academic_year))
    year, month = period[0].year, period[0].month
    months: list[tuple[int, int]] = []
    while len(months) < 12:
        months.append((year, month))
        month += 1
        if month > 12:
            year, month = year + 1, 1
    for offset in range(0, 12, 3):
        columns = st.columns(3, gap="medium")
        for column, (year, month) in zip(columns, months[offset : offset + 3]):
            with column:
                st.markdown('<div class="kp-cal-month-anchor"></div>', unsafe_allow_html=True)
                _render_calendar_month(
                    academic_year=academic_year,
                    year=year,
                    month=month,
                    weeks=weeks,
                    study=study,
                    short=short,
                    breaks=breaks,
                )


@st.dialog("Календарь учебного года", width="large")
def _show_year_calendar_dialog() -> None:
    academic_year = str(st.session_state.get("kp_calendar_year") or "")
    if not academic_year:
        return
    _install_calendar_translate_guard()
    selected = st.session_state.get("kp_selected_week")
    if selected:
        _render_week_detail(academic_year, int(selected))
        return
    selected_month = st.session_state.get("kp_selected_month")
    if selected_month:
        selected_year, selected_month_number = selected_month
        _render_month_detail(
            academic_year,
            int(selected_year),
            int(selected_month_number),
        )
        return
    st.markdown(
        f'<p class="kp-cal-dialog-lead">{html.escape(_academic_period_caption(academic_year))}</p>',
        unsafe_allow_html=True,
    )
    _render_calendar_months(academic_year)
    st.markdown(
        _recommended_break_cards_html(
            academic_year,
            str(st.session_state.get("class_name") or ""),
        ),
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="kp-cal-nav-hint">'
        'Нажмите на название месяца — план на месяц.<br>'
        'Нажмите на № недели — занятия недели.'
        '</div>',
        unsafe_allow_html=True,
    )
    links = "".join(
        f'<a href="{url}" target="_blank" rel="noopener noreferrer">{html.escape(label)}</a>'
        for label, url in (*_OFFICIAL_HOLIDAY_SOURCES, *_TEACHER_LINKS)
    )
    st.markdown(
        '<div class="kp-cal-footer">'
        f'<p>{html.escape(_short_weeks_note(academic_year))}</p>'
        '<div class="kp-cal-legend">'
        '<span><i class="kp-cal-week"></i>№ — номер недели календарного плана</span>'
        '<span><i class="kp-cal-holiday"></i>Красный — выходной / официальный праздник</span>'
        '<span><i class="kp-cal-break"></i>Бежевый — рекомендуемые каникулы / перерыв</span>'
        '<span><i class="kp-cal-professional"></i>Профессиональная дата</span>'
        '</div>'
        '<p class="kp-cal-source-note">Официальные праздники отмечены по статье 112 ТК РФ; переносы выходных 2027 года не размечены без опубликованного постановления.</p>'
        '<p class="kp-useful-title">Полезно педагогу</p>'
        f"{links}</div>",
        unsafe_allow_html=True,
    )


def _render_year_calendar_card(academic_year: str, *, owner: str) -> None:
    st.markdown(
        '<div class="kp-cal-card">'
        '<div class="kp-cal-card-icon" aria-hidden="true">📅</div>'
        '<div class="kp-cal-card-copy">'
        f'<div class="kp-step-title">Календарь {html.escape(academic_year)} учебного года</div>'
        '<div class="kp-cal-card-lead">Недели №1–36 соответствуют строкам календарного плана</div>'
        f'<div class="kp-step-note">{html.escape(_academic_period_caption(academic_year))}</div>'
        "</div></div>",
        unsafe_allow_html=True,
    )
    if st.button(
        "Открыть календарь",
        type="secondary",
        use_container_width=True,
        key=f"kp_toggle_year_calendar_{owner}",
    ):
        st.session_state["kp_calendar_year"] = academic_year
        st.session_state["kp_calendar_open"] = True
        st.session_state["kp_calendar_owner"] = owner
        st.session_state.pop("kp_selected_week", None)
        st.session_state.pop("kp_selected_month", None)
        st.rerun()
    if (
        st.session_state.get("kp_calendar_open")
        and st.session_state.get("kp_calendar_owner") == owner
    ):
        st.session_state["kp_calendar_year"] = academic_year
        _show_year_calendar_dialog()


def _inject_landing_styles() -> None:
    st.markdown(
        """
        <style>
        html, body, .stApp, [data-testid="stAppViewContainer"] {
            background: #ffffff !important;
            line-height: 1.56;
        }
        header[data-testid="stHeader"] {
            display: none !important;
        }
        #MainMenu, footer, .stAppDeployButton, [data-testid="stToolbar"],
        [data-testid="stDecoration"], [data-testid="stStatusWidget"],
        [data-testid="stSkillsNudgeAnchor"], [data-testid="stSkillsNudge"],
        [data-testid="stAppDeployButton"], .stDeployButton,
        .stAppHeader, header {
            display: none !important;
            visibility: hidden !important;
            height: 0 !important;
        }
        .block-container,
        [data-testid="stMainBlockContainer"] {
            padding-top: 0.7rem !important;
            padding-bottom: 0.4rem !important;
            padding-left: 1.15rem !important;
            padding-right: 1.15rem !important;
            max-width: 1220px;
        }
        .stMainBlockContainer [data-testid="stVerticalBlock"] {
            gap: 0.42rem !important;
        }
        [data-testid="stHorizontalBlock"] {
            align-items: stretch !important;
        }
        [data-testid="stElementContainer"]:has(.kp-step-title),
        [data-testid="stElementContainer"]:has(.kp-card-head) {
            height: auto !important;
            min-height: 2.35rem !important;
            flex: 0 0 auto !important;
        }
        [data-testid="stMarkdown"]:has(.kp-step-title),
        [data-testid="stMarkdown"]:has(.kp-step-title) > div {
            height: auto !important;
            min-height: fit-content !important;
            flex-shrink: 0 !important;
            display: block !important;
        }
        [data-testid="stVerticalBlock"]:has(> [data-testid="stElementContainer"] .kp-left-card) {
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            background: #ffffff;
            box-shadow: 0 1px 2px rgba(15, 23, 42, 0.04);
            padding: 0.5rem 0.75rem 0.45rem;
            flex-grow: 0 !important;
            flex-basis: auto !important;
        }
        [data-testid="stElementContainer"]:has(.kp-card-shell) {
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-bottom: none;
            border-radius: 10px 10px 0 0;
            box-shadow: 0 1px 2px rgba(15, 23, 42, 0.04);
            padding: 0.5rem 0.75rem 0.3rem;
        }
        [data-testid="stElementContainer"]:has(.kp-card-shell) + [data-testid="stElementContainer"] {
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-top: none;
            border-radius: 0 0 10px 10px;
            box-shadow: 0 1px 2px rgba(15, 23, 42, 0.04);
            padding: 0 0.75rem 0.5rem;
            margin-bottom: 0.45rem;
        }
        [data-testid="stElementContainer"]:has(.kp-cal-card) {
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-bottom: none;
            border-radius: 10px 10px 0 0;
            box-shadow: 0 1px 2px rgba(15, 23, 42, 0.04);
            padding: 0.5rem 0.75rem 0.25rem;
        }
        [data-testid="stElementContainer"]:has(.kp-cal-card) + [data-testid="stElementContainer"] {
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-top: none;
            border-radius: 0 0 10px 10px;
            box-shadow: 0 1px 2px rgba(15, 23, 42, 0.04);
            padding: 0 0.75rem 0.5rem;
        }
        .kp-hero-top {
            margin: 0 !important;
            padding: 0 !important;
        }
        [data-testid="stHeading"],
        .stHeading {
            height: auto !important;
            min-height: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
        }
        .element-container:has([data-testid="stHeading"]),
        .element-container:has(h1) {
            min-height: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
        }
        .kp-hero-title, h1 {
            font-size: 25px !important;
            font-weight: 700 !important;
            line-height: 1.2 !important;
            margin: 0 0 0.2rem 0 !important;
            color: #1f2937 !important;
            padding: 0 !important;
            min-height: 0 !important;
        }
        .kp-hero-subtitle {
            font-size: 14.5px;
            line-height: 1.55;
            color: #4b5563;
            margin: 0 0 0.15rem 0;
        }
        .kp-card-head,
        .kp-step-title,
        .kp-step-note {
            display: block;
        }
        .kp-card-head {
            padding: 0 0 0.3rem 0;
        }
        .kp-step-title {
            font-size: 15.5px;
            font-weight: 600;
            color: #111827;
            margin: 0;
            line-height: 1.45;
        }
        .kp-panel-title {
            color: #0f3f75;
            font-size: 17px;
            font-weight: 700;
            line-height: 1.5;
            padding: 0.58rem 0.75rem;
            margin: 0 0 0.35rem;
            border-radius: 8px;
            background: #edf6ff;
        }
        .kp-field-title {
            color: #1f2937;
            font-size: 14.5px;
            font-weight: 600;
            line-height: 1.45;
            margin: 0.35rem 0 0.1rem;
        }
        .kp-field-note {
            color: #6b7280;
            font-size: 12.5px;
            line-height: 1.4;
            margin: 0 0 0.15rem;
        }
        .kp-step-note {
            font-size: 14.5px;
            color: #6b7280;
            margin: 0.12rem 0 0 0;
            line-height: 1.3;
        }
        .kp-year-spacer {
            font-size: 14.5px;
            color: #31333f;
            font-weight: 500;
            margin: 0 0 0.1rem 0;
            line-height: 1.3;
        }
        .kp-year-canon {
            font-size: 14.5px;
            color: #1d4ed8;
            font-weight: 600;
            margin: 0.05rem 0 0.35rem;
            line-height: 1.4;
        }
        .kp-card-divider {
            height: 1px;
            background: #eef2f7;
            margin: 0.35rem 0 0.3rem 0;
        }
        .kp-badge {
            display: inline-block;
            font-size: 0.68rem;
            font-weight: 700;
            letter-spacing: 0.04em;
            text-transform: uppercase;
            border-radius: 999px;
            padding: 0.12rem 0.45rem;
            margin-left: 0.35rem;
            vertical-align: middle;
        }
        .kp-badge-required {
            background: #eff6ff;
            color: #1d4ed8;
        }
        .kp-badge-optional {
            background: #f3f4f6;
            color: #6b7280;
        }
        [data-testid="stWidgetLabel"] p,
        [data-testid="stWidgetLabel"] label,
        .stTextInput label, .stNumberInput label {
            font-size: 14.5px !important;
            font-weight: 500 !important;
        }
        [data-testid="stWidgetLabel"] {
            margin-bottom: 0.1rem !important;
        }
        [data-testid="stTextInput"] input,
        [data-testid="stNumberInput"] input {
            font-size: 14.5px !important;
            min-height: 2.05rem !important;
            padding-top: 0.2rem !important;
            padding-bottom: 0.2rem !important;
        }
        [data-testid="stCaptionContainer"] {
            margin-top: 0 !important;
            font-size: 14px !important;
        }
        .kp-normative {
            margin: 0;
        }
        .kp-normative [data-testid="stExpander"] {
            margin: 0 !important;
            border: none !important;
            background: transparent !important;
        }
        .kp-normative details {
            border: none !important;
            background: transparent !important;
            padding: 0 !important;
        }
        .kp-normative summary {
            color: #6b7280;
            font-size: 13px;
            cursor: pointer;
            list-style: none;
        }
        .kp-normative summary::-webkit-details-marker {
            display: none;
        }
        div[data-testid="stFileUploader"] {
            margin: 0.15rem 0 0 0;
        }
        div[data-testid="stFileUploader"] label {
            display: none !important;
        }
        div[data-testid="stFileUploader"] section {
            padding-top: 0 !important;
            padding-bottom: 0 !important;
        }
        [data-testid="stFileUploaderDropzoneInstructions"],
        [data-testid="stFileUploaderDropzoneInstructions"] span,
        div[data-testid="stFileUploader"] small {
            display: none !important;
            visibility: hidden !important;
            height: 0 !important;
            max-height: 0 !important;
            overflow: hidden !important;
            margin: 0 !important;
            padding: 0 !important;
            font-size: 0 !important;
            line-height: 0 !important;
        }
        [data-testid="stFileUploaderDropzone"] {
            min-height: 2.15rem !important;
            padding: 0.2rem 0.45rem !important;
        }
        [data-testid="stFileUploaderDropzone"] button {
            min-height: 1.85rem !important;
            font-size: 14px !important;
            padding: 0.15rem 0.65rem !important;
        }
        [data-testid="stExpander"] {
            margin-top: 0;
            margin-bottom: 0;
        }
        [data-testid="stExpander"] details {
            border: none;
            background: transparent;
        }
        [data-testid="stExpander"] summary {
            color: #6b7280;
            font-size: 13px;
        }
        div[data-testid="stButton"] {
            margin-top: 0;
        }
        .stButton > button[kind="secondary"] {
            background: #ffffff;
            color: #2563eb;
            border: 1px solid #bfdbfe;
            font-size: 14.5px;
            font-weight: 600;
            min-height: 2.15rem;
            padding: 0.3rem 0.7rem;
            border-radius: 8px;
        }
        .stButton > button[kind="secondary"]:hover {
            background: #eff6ff;
            color: #1d4ed8;
            border-color: #93c5fd;
        }
        .stButton > button[kind="primary"] {
            background: #2563eb;
            border: none;
            font-size: 16px;
            font-weight: 600;
            padding: 0.7rem 1.1rem;
            min-height: 50px;
            border-radius: 8px;
        }
        .stButton > button[kind="primary"]:hover {
            background: #1d4ed8;
            border: none;
            color: #ffffff;
        }
        .kp-cal-month-title {
            font-size: 18px;
            font-weight: 750;
            line-height: 1.4;
            color: #172554;
            margin: 0 0 0.32rem;
        }
        .kp-cal-days-head {
            display: grid;
            grid-template-columns: 0.95fr repeat(7, 1fr);
            gap: 2px;
            align-items: center;
            min-height: 22px;
            line-height: 22px;
            margin-bottom: 9px;
            color: #475569;
            font-size: 12px;
            font-weight: 750;
            text-align: center;
        }
        .kp-cal-week-head {
            color: #1d4ed8;
            font-family: Arial, sans-serif;
            font-size: 12px;
            font-weight: 800;
            white-space: nowrap;
        }
        .kp-cal-days-row {
            display: grid;
            grid-template-columns: repeat(7, 1fr);
            gap: 2px;
            height: 27px;
            align-items: stretch;
        }
        .kp-cal-day {
            display: flex;
            align-items: center;
            justify-content: center;
            border: 1px solid #d8e0e8;
            border-radius: 4px;
            background: #ffffff;
            color: #1f2937;
            font-size: 13px;
            font-weight: 650;
            line-height: 1;
        }
        .kp-cal-day-empty {
            background: transparent;
            border-color: transparent;
        }
        .kp-cal-day.kp-cal-study {
            background: #ffffff;
            color: #172b3f;
            border-color: #cfd9e4;
        }
        .kp-cal-day.kp-cal-wknd {
            background: #f3f4f6;
            color: #4b5563;
        }
        .kp-cal-day.kp-cal-break {
            background: #f3dfbf;
            color: #6b4d22;
        }
        .kp-cal-week-empty {
            height: 27px;
        }
        [role="dialog"] [data-testid="stHorizontalBlock"]:has(.kp-cal-month-anchor) {
            gap: 1.5rem !important;
            margin-bottom: 1.5rem;
        }
        [role="dialog"] [data-testid="stVerticalBlock"]:has(.kp-cal-month-anchor) {
            gap: 3px !important;
            padding: 0.6rem !important;
            border: 1px solid #d9e3ef;
            border-radius: 10px;
            background: #ffffff;
            box-shadow: 0 2px 8px rgba(15, 63, 117, 0.06);
        }
        [role="dialog"] [data-testid="stVerticalBlock"]:has(.kp-cal-month-anchor)
        [data-testid="stHorizontalBlock"] {
            gap: 0.28rem !important;
            min-height: 27px;
        }
        [role="dialog"] [data-testid="stVerticalBlock"]:has(.kp-cal-month-anchor)
        [data-testid="stButton"] {
            margin: 0 !important;
        }
        [role="dialog"] [data-testid="stVerticalBlock"]:has(.kp-cal-month-anchor)
        .stButton > button {
            height: 27px !important;
            min-height: 27px !important;
            padding: 0 !important;
            border-radius: 4px !important;
            border: 1px solid #93c5fd !important;
            background: #dbeafe !important;
            color: #123a70 !important;
            font-size: 12.5px !important;
            font-weight: 800 !important;
        }
        [role="dialog"] [data-testid="stVerticalBlock"]:has(.kp-cal-month-anchor)
        .stButton > button:hover {
            background: #2563eb !important;
            color: #ffffff !important;
            border-color: #2563eb !important;
        }        .kp-cal-legend {
            display: flex;
            flex-wrap: wrap;
            gap: 0.45rem 1rem;
            margin-top: 0.7rem;
            font-size: 13px;
            color: #4b5563;
        }
        .kp-cal-legend i {
            display: inline-block;
            width: 0.7rem;
            height: 0.7rem;
            border-radius: 2px;
            margin-right: 0.35rem;
            vertical-align: -1px;
        }
        .kp-cal-legend i.kp-cal-study { background: #ffffff; border: 1px solid #cbd5e1; }
        .kp-cal-legend i.kp-cal-wknd { background: #f3f4f6; }
        .kp-cal-legend i.kp-cal-break { background: #f3eee4; }
        .kp-cal-legend i.kp-cal-week {
            background: #2563eb;
            border-radius: 999px;
        }

        .kp-cal-dialog-title {
            font-size: 15.5px;
            font-weight: 600;
            color: #111827;
            margin: 0 0 0.15rem 0;
        }
        [role="dialog"] {
            width: min(1310px, 96vw) !important;
            max-width: min(1310px, 96vw) !important;
        }
        [role="dialog"] h2 {
            line-height: 1.4 !important;
        }
        [role="dialog"] [data-testid="stVerticalBlock"] {
            gap: 0.25rem !important;
        }
        [role="dialog"] .stButton > button {
            min-height: 1.7rem !important;
            height: 1.7rem !important;
            padding: 0.05rem 0.12rem !important;
            border-radius: 999px !important;
            font-size: 11px !important;
            color: #1d4ed8 !important;
            background: #eff6ff !important;
            border: 1px solid #bfdbfe !important;
        }
        .kp-cal-dialog-lead {
            color: #172554;
            font-size: 17px;
            font-weight: 650;
            line-height: 1.4;
            margin: 0.5rem 0 1.25rem;
        }
        .kp-cal-footer {
            margin-top: 0.5rem;
            padding: 0.55rem 0.7rem;
            border-radius: 8px;
            background: #fff8ed;
            color: #5f4a31;
            font-size: 13px;
            line-height: 1.45;
        }
        .kp-cal-footer p { margin: 0 0 0.3rem; }
        .kp-cal-vacations {
            margin: 0.35rem 0 0.8rem;
            padding: 0.8rem;
            border-radius: 10px;
            background: #fffaf2;
        }
        .kp-cal-vacations h3 {
            margin: 0 0 0.65rem;
            color: #4f3d28;
            font-size: 17px;
            font-weight: 750;
            line-height: 1.4;
        }
        .kp-cal-vacation-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: 0.65rem;
        }
        .kp-cal-vacation-card {
            min-width: 0;
            padding: 0.65rem 0.7rem;
            border: 1px solid #eadcc4;
            border-radius: 8px;
            background: #ffffff;
            box-shadow: 0 1px 4px rgba(95, 74, 49, 0.05);
        }
        .kp-cal-vacation-card strong,
        .kp-cal-vacation-card span {
            display: block;
            line-height: 1.4;
        }
        .kp-cal-vacation-card strong {
            color: #4f3d28;
            font-size: 13.5px;
            font-weight: 750;
            margin-bottom: 0.18rem;
        }
        .kp-cal-vacation-card span {
            color: #7a6852;
            font-size: 12.5px;
        }
        .kp-cal-vacation-extra {
            margin: 0.65rem 0 0 !important;
            color: #6b5234;
            font-size: 12.5px;
            font-weight: 600;
            line-height: 1.45;
        }
        @media (max-width: 900px) {
            .kp-cal-vacation-grid {
                grid-template-columns: repeat(2, minmax(0, 1fr));
            }
        }
        .kp-cal-nav-hint {
            margin: 0.3rem 0 0.75rem;
            padding: 0.55rem 0.75rem;
            border: 1px solid #dbeafe;
            border-radius: 8px;
            background: #f0f7ff;
            color: #29496f;
            font-size: 13.5px;
            font-weight: 600;
            line-height: 1.5;
        }
        .kp-useful-title {
            color: #374151;
            font-weight: 700;
            margin-top: 0.55rem !important;
        }
        .kp-cal-footer a {
            display: inline-block;
            margin-right: 1.15rem;
            color: #1d4ed8;
        }
        .kp-week-title {
            color: #0f3f75;
            font-size: 21px;
            font-weight: 700;
            line-height: 1.35;
            margin: 0 0 0.7rem;
        }
        .kp-week-source {
            color: #64748b;
            font-size: 13px;
            line-height: 1.5;
            margin: 0 0 0.55rem;
        }
        .kp-month-weeks {
            color: #475569;
            font-size: 14px;
            font-weight: 650;
            line-height: 1.4;
            margin: -0.35rem 0 0.75rem;
        }
        .kp-week-part {
            color: #0f3f75;
            font-size: 14px;
            font-weight: 700;
            line-height: 1.5;
            margin: 0.8rem 0 0.15rem;
        }
        .kp-week-fields {
            border-top: 1px solid #dbe5ef;
        }
        .kp-week-field {
            padding: 0.55rem 0;
            border-bottom: 1px solid #e5e7eb;
        }
        .kp-week-label {
            color: #526579;
            font-size: 12px;
            font-weight: 700;
            letter-spacing: 0.03em;
            text-transform: uppercase;
            margin-bottom: 0.15rem;
        }
        .kp-week-value {
            color: #1f2937;
            font-size: 15px;
            line-height: 1.5;
            overflow-wrap: anywhere;
        }
        [data-testid="stDownloadButton"] button {
            min-height: 52px !important;
            font-size: 1.05rem !important;
            font-weight: 600 !important;
            padding: 0.75rem 1.15rem !important;
            border: none !important;
            color: #ffffff !important;
            background: #14866d !important;
            border-radius: 8px !important;
        }
        [data-testid="stDownloadButton"] button:hover {
            background: #0f6f5b !important;
            color: #ffffff !important;
        }
        .kp-status-card {
            margin: 1.75rem auto 0 auto;
            max-width: 860px;
            padding: 0.25rem 0 0.5rem 0;
        }
        .kp-status-title {
            font-size: 23px;
            font-weight: 700;
            color: #111827;
            margin: 0 0 0.4rem 0;
            line-height: 1.3;
        }
        .kp-status-lead {
            font-size: 16px;
            color: #4b5563;
            margin: 0 0 1.35rem 0;
            line-height: 1.45;
        }
        .kp-program-name {
            font-size: 19px;
            font-weight: 600;
            color: #111827;
            margin: 0 0 0.4rem 0;
            line-height: 1.35;
        }
        .kp-status-meta,
        .kp-status-kpi,
        .kp-status-check {
            font-size: 16px;
            color: #374151;
            margin: 0 0 0.45rem 0;
            line-height: 1.45;
        }
        .kp-status-kpi {
            margin: 1rem 0 1.15rem 0;
            color: #1f2937;
        }
        .kp-status-checks {
            margin: 0 0 1.1rem 0;
        }
        .kp-status-check.warn {
            color: #92400e;
        }
        .kp-edit-slot [data-testid="stButton"] {
            margin-top: 0;
        }
        .kp-edit-slot button {
            background: transparent !important;
            color: #2563eb !important;
            border: none !important;
            box-shadow: none !important;
            min-height: 2.25rem !important;
            font-size: 15px !important;
            font-weight: 600 !important;
            padding: 0.25rem 0 !important;
        }
        .kp-edit-slot button:hover {
            color: #1d4ed8 !important;
            background: transparent !important;
        }
        .block-container:has(.kp-status-card)
        [data-testid="stVerticalBlock"]:has(> [data-testid="stElementContainer"] .kp-form-park):not(:has(.kp-status-card)),
        .block-container:has(.kp-status-card)
        [data-testid="stVerticalBlock"]:has(> .element-container .kp-form-park):not(:has(.kp-status-card)) {
            display: none !important;
        }
        .block-container:has(.kp-status-card) .kp-hero-top,
        .block-container:has(.kp-status-card) .kp-step-card,
        .block-container:has(.kp-status-card) .kp-card-head,
        .block-container:has(.kp-status-card) .kp-left-card,
        .block-container:has(.kp-status-card) .kp-cal-card,
        .block-container:has(.kp-status-card) .kp-cal-grid,
        .block-container:has(.kp-status-card) .kp-uploaded-name,
        .block-container:has(.kp-status-card) .kp-normative,
        .block-container:has(.kp-status-card) [data-testid="stFileUploader"],
        .block-container:has(.kp-status-card) [data-testid="stTextInput"],
        .block-container:has(.kp-status-card) [data-testid="stNumberInput"],
        .block-container:has(.kp-status-card) [data-testid="stCaptionContainer"],
        .block-container:has(.kp-status-card) [data-testid="stVerticalBlockBorderWrapper"],
        .block-container:has(.kp-status-card) [data-testid="stHorizontalBlock"]:has([data-testid="stFileUploader"]),
        .block-container:has(.kp-status-card) .element-container:has(.kp-form-actions),
        .block-container:has(.kp-status-card) .element-container:has(.kp-form-actions) + .element-container,
        .block-container:has(.kp-status-card) .element-container:has(.kp-check-slot),
        .block-container:has(.kp-status-card) .element-container:has(.kp-check-slot) + .element-container,
        .block-container:has(.kp-status-card) [data-testid="stElementContainer"]:has(.kp-form-actions),
        .block-container:has(.kp-status-card) [data-testid="stElementContainer"]:has(.kp-form-actions) + [data-testid="stElementContainer"],
        .block-container:has(.kp-status-card) [data-testid="stElementContainer"]:has(.kp-check-slot),
        .block-container:has(.kp-status-card) [data-testid="stElementContainer"]:has(.kp-check-slot) + [data-testid="stElementContainer"] {
            display: none !important;
        }
        .kp-normative-check {
            background: #ffffff;
            border: none;
            border-radius: 0;
            padding: 0.35rem 0 0.15rem 0;
            margin: 0.15rem 0 0.15rem 0;
        }
        .kp-normative-check-title {
            font-size: 1rem;
            font-weight: 600;
            color: #374151;
            margin: 0 0 0.2rem 0;
        }
        .kp-normative-check-lead {
            font-size: 0.95rem;
            color: #6b7280;
            margin: 0 0 0.45rem 0;
            line-height: 1.4;
        }
        .kp-normative-layer {
            font-size: 0.94rem;
            font-weight: 600;
            color: #111827;
            margin: 0.55rem 0 0.1rem 0;
        }
        .kp-normative-layer-lead {
            font-size: 0.82rem;
            color: #6b7280;
            margin: 0 0 0.2rem 0;
            line-height: 1.35;
        }
        .kp-normative-note {
            font-size: 0.9rem;
            color: #4b5563;
            margin: 0.2rem 0 0.35rem 0;
            line-height: 1.4;
        }
        .kp-normative-check-label {
            font-size: 0.92rem;
            font-weight: 600;
            margin: 0.4rem 0 0.15rem 0;
        }
        .kp-normative-check-label.ok { color: #166534; }
        .kp-normative-check-label.warn { color: #92400e; }
        .kp-normative-check-label.skip { color: #6b7280; }
        .kp-normative-check ul {
            margin: 0 0 0.15rem 1.1rem;
            padding: 0;
            color: #374151;
        }
        .kp-normative-check li {
            margin: 0.12rem 0;
            line-height: 1.4;
        }
        [data-testid="stFileUploaderFile"] {
            display: none !important;
        }
        .kp-uploaded-name {
            margin: 0.2rem 0 0.1rem 0;
            color: #374151;
            font-size: 0.9rem;
            line-height: 1.3;
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
        }
        button[title="Удалить файл"] {
            min-width: 2rem !important;
            height: 2rem !important;
            padding: 0 !important;
            font-size: 1.15rem !important;
            line-height: 1 !important;
            color: #6b7280 !important;
            background: #ffffff !important;
            border: 1px solid #e5e7eb !important;
            border-radius: 8px !important;
        }
        button[title="Удалить файл"]:hover {
            color: #b91c1c !important;
            border-color: #fecaca !important;
            background: #fef2f2 !important;
        }
        .kp-cal-card {
            display: flex;
            align-items: flex-start;
            gap: 0.8rem;
        }
        .kp-cal-card-icon {
            display: flex;
            align-items: center;
            justify-content: center;
            width: 2.35rem;
            height: 2.35rem;
            flex: 0 0 2.35rem;
            border-radius: 9px;
            background: #dbeafe;
            color: #174a87;
            font-size: 1.2rem;
        }
        .kp-cal-card-copy {
            min-width: 0;
        }
        .kp-cal-card-lead {
            color: #244766;
            font-size: 14px;
            font-weight: 600;
            line-height: 1.5;
            margin-top: 0.18rem;
        }
        [data-testid="stElementContainer"]:has(.kp-cal-card),
        [data-testid="stElementContainer"]:has(.kp-cal-card) + [data-testid="stElementContainer"] {
            background: #edf6ff;
            border-color: #bcd7f2;
        }
        [data-testid="stElementContainer"]:has(.kp-cal-card) + [data-testid="stElementContainer"]
        .stButton > button {
            background: #ffffff;
            border-color: #74a9df;
            color: #174a87;
            font-weight: 700;
        }
        [data-testid="stElementContainer"]:has(.kp-cal-card) + [data-testid="stElementContainer"]
        .stButton > button:hover {
            background: #dbeafe;
            border-color: #2563eb;
            color: #123a70;
        }
        [role="dialog"] [data-testid="stVerticalBlock"]:has(.kp-cal-month-anchor)
        .stButton > button p {
            white-space: nowrap !important;
            word-break: keep-all !important;
            overflow-wrap: normal !important;
            line-height: 1 !important;
        }
        [role="dialog"] .stButton > button,
        [role="dialog"] .stButton > button * {
            white-space: nowrap !important;
            word-break: normal !important;
            overflow-wrap: normal !important;
        }
        [class*="st-key-kp_week_"] button {
            height: 27px !important;
            min-height: 27px !important;
            padding: 0 !important;
            border-radius: 4px !important;
            border: 1px solid #93c5fd !important;
            background: #dbeafe !important;
            color: #123a70 !important;
            font-size: 12.5px !important;
            font-weight: 800 !important;
        }
        [class*="st-key-kp_week_"] button,
        [class*="st-key-kp_week_"] button * {
            white-space: nowrap !important;
            word-break: normal !important;
            overflow-wrap: normal !important;
            line-height: 1 !important;
        }
        [class*="st-key-kp_calendar_month_title_"] button {
            height: auto !important;
            min-height: 38px !important;
            padding: 0.45rem 0.65rem !important;
            border: 1px solid #bfdbfe !important;
            background: #eaf4ff !important;
            color: #172554 !important;
            font-size: 18px !important;
            font-weight: 750 !important;
            line-height: 1.4 !important;
            box-shadow: none !important;
            cursor: pointer !important;
            border-radius: 7px !important;
            display: flex !important;
            align-items: center !important;
            text-align: left !important;
        }
        [class*="st-key-kp_calendar_month_title_"] button::after {
            content: "›";
            margin-left: auto;
            color: #1d4ed8;
            font-size: 22px;
            font-weight: 700;
            line-height: 1;
        }
        [class*="st-key-kp_calendar_month_title_"] button:hover {
            background: #dbeeff !important;
            color: #1e40af !important;
            border-color: #93c5fd !important;
        }
        .kp-cal-days-head .kp-cal-weekend-head {
            color: #b42318;
        }
        .kp-cal-day.kp-cal-wknd {
            background: #ffffff;
            color: #b42318;
            border-color: #e2e8f0;
        }
        .kp-cal-day.kp-cal-break.kp-cal-wknd {
            background: #f3dfbf;
            color: #b42318;
            border-color: #ead4ae;
        }
        .kp-cal-day.kp-cal-professional {
            position: relative;
            background: #f1e8ff;
            color: #6842a8;
            border-color: #d8c2f3;
            font-weight: 800;
        }
        .kp-cal-day.kp-cal-holiday {
            position: relative;
            background: #fff1f0;
            color: #b42318;
            border-color: #fecaca;
            font-weight: 800;
        }
        .kp-cal-day.kp-cal-holiday:hover::after {
            content: attr(data-holiday);
            position: absolute;
            left: 50%;
            bottom: calc(100% + 7px);
            transform: translateX(-50%);
            z-index: 20;
            width: max-content;
            max-width: 190px;
            padding: 0.35rem 0.5rem;
            border-radius: 6px;
            background: #172554;
            color: #ffffff;
            font-size: 11px;
            font-weight: 600;
            line-height: 1.3;
            box-shadow: 0 4px 12px rgba(15, 23, 42, 0.2);
        }
        .kp-cal-day.kp-cal-professional:hover::after {
            content: attr(data-professional);
            position: absolute;
            left: 50%;
            bottom: calc(100% + 7px);
            transform: translateX(-50%);
            z-index: 20;
            width: max-content;
            max-width: 190px;
            padding: 0.35rem 0.5rem;
            border-radius: 6px;
            background: #3f2a66;
            color: #ffffff;
            font-size: 11px;
            font-weight: 600;
            line-height: 1.3;
            box-shadow: 0 4px 12px rgba(15, 23, 42, 0.2);
        }
        .kp-cal-legend i.kp-cal-holiday {
            background: #fff1f0;
            border: 1px solid #fca5a5;
        }
        .kp-cal-legend i.kp-cal-professional {
            background: #f1e8ff;
            border: 1px solid #c4a5eb;
        }
        .kp-cal-source-note {
            color: #6b7280;
            font-size: 12px;
            line-height: 1.45;
            margin-top: 0.55rem !important;
        }
        .kp-week-table-scroll {
            width: 100%;
            overflow-x: auto;
            border: 1px solid #d9e2ec;
            border-radius: 8px;
            background: #ffffff;
            margin: 0 0 0.75rem;
        }
        .kp-week-table {
            width: 1780px;
            min-width: 1780px;
            border-collapse: separate;
            border-spacing: 0;
            color: #1f2937;
            font-size: 13.5px;
            line-height: 1.45;
        }
        .kp-week-table th,
        .kp-week-table td {
            padding: 0.65rem 0.7rem;
            border-right: 1px solid #e2e8f0;
            border-bottom: 1px solid #e2e8f0;
            text-align: left;
            vertical-align: top;
            white-space: normal;
            overflow-wrap: anywhere;
        }
        .kp-week-table th {
            position: sticky;
            top: 0;
            z-index: 2;
            background: #eaf3fc;
            color: #173f6b;
            font-size: 12.5px;
            font-weight: 750;
        }
        .kp-week-table th:nth-child(1),
        .kp-week-table td:nth-child(1) {
            position: sticky;
            left: 0;
            width: 120px;
            min-width: 120px;
            z-index: 3;
            background: #ffffff;
        }
        .kp-week-table th:nth-child(2),
        .kp-week-table td:nth-child(2) {
            position: sticky;
            left: 120px;
            width: 260px;
            min-width: 260px;
            z-index: 3;
            background: #ffffff;
            box-shadow: 5px 0 8px -8px #64748b;
        }
        .kp-week-table th:nth-child(1),
        .kp-week-table th:nth-child(2) {
            z-index: 4;
            background: #eaf3fc;
        }
        .kp-week-table th:nth-child(3),
        .kp-week-table th:nth-child(4),
        .kp-week-table td:nth-child(3),
        .kp-week-table td:nth-child(4) {
            width: 250px;
            min-width: 250px;
        }
        .kp-week-table th:nth-child(5),
        .kp-week-table td:nth-child(5) {
            width: 170px;
            min-width: 170px;
        }
        .kp-week-table th:nth-child(6),
        .kp-week-table th:nth-child(7),
        .kp-week-table td:nth-child(6),
        .kp-week-table td:nth-child(7) {
            width: 245px;
            min-width: 245px;
        }
        .kp-week-table th:nth-child(8),
        .kp-week-table td:nth-child(8) {
            width: 210px;
            min-width: 210px;
        }
        .kp-month-table {
            width: 2150px;
            min-width: 2150px;
        }
        .kp-month-table th:nth-child(1),
        .kp-month-table td:nth-child(1) {
            width: 80px;
            min-width: 80px;
        }
        .kp-month-table th:nth-child(2),
        .kp-month-table td:nth-child(2) {
            left: 80px;
            width: 125px;
            min-width: 125px;
        }
        .kp-month-table th:nth-child(3),
        .kp-month-table td:nth-child(3) {
            position: sticky;
            left: 205px;
            width: 275px;
            min-width: 275px;
            z-index: 3;
            background: #ffffff;
            box-shadow: 5px 0 8px -8px #64748b;
        }
        .kp-month-table th:nth-child(1),
        .kp-month-table th:nth-child(2),
        .kp-month-table th:nth-child(3) {
            z-index: 4;
            background: #eaf3fc;
        }
        .kp-month-table th:nth-child(4),
        .kp-month-table td:nth-child(4),
        .kp-month-table th:nth-child(5),
        .kp-month-table td:nth-child(5) {
            width: 250px;
            min-width: 250px;
        }
        .kp-month-table th:nth-child(6),
        .kp-month-table td:nth-child(6) {
            width: 180px;
            min-width: 180px;
        }
        .kp-month-table th:nth-child(7),
        .kp-month-table td:nth-child(7),
        .kp-month-table th:nth-child(8),
        .kp-month-table td:nth-child(8) {
            width: 270px;
            min-width: 270px;
        }
        .kp-month-table th:nth-child(9),
        .kp-month-table td:nth-child(9) {
            width: 210px;
            min-width: 210px;
        }
        [class*="st-key-kp_week_back"] button {
            width: auto !important;
            min-width: 150px !important;
            padding: 0.35rem 0.8rem !important;
            border-radius: 7px !important;
            font-size: 13px !important;
        }        /* STATE 1: readable vertical rhythm; scrolling is preferable to compression. */
        .block-container,
        [data-testid="stMainBlockContainer"] {
            padding-top: 1.25rem !important;
            padding-bottom: 2rem !important;
            max-width: 1240px;
        }
        .stMainBlockContainer > [data-testid="stVerticalBlock"] {
            gap: 0.72rem !important;
        }
        .kp-hero-title, h1 {
            font-size: 29px !important;
            line-height: 1.3 !important;
            margin: 0 0 0.35rem !important;
            color: #102a43 !important;
        }
        .kp-hero-subtitle {
            font-size: 15.5px;
            line-height: 1.65;
            margin: 0 0 0.85rem;
            color: #526579;
        }
        [data-testid="stColumn"]:has(.kp-panel-title) {
            padding: 0.9rem 1rem 1.05rem;
            border-radius: 12px;
            background: #ffffff;
            border: 1px solid #e8eef5;
            box-shadow: 0 3px 14px rgba(30, 64, 105, 0.07);
        }
        .kp-panel-title {
            color: #103f75;
            font-size: 20px;
            line-height: 1.33;
            padding: 0.85rem 0.82rem;
            margin: 0 0 0.72rem;
            background: #dceeff;
        }
        .kp-field-title {
            color: #172b3f;
            font-size: 15px;
            line-height: 1.55;
            margin: 0.68rem 0 0.22rem;
        }
        .kp-field-note {
            color: #68788a;
            font-size: 13.5px;
            line-height: 1.58;
            margin: 0 0 0.32rem;
        }
        div[data-testid="stFileUploader"] {
            margin: 0.45rem 0 0.65rem;
        }
        [data-testid="stFileUploaderDropzone"] {
            min-height: 2.75rem !important;
            padding: 0.34rem 0.55rem !important;
            background: #f4f7fa;
        }
        [data-testid="stWidgetLabel"] {
            margin-bottom: 0.3rem !important;
        }
        [data-testid="stTextInput"],
        [data-testid="stNumberInput"] {
            margin-bottom: 0.45rem;
        }
        [data-testid="stTextInput"] input,
        [data-testid="stNumberInput"] input {
            min-height: 2.6rem !important;
            padding-top: 0.45rem !important;
            padding-bottom: 0.45rem !important;
            color: #172b3f !important;
            background: #f4f7fa !important;
        }
        .kp-year-canon {
            line-height: 1.55;
            margin: 0.18rem 0 0.7rem;
        }
        [data-testid="stElementContainer"]:has(.kp-cal-card) {
            padding: 0.75rem 1rem 0.35rem;
            border-color: #e8eef5;
            box-shadow: 0 3px 14px rgba(30, 64, 105, 0.06);
        }
        [data-testid="stElementContainer"]:has(.kp-cal-card) + [data-testid="stElementContainer"] {
            padding: 0.15rem 1rem 0.85rem;
            border-color: #e8eef5;
            box-shadow: 0 3px 14px rgba(30, 64, 105, 0.06);
        }
        .kp-step-title {
            color: #172b3f;
            line-height: 1.55;
        }
        .kp-step-note {
            color: #68788a;
            line-height: 1.55;
            margin: 0.2rem 0 0;
        }
        .kp-check-slot + [data-testid="stElementContainer"] {
            margin-top: 0.35rem;
        }

        /* Readability pass: explicit component hierarchy without scaling the calendar grid. */
        .kp-badge {
            font-size: 11.5px;
            line-height: 1.35;
        }
        [data-testid="stWidgetLabel"] p,
        [data-testid="stWidgetLabel"] label,
        .stTextInput label,
        .stNumberInput label {
            font-size: 15px !important;
            font-weight: 600 !important;
            line-height: 1.45 !important;
        }
        [data-testid="stTextInput"] input,
        [data-testid="stNumberInput"] input {
            min-height: 44px !important;
            font-size: 15px !important;
            line-height: 1.4 !important;
        }
        [data-testid="stFileUploaderDropzone"] {
            min-height: 44px !important;
        }
        [data-testid="stFileUploaderDropzone"] button {
            min-height: 36px !important;
            font-size: 14.5px !important;
            font-weight: 600 !important;
        }
        .kp-cal-card .kp-step-title {
            font-size: 18px;
            font-weight: 700;
        }
        .kp-cal-card-lead,
        .kp-cal-card .kp-step-note {
            font-size: 14px;
        }
        [data-testid="stMarkdown"] p.kp-status-title {
            font-size: 24px !important;
            font-weight: 700 !important;
            line-height: 1.35 !important;
            margin-bottom: 0.5rem !important;
        }
        [data-testid="stMarkdown"] p.kp-status-lead {
            font-size: 16.5px !important;
            line-height: 1.5 !important;
        }
        [data-testid="stMarkdown"] p.kp-program-name {
            font-size: 19px !important;
            font-weight: 600 !important;
            line-height: 1.4 !important;
        }
        [data-testid="stMarkdown"] p.kp-status-meta,
        [data-testid="stMarkdown"] p.kp-status-kpi,
        [data-testid="stMarkdown"] p.kp-status-check {
            font-size: 16.5px !important;
            line-height: 1.5 !important;
        }
        [class*="st-key-kp_toggle_year_calendar_"] button,
        [class*="st-key-regenerate_calendar"] button {
            min-height: 44px !important;
            font-size: 15px !important;
            font-weight: 600 !important;
        }
        .kp-edit-slot button {
            min-height: 42px !important;
            font-size: 15px !important;
        }
        [data-testid="stElementContainer"]:has(.kp-check-slot)
        + [data-testid="stElementContainer"] {
            margin-top: 0.45rem;
            margin-bottom: 0.7rem;
        }
        [class*="st-key-kp_toggle_year_calendar_inputs"] {
            margin-bottom: 0.45rem;
        }
        .kp-cal-vacation-card strong {
            font-size: 13.5px;
        }
        .kp-cal-vacation-card span,
        .kp-cal-vacation-extra,
        .kp-cal-source-note {
            font-size: 13px;
            line-height: 1.45;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _render_group_class_fields() -> tuple[str, str, str]:
    group_col, class_col = st.columns(2)
    with group_col:
        group_number = st.text_input(
            "Группа №",
            key="group_number",
            help="Необязательно. Если не заполнить, в документе останется линия.",
        )
    with class_col:
        class_name = st.text_input(
            "Класс",
            key="class_name",
            help="Необязательно. Если не заполнить, в документе останется линия.",
        )
    teacher_name = st.text_input(
        "ФИО педагога",
        key="teacher_name",
        help="Необязательно. Если не заполнить, в документе ничего не добавится.",
    )
    return group_number or "", class_name or "", teacher_name or ""


def _peek_academic_year_resolution(
    utp_file, program_file,
) -> AcademicYearResolution:
    utp_mentions: tuple = ()
    program_mentions: tuple = ()
    if utp_file is not None:
        try:
            utp_mentions = mentions_from_utp(parse_utp(utp_file.getvalue()))
        except Exception:
            utp_mentions = ()
    if program_file is not None:
        filename = getattr(program_file, "name", "") or ""
        try:
            if filename.lower().endswith(".docx"):
                program = parse_program(program_file.getvalue(), filename)
                program_mentions = mentions_from_program(program)
                if utp_file is None:
                    try:
                        utp_mentions = mentions_from_utp(parse_utp(program_file.getvalue()))
                    except Exception:
                        pass
        except Exception:
            program_mentions = ()
    return resolve_academic_year(utp_mentions, program_mentions)


def _render_academic_year_input(utp_file, program_file) -> str:
    resolution = _peek_academic_year_resolution(utp_file, program_file)
    docs_fp = _inputs_fingerprint(utp_file, program_file)
    suggested = academic_year_start(resolution.suggested)
    if st.session_state.get("academic_year_docs_fp") != docs_fp:
        st.session_state["academic_year_start"] = (
            suggested if suggested is not None else default_academic_year_start()
        )
        st.session_state["academic_year_docs_fp"] = docs_fp
    st.session_state.setdefault("academic_year_start", default_academic_year_start())

    st.markdown('<div class="kp-field-title">Учебный год</div>', unsafe_allow_html=True)
    start = int(
        st.number_input(
            "Начало учебного года",
            min_value=1990,
            max_value=2100,
            step=1,
            key="academic_year_start",
            help="Укажите первый календарный год учебного периода.",
            label_visibility="collapsed",
        )
    )
    academic_year = format_academic_year(start)
    st.markdown(
        f'<p class="kp-year-canon">{html.escape(academic_year)} учебный год</p>',
        unsafe_allow_html=True,
    )
    if resolution.status is AcademicYearStatus.CONFLICT:
        st.warning(resolution.message)
    elif resolution.status in {AcademicYearStatus.AUTO, AcademicYearStatus.SINGLE}:
        st.info(resolution.message)
    elif resolution.status is AcademicYearStatus.MISSING and (utp_file or program_file):
        st.caption(resolution.message)
    return academic_year


def _form_is_open() -> bool:
    if st.session_state.get("ui_edit_inputs"):
        return True
    return not bool(st.session_state.get("analysis_ready"))


def _open_input_form() -> None:
    st.session_state["ui_edit_inputs"] = True
    st.rerun()


def _render_upload_fields() -> tuple[object | None, object | None, object | None, str, str, str, str]:
    left_col, right_col = st.columns((1.06, 0.94), gap="medium")
    with left_col:
        st.markdown(
            '<div class="kp-panel-title">1. Документы</div>'
            '<div class="kp-field-title">Программа обучения'
            '<span class="kp-badge kp-badge-required">обязательно</span></div>',
            unsafe_allow_html=True,
        )
        program_file = _file_uploader_with_clear(
            "program",
            label="Загрузите образовательную программу",
            type=("doc", "docx"),
            help="Программа — образовательная программа, DOC/DOCX, до 10 МБ.",
        )
        st.markdown(
            '<div class="kp-field-title">Учебно-тематический план'
            '<span class="kp-badge kp-badge-optional">необязательно</span></div>'
            '<div class="kp-field-note">Загрузите отдельно, только если УТП находится в другом файле.</div>',
            unsafe_allow_html=True,
        )
        utp_file = _file_uploader_with_clear(
            "utp",
            label="Загрузите УТП",
            type=("docx",),
            help="УТП — учебно-тематический план, DOCX, до 10 МБ.",
        )
        st.markdown(
            '<div class="kp-field-title">Шаблон календарного плана'
            '<span class="kp-badge kp-badge-optional">необязательно</span></div>'
            '<div class="kp-field-note">Если есть образец вашей организации — загрузите его; иначе используем стандартный.</div>',
            unsafe_allow_html=True,
        )
        organization_template_file = _file_uploader_with_clear(
            "template",
            label="Шаблон календарного плана вашей организации",
            type=("docx",),
            help=(
                "Шаблон — только образец календарного плана организации, "
                "DOCX, до 10 МБ."
            ),
        )
    with right_col:
        st.markdown(
            '<div class="kp-panel-title">2. Сведения для плана</div>',
            unsafe_allow_html=True,
        )
        academic_year = _render_academic_year_input(utp_file, program_file)
        group_number, class_name, teacher_name = _render_group_class_fields()

    return (
        utp_file,
        program_file,
        organization_template_file,
        academic_year,
        group_number,
        class_name,
        teacher_name,
    )


def _render_upload_screen() -> tuple[object | None, object | None, object | None, str, str, str, str, bool]:
    _inject_landing_styles()
    form_open = _form_is_open()

    with st.container():
        st.markdown('<div class="kp-form-park"></div>', unsafe_allow_html=True)
        st.markdown('<div class="kp-hero-top">', unsafe_allow_html=True)
        st.title("Календарь педагога")
        if form_open:
            st.markdown(
                '<p class="kp-hero-subtitle">Загрузите документы — приложение проверит часы, '
                "составит расписание занятий и подготовит календарный план в Word.</p>",
                unsafe_allow_html=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)
        fields = _render_upload_fields()
        st.markdown('<div class="kp-form-actions"></div>', unsafe_allow_html=True)
        st.markdown('<div class="kp-check-slot"></div>', unsafe_allow_html=True)
        check_clicked = st.button(
            "Проверить документы",
            type="primary",
            use_container_width=True,
        )
        if not st.session_state.get("analysis_ready") or form_open:
            _render_year_calendar_card(str(fields[3]), owner="inputs")
        _render_normative_panel()
    return (*fields, check_clicked)


def _render_normative_panel() -> None:
    normative_registry = get_builtin_normative_registry()
    registry_snapshot = normative_registry.current
    st.markdown('<div class="kp-normative">', unsafe_allow_html=True)
    with st.expander("Нормативная база", expanded=False):
        st.caption(
            "Справочник документов из реестра. Это не проверка вашей программы "
            "и на календарь автоматически не влияет."
        )
        st.write(
            f"Версия: **{registry_snapshot.registry_version}** · действует с: "
            f"**{registry_snapshot.effective_from:%d.%m.%Y}** · проверена: "
            f"**{registry_snapshot.verified_on:%d.%m.%Y}**."
        )
        st.write(f"Документов в реестре: **{len(registry_snapshot.documents)}**.")
        st.dataframe(
            [
                {
                    "Документ": document.title,
                    "Номер": document.number,
                    "Дата": document.document_date.strftime("%d.%m.%Y"),
                    "Статус": document.status.value,
                    "Что регулирует": document.regulates,
                    "Действует с": (
                        document.effective_from.strftime("%d.%m.%Y")
                        if document.effective_from
                        else ""
                    ),
                    "Действует до": (
                        document.effective_until.strftime("%d.%m.%Y")
                        if document.effective_until
                        else ""
                    ),
                    "Официальный URL": document.official_url,
                    "Проверено": (
                        document.verified_on.strftime("%d.%m.%Y")
                        if document.verified_on
                        else ""
                    ),
                }
                for document in registry_snapshot.documents
            ],
            hide_index=True,
            use_container_width=True,
        )
        saved_version = st.session_state.setdefault(
            "calendar_registry_version", registry_snapshot.registry_version
        )
        notice = get_update_notice(
            CalendarRegistryReference(saved_version), normative_registry
        )
        if notice.update_available:
            st.warning(
                f"Доступна нормативная база {notice.available_version}; "
                f"календарь использует {notice.calendar_version}."
            )
            choice_label = st.radio(
                "Выберите нормативную версию",
                ("Оставить текущую", "Применить актуальную версию"),
            )
            if st.button("Подтвердить выбор нормативной версии"):
                choice = (
                    NormativeUpdateChoice.APPLY_CURRENT
                    if choice_label == "Применить актуальную версию"
                    else NormativeUpdateChoice.KEEP_EXISTING
                )
                reference = resolve_registry_reference(
                    CalendarRegistryReference(saved_version), normative_registry, choice
                )
                st.session_state["calendar_registry_version"] = reference.registry_version
    st.markdown("</div>", unsafe_allow_html=True)


def _fact(value: object | None) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


_TABLE_WIDTH_WARNING = (
    "Ширина таблицы не задана явно; возможны переносы на новые страницы."
)


_INTERNAL_SLOT_WARNINGS = frozenset(
    {
        _TABLE_WIDTH_WARNING,
        SLOT_PACK_WARNING,
        SLOT_CONTINUE_WARNING,
    }
)

_CE2_SAFE_WARNING_RE = re.compile(
    r"^Безопасный шаблон CE2:\s*([a-z][a-z0-9_]*)\.$"
)
_CE2_SAFE_USER_MESSAGE = (
    "Некоторые формулировки автоматически приведены "
    "к безопасному нейтральному виду."
)


def _teacher_generation_warnings(warnings: tuple[str, ...]) -> tuple[str, ...]:
    visible: list[str] = []
    safe_codes: list[str] = []
    for warning in warnings:
        if warning in _INTERNAL_SLOT_WARNINGS:
            continue
        match = _CE2_SAFE_WARNING_RE.fullmatch(warning)
        if match:
            safe_codes.append(match.group(1))
            continue
        visible.append(warning)
    if safe_codes:
        logging.getLogger(__name__).info(
            "CE2 safe fallback diagnostics: %s",
            ", ".join(dict.fromkeys(safe_codes)),
        )
        visible.append(_CE2_SAFE_USER_MESSAGE)
    return tuple(visible)


def _collect_analysis_warnings(
    utp: UtpParseResult,
    matches: tuple[ContentMatch, ...],
    content_rows: tuple[CalendarContentRow, ...],
    lessons: tuple[LessonContentRow, ...],
) -> tuple[str, ...]:
    collected: list[str] = list(utp.warnings)
    for match in matches:
        if match.ambiguous_candidates:
            collected.append(
                f"Неоднозначное соответствие для «{match.utp_position.title}»: "
                + "; ".join(match.ambiguous_candidates)
            )
    collected.extend(warning for row in content_rows for warning in row.warnings)
    collected.extend(warning for row in lessons for warning in row.warnings)
    return tuple(dict.fromkeys(collected))


_LAYER_TITLES = (
    (NormativeLayer.FEDERAL, "Документы закона"),
    (NormativeLayer.LOCAL, "Календарь учреждения"),
    (NormativeLayer.METHODICAL, "Сверка ваших часов"),
)
_SHORT_WEEK_NOTE = (
    "1–6 сентября и 28–30 декабря — короткие недели, часы на них остаются. "
    "Даты приложение не сдвигает."
)
_METHODICAL_PASS_LINE = "Часы программы, УТП и плана совпадают."
_ATTESTATION_MISSING_UI = (
    "В программе указана аттестация, но в темах УТП она не найдена."
)
_YEAR_NO_DURATION_UI = "Срок программы не указан, сравнить год со сроком нельзя."


def _lesson_views_for_normative(
    content_rows: tuple[CalendarContentRow, ...],
) -> tuple[NormativeLessonView, ...]:
    rows = build_lesson_content_v2(content_rows)
    return tuple(
        NormativeLessonView(
            theory_hours=row.source.theory_hours,
            practice_hours=row.source.practice_hours,
            lesson_type=row.lesson_type,
            assessment_method=row.assessment_method,
            topic_title=row.source.topic_title,
        )
        for row in rows
    )


def _year_is_found(report: NormativeReport) -> bool:
    return any(
        item.check_id == "study_year_found" and item.verdict is NormativeVerdict.PASS
        for item in report.checks
    )


def _normative_teacher_text(item: NormativeCheck, report: NormativeReport) -> str:
    if item.check_id == "attestation" and "календарном плане и УТП" in item.teacher_text:
        return _ATTESTATION_MISSING_UI
    if (
        item.check_id == "year_within_duration"
        and item.verdict is NormativeVerdict.NOT_CHECKED
        and _year_is_found(report)
    ):
        return _YEAR_NO_DURATION_UI
    return item.teacher_text


def _is_short_week_warning(item: NormativeCheck) -> bool:
    return (
        item.check_id == "short_week_full_load"
        and item.verdict is NormativeVerdict.WARNING
    )


def _is_hidden_duration_warning(item: NormativeCheck, report: NormativeReport) -> bool:
    return (
        item.check_id == "duration_found"
        and item.verdict is NormativeVerdict.WARNING
        and _year_is_found(report)
        and any(
            other.check_id == "year_within_duration"
            and other.verdict is NormativeVerdict.NOT_CHECKED
            for other in report.checks
        )
    )


def _collapsed_pass_line(
    layer: NormativeLayer,
    passed: tuple[NormativeCheck, ...],
) -> str:
    if layer is NormativeLayer.METHODICAL:
        return _METHODICAL_PASS_LINE
    return passed[0].teacher_text


def _short_week_note(academic_year: str) -> str:
    if academic_year == APPROVED_ACADEMIC_YEAR:
        return _SHORT_WEEK_NOTE
    return "Короткие недели получили полную нагрузку. Даты приложение не сдвигает."


def _hours_match(report: NormativeReport) -> bool:
    return any(
        item.verdict is NormativeVerdict.PASS and item.check_id != "short_week_full_load"
        for item in report.for_layer(NormativeLayer.METHODICAL)
    )


def _calendar_checked(report: NormativeReport) -> bool:
    return any(
        item.verdict is NormativeVerdict.PASS
        for item in report.for_layer(NormativeLayer.LOCAL)
    )


def _visible_normative_remarks(
    report: NormativeReport,
) -> tuple[NormativeCheck, ...]:
    remarks: list[NormativeCheck] = []
    for item in report.checks:
        if _is_short_week_warning(item) or _is_hidden_duration_warning(item, report):
            continue
        if item.verdict in {NormativeVerdict.WARNING, NormativeVerdict.NOT_CHECKED}:
            remarks.append(item)
    return tuple(remarks)


def _hours_phrase(count: int) -> str:
    remainder_ten = count % 10
    remainder_hundred = count % 100
    if remainder_ten == 1 and remainder_hundred != 11:
        word = "час"
    elif remainder_ten in {2, 3, 4} and remainder_hundred not in {12, 13, 14}:
        word = "часа"
    else:
        word = "часов"
    return f"{count} {word}"


def _weeks_phrase(count: int) -> str:
    remainder_ten = count % 10
    remainder_hundred = count % 100
    if remainder_ten == 1 and remainder_hundred != 11:
        word = "неделя"
    elif remainder_ten in {2, 3, 4} and remainder_hundred not in {12, 13, 14}:
        word = "недели"
    else:
        word = "недель"
    return f"{count} {word}"


def _render_status_checks(report: NormativeReport) -> None:
    lines: list[str] = []
    if _hours_match(report):
        lines.append('<p class="kp-status-check">✓ Часы совпадают</p>')
    if _calendar_checked(report):
        lines.append('<p class="kp-status-check">✓ Календарь проверен</p>')
    remarks = _visible_normative_remarks(report)
    if remarks:
        lines.append(
            f'<p class="kp-status-check warn">⚠ {len(remarks)} замечаний '
            "— не мешают формированию</p>"
        )
    if lines:
        st.markdown(
            '<div class="kp-status-checks">' + "".join(lines) + "</div>",
            unsafe_allow_html=True,
        )


def _render_normative_report(report: NormativeReport, *, academic_year: str) -> None:
    sections: list[str] = [
        '<div class="kp-normative-check">',
        '<p class="kp-normative-check-title">Нормативная и методическая проверка</p>',
        '<p class="kp-normative-check-lead">'
        "Эта проверка не изменяет Word автоматически.</p>",
    ]
    for layer, title in _LAYER_TITLES:
        layer_checks = report.for_layer(layer)
        if not layer_checks:
            continue
        sections.append(f'<p class="kp-normative-layer">{html.escape(title)}</p>')
        passed = tuple(
            item
            for item in layer_checks
            if item.verdict is NormativeVerdict.PASS
            and item.check_id != "short_week_full_load"
        )
        warnings = tuple(
            item
            for item in layer_checks
            if item.verdict is NormativeVerdict.WARNING
            and not _is_short_week_warning(item)
            and not _is_hidden_duration_warning(item, report)
        )
        unchecked = tuple(
            item
            for item in layer_checks
            if item.verdict is NormativeVerdict.NOT_CHECKED
        )
        short_week_notes = tuple(
            item for item in layer_checks if _is_short_week_warning(item)
        )
        if passed:
            sections.append("<ul>")
            sections.append(
                f"<li>✓ {html.escape(_collapsed_pass_line(layer, passed))}</li>"
            )
            sections.append("</ul>")
        if short_week_notes:
            sections.append(
                f'<p class="kp-normative-note">{html.escape(_short_week_note(academic_year))}</p>'
            )
        if warnings:
            sections.append(
                '<p class="kp-normative-check-label warn">⚠ На что обратить внимание</p>'
            )
            sections.append("<ul>")
            sections.extend(
                f"<li>{html.escape(_normative_teacher_text(item, report))}</li>"
                for item in warnings
            )
            sections.append("</ul>")
        if unchecked:
            sections.append(
                '<p class="kp-normative-check-label skip">— Что не удалось проверить</p>'
            )
            sections.append("<ul>")
            sections.extend(
                f"<li>{html.escape(_normative_teacher_text(item, report))}</li>"
                for item in unchecked
            )
            sections.append("</ul>")
    sections.append("</div>")
    st.markdown("".join(sections), unsafe_allow_html=True)


def _render_teacher_analysis_screen(
    *,
    utp: UtpParseResult,
    program: ProgramData | None,
    schedule: ScheduleResult,
    matches: tuple[ContentMatch, ...],
    academic_year: str,
    source_utp_name: str | None = None,
    program_filename: str | None = None,
    content_rows: tuple[CalendarContentRow, ...] = (),
    after_summary: Callable[[], None] | None = None,
) -> None:
    metadata = utp.metadata
    totals = utp.table_totals
    weeks = metadata.study_weeks or len(schedule.weeks)
    report = evaluate_normative_mvp(
        utp,
        program,
        academic_year=academic_year,
        study_year_hints=(source_utp_name, program_filename),
        schedule=schedule,
        lessons=_lesson_views_for_normative(content_rows),
    )
    generated = bool(
        st.session_state.get("calendar_generation_succeeded")
        and st.session_state.get("calendar_download")
    )
    program_name = _fact(metadata.program_name or (program.title if program else None))
    study_year = _fact(
        study_year_label(
            metadata.study_year,
            source_utp_name,
            program_filename,
        )
    )

    st.markdown('<div class="kp-status-card">', unsafe_allow_html=True)
    title_col, edit_col = st.columns((3.4, 1.1), gap="small")
    with title_col:
        if generated:
            st.markdown(
                '<p class="kp-status-title">✓ Календарный план готов</p>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<p class="kp-status-title">✓ Документы проверены</p>',
                unsafe_allow_html=True,
            )
            st.markdown(
                '<p class="kp-status-lead">Можно сформировать календарный план</p>',
                unsafe_allow_html=True,
            )
    with edit_col:
        st.markdown('<div class="kp-edit-slot">', unsafe_allow_html=True)
        if st.button("Изменить данные", key="kp_edit_data"):
            _open_input_form()
        st.markdown("</div>", unsafe_allow_html=True)

    if program_name:
        st.markdown(
            f'<p class="kp-program-name">{html.escape(program_name)}</p>',
            unsafe_allow_html=True,
        )
    meta = academic_year
    if study_year:
        meta = f"{academic_year} · {study_year}"
    st.markdown(f'<p class="kp-status-meta">{html.escape(meta)}</p>', unsafe_allow_html=True)
    if totals:
        if generated:
            kpi = f"{_weeks_phrase(weeks)} · {_hours_phrase(totals.total)}"
        else:
            kpi = (
                f"{_weeks_phrase(weeks)} | {_hours_phrase(totals.total)} | "
                f"{totals.theory} ч теория | {totals.practice} ч практика"
            )
    else:
        kpi = _weeks_phrase(weeks)
    st.markdown(f'<p class="kp-status-kpi">{html.escape(kpi)}</p>', unsafe_allow_html=True)

    if not generated:
        _render_status_checks(report)
        with st.expander("Подробнее о проверке", expanded=False):
            _render_normative_report(report, academic_year=academic_year)

    if after_summary is not None:
        after_summary()
    _render_year_calendar_card(academic_year, owner="analysis")
    st.markdown("</div>", unsafe_allow_html=True)


def _store_analysis_context(
    *,
    validated_utp: ValidatedUpload,
    validated_program: ValidatedUpload | None,
    template_selection: CalendarTemplateSelection,
    academic_year: str,
) -> None:
    st.session_state["calendar_context"] = {
        "validated_utp": validated_utp,
        "validated_program": validated_program,
        "template_selection": template_selection,
        "academic_year": academic_year,
    }


def _show_generation_controls(
    *,
    validated_utp: ValidatedUpload,
    validated_program: ValidatedUpload | None,
    template_selection: CalendarTemplateSelection,
    academic_year: str,
    group_number: str,
    class_name: str,
    teacher_name: str,
) -> None:
    current_revision = _generator_revision()
    if current_revision != _LOADED_GENERATOR_REVISION:
        logging.getLogger(__name__).warning(
            "Generator revision mismatch: loaded=%s current=%s",
            _LOADED_GENERATOR_REVISION, current_revision,
        )
        st.button(
            "Сформировать календарный план",
            type="primary",
            use_container_width=True,
            disabled=True,
            key="generate_calendar",
        )
        st.info("Приложение обновилось. Обновите страницу, чтобы продолжить.")
        _show_generation_result()
        return

    generation_pending = bool(st.session_state.get("calendar_generation_pending"))
    generated = bool(
        st.session_state.get("calendar_generation_succeeded")
        and st.session_state.get("calendar_download")
    )
    if generated:
        _show_generation_result()
        regenerate = st.button(
            "Сформировать заново",
            use_container_width=True,
            disabled=generation_pending,
            key="regenerate_calendar",
        )
        generation_requested = regenerate
    else:
        generation_requested = st.button(
            "Сформировать календарный план",
            type="primary",
            use_container_width=True,
            disabled=generation_pending,
            key="generate_calendar",
        )
    if generation_requested:
        st.session_state.pop("calendar_generation_invalidated", None)
        st.session_state["calendar_generation_pending"] = True
        st.session_state.pop("calendar_generation_error", None)
        st.session_state.pop("calendar_generation_succeeded", None)
        st.session_state.pop("calendar_resolved_lessons", None)
        st.session_state.pop("calendar_plan_snapshot", None)
        st.session_state.pop("calendar_download", None)
        st.session_state.pop("calendar_warnings", None)
        st.session_state.pop("calendar_ai_usage", None)
        st.rerun()

    if generation_pending:
        utp = validated_utp.parsed
        assert isinstance(utp, UtpParseResult)
        program = None
        if validated_program is not None:
            program = validated_program.parsed
            assert isinstance(program, ProgramData)

        try:
            with st.spinner("Формируем календарный план…"):
                with TransientDocumentSession() as operation:
                    result = run_calendar_pipeline(
                        utp,
                        program,
                        academic_year=academic_year,
                        template=template_selection,
                        source_utp_name=validated_utp.filename,
                        use_ai=False,
                        program_filename=(
                            validated_program.filename
                            if validated_program is not None
                            else None
                        ),
                        group_number=group_number,
                        class_name=class_name,
                        teacher_name=teacher_name,
                    )
                    operation.publish_result(result.filename, result.content)
                    st.session_state["calendar_download"] = operation.take_result_for_download()
                    st.session_state["calendar_warnings"] = result.warnings
                    resolved_lessons = tuple(
                        getattr(result, "resolved_lessons", ())
                    )
                    st.session_state["calendar_resolved_lessons"] = resolved_lessons
                    st.session_state["calendar_plan_snapshot"] = (
                        _calendar_plan_snapshot(resolved_lessons, result.content)
                    )
        except (PipelineError, ScheduleValidationError, ValueError) as error:
            st.session_state["calendar_generation_error"] = str(error)
        else:
            st.session_state["calendar_generation_succeeded"] = True
        finally:
            st.session_state["calendar_generation_pending"] = False
        st.rerun()

    if not generated:
        _show_generation_result()


@st.fragment(run_every="2s")
def _show_generation_result() -> None:
    inputs = st.session_state.get("calendar_generation_inputs", "")
    if _sync_generation_fingerprint((inputs, _generator_revision())):
        _reset_analysis_state()
    if st.session_state.get("calendar_generation_invalidated"):
        st.info("Сформируйте календарный план заново.")
    generation_error = st.session_state.get("calendar_generation_error")
    if generation_error:
        st.error(f"Не удалось сформировать календарный план: {generation_error}")

    for warning in _teacher_generation_warnings(
        tuple(st.session_state.get("calendar_warnings", ()))
    ):
        st.warning(warning)

    download = st.session_state.get("calendar_download")
    if download is not None:
        context = st.session_state.get("calendar_context") or {}
        academic_year = str(context.get("academic_year") or APPROVED_ACADEMIC_YEAR)
        st.download_button(
            f"Скачать план за {academic_year} учебный год",
            data=download.content,
            file_name=download.filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )


def run_app() -> None:
    """Показать экран загрузки, анализа и формирования календарного плана."""
    st.set_page_config(
        page_title="Календарь педагога",
        page_icon="📅",
        layout="wide",
        initial_sidebar_state="collapsed",
    )

    (
        utp_file,
        program_file,
        organization_template_file,
        academic_year,
        group_number,
        class_name,
        teacher_name,
        check_clicked,
    ) = _render_upload_screen()

    _refresh_generation_inputs(
        utp_file, program_file, organization_template_file,
        academic_year, group_number, class_name, teacher_name,
    )
    if st.session_state.get("calendar_generation_invalidated") and not st.session_state.get("analysis_ready"):
        st.info("План устарел. Нажмите «Проверить документы» и сформируйте календарный план заново.")

    if check_clicked:
        if program_file is None:
            st.error("Загрузите программу обучения.")
            return

        with TransientDocumentSession() as uploads:
            uploads.replace(
                UploadPurpose.PROGRAM,
                program_file.name,
                program_file.getvalue(),
            )
            if utp_file is not None:
                uploads.replace(UploadPurpose.UTP, utp_file.name, utp_file.getvalue())
            if organization_template_file is not None:
                uploads.replace(
                    UploadPurpose.CALENDAR_TEMPLATE,
                    organization_template_file.name,
                    organization_template_file.getvalue(),
                )
            try:
                transient_program = uploads.get(UploadPurpose.PROGRAM)
                assert transient_program is not None
                validated_program = validate_upload(
                    UploadPurpose.PROGRAM,
                    transient_program.filename,
                    transient_program.content,
                )
                transient_utp = uploads.get(UploadPurpose.UTP)
                validated_utp_upload = (
                    validate_upload(
                        UploadPurpose.UTP,
                        transient_utp.filename,
                        transient_utp.content,
                    )
                    if transient_utp is not None
                    else None
                )
                transient_template = uploads.get(UploadPurpose.CALENDAR_TEMPLATE)
                validated_template = (
                    validate_upload(
                        UploadPurpose.CALENDAR_TEMPLATE,
                        transient_template.filename,
                        transient_template.content,
                    )
                    if transient_template is not None
                    else None
                )
                resolved_utp = resolve_utp(validated_utp_upload, validated_program)
            except UploadValidationError as error:
                st.error(str(error))
                return
            except UtpResolutionError as error:
                st.error(str(error))
                return

        template_selection = select_calendar_template()
        if validated_template is not None:
            try:
                template_selection = select_calendar_template(
                    validated_template.filename,
                    validated_template.content,
                )
            except OrganizationTemplateError:
                st.error(ORG_TEMPLATE_UNSUPPORTED_MESSAGE)
                return

        validated_utp = ValidatedUpload(
            UploadPurpose.UTP,
            (
                validated_utp_upload.filename
                if validated_utp_upload is not None
                else f"УТП из файла «{validated_program.filename}»"
            ),
            (
                validated_utp_upload.content
                if validated_utp_upload is not None
                else validated_program.content
            ),
            resolved_utp,
        )
        program = parse_program(
            validated_program.content,
            validated_program.filename,
            study_year=infer_study_year_number(resolved_utp.metadata.study_year),
        )
        validated_program = ValidatedUpload(
            validated_program.purpose,
            validated_program.filename,
            validated_program.content,
            program,
        )
        utp = resolved_utp

        try:
            build_schedule(utp, academic_year)
        except (ScheduleValidationError, ValueError) as error:
            st.error(f"Не удалось построить календарное распределение: {error}")
            return

        _store_analysis_context(
            validated_utp=validated_utp,
            validated_program=validated_program,
            template_selection=template_selection,
            academic_year=academic_year,
        )
        st.session_state["analysis_ready"] = True
        st.session_state.pop("analysis_warnings", None)
        st.session_state.pop("calendar_download", None)
        st.session_state.pop("calendar_warnings", None)
        st.session_state.pop("calendar_ai_usage", None)
        st.session_state.pop("calendar_generation_pending", None)
        st.session_state.pop("calendar_generation_error", None)
        st.session_state.pop("calendar_generation_succeeded", None)
        st.session_state["ui_edit_inputs"] = False
        st.rerun()

    if (
        st.session_state.get("analysis_ready")
        and "calendar_context" in st.session_state
        and not _form_is_open()
    ):
        context = st.session_state["calendar_context"]
        validated_utp = context["validated_utp"]
        validated_program = context["validated_program"]
        template_selection = context["template_selection"]
        academic_year = context["academic_year"]

        utp = validated_utp.parsed
        assert isinstance(utp, UtpParseResult)
        program = None
        if validated_program is not None:
            program = validated_program.parsed
            assert isinstance(program, ProgramData)

        try:
            schedule = build_schedule(utp, academic_year)
            content_rows = build_content_model(
                schedule,
                utp,
                program,
                validated_utp.filename,
            )
            lessons = build_lesson_content(content_rows)
            matches = (
                tuple(match_utp_to_program(utp.topics, program.content_items))
                if program is not None
                else ()
            )
        except (ScheduleValidationError, ValueError) as error:
            st.error(f"Не удалось построить календарное распределение: {error}")
            st.session_state["analysis_ready"] = False
            return

        detail_warnings = st.session_state.get("analysis_warnings")
        if detail_warnings is None:
            detail_warnings = _collect_analysis_warnings(
                utp,
                matches,
                content_rows,
                lessons,
            )
            st.session_state["analysis_warnings"] = detail_warnings

        _render_teacher_analysis_screen(
            utp=utp,
            program=program,
            schedule=schedule,
            matches=matches,
            academic_year=academic_year,
            source_utp_name=validated_utp.filename,
            program_filename=(
                validated_program.filename if validated_program is not None else None
            ),
            content_rows=content_rows,
            after_summary=lambda: _show_generation_controls(
                validated_utp=validated_utp,
                validated_program=validated_program,
                template_selection=template_selection,
                academic_year=academic_year,
                group_number=group_number,
                class_name=class_name,
                teacher_name=teacher_name,
            ),
        )
