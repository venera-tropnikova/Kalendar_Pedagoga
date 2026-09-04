"""Безопасное чтение образовательных программ DOCX и legacy DOC."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
import shutil
import subprocess
import tempfile

from docx import Document


class LegacyDocUnsupportedError(RuntimeError):
    """Legacy DOC нельзя безопасно прочитать без LibreOffice."""


@dataclass(frozen=True)
class ProgramContentItem:
    number: str | None
    title: str
    content: str
    parent_section: str | None = None


@dataclass(frozen=True)
class ProgramData:
    title: str | None
    duration: str | None
    student_age: str | None
    goal: str | None
    tasks: tuple[str, ...]
    lesson_forms: tuple[str, ...]
    teaching_methods: tuple[str, ...]
    expected_results: tuple[str, ...]
    knowledge_outcomes: tuple[str, ...]
    skill_outcomes: tuple[str, ...]
    content_items: tuple[ProgramContentItem, ...]
    duration_years: int | None = None
    age_min: int | None = None
    age_max: int | None = None
    attestation_statements: tuple[str, ...] = ()
    academic_year: str | None = None
    academic_year_mentions: tuple = ()


def find_libreoffice() -> Path | None:
    """Найти LibreOffice без изменения системной конфигурации."""
    executable = shutil.which("soffice") or shutil.which("soffice.exe")
    candidates = [
        executable,
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return Path(candidate)
    return None


def convert_legacy_doc(data: bytes, soffice: Path | None = None) -> bytes:
    """Конвертировать временную копию DOC в DOCX и удалить все временные файлы."""
    converter = soffice or find_libreoffice()
    if converter is None:
        raise LegacyDocUnsupportedError(
            "Формат legacy .DOC пока не поддерживается: LibreOffice не найден. "
            "Исходный файл не изменён."
        )
    with tempfile.TemporaryDirectory(prefix="calendar_pedagoga_") as temp_name:
        temp_dir = Path(temp_name)
        source = temp_dir / "program.doc"
        source.write_bytes(data)
        profile = temp_dir / "lo_profile"
        output = temp_dir / "program.docx"
        command = [
            str(converter),
            "--headless",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to",
            "docx",
            "--outdir",
            str(temp_dir),
            str(source),
        ]
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        if completed.returncode != 0 or not output.is_file():
            message = (completed.stderr or completed.stdout).strip()
            raise LegacyDocUnsupportedError(
                "LibreOffice не смог преобразовать legacy .DOC в DOCX"
                + (f": {message}" if message else ".")
            )
        return output.read_bytes()


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _first_match(text: str, patterns: tuple[str, ...]) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return _clean(match.group(1))
    return None


def _number_title(value: str) -> tuple[str | None, str]:
    match = re.match(r"^(\d+(?:\.\d+)*)\.?\s+(.+)$", _clean(value))
    return (match.group(1), _clean(match.group(2))) if match else (None, _clean(value))


def _collect_after(
    paragraphs: list[str],
    heading_pattern: str,
    stop_pattern: str,
) -> tuple[str, ...]:
    result: list[str] = []
    active = False
    for paragraph in paragraphs:
        if re.search(heading_pattern, paragraph, re.IGNORECASE):
            active = True
            remainder = re.sub(
                rf"^.*?{heading_pattern}\s*[:.-]?\s*",
                "",
                paragraph,
                flags=re.IGNORECASE,
            )
            if _clean(remainder):
                result.append(_clean(remainder))
            continue
        if active and re.search(stop_pattern, paragraph, re.IGNORECASE):
            break
        if active and _clean(paragraph):
            result.append(_clean(re.sub(r"^[•\-–—]\s*", "", paragraph)))
    return tuple(result)


def _is_bold_heading(paragraph) -> bool:
    runs = [run for run in paragraph.runs if _clean(run.text)]
    return bool(runs) and all(run.bold is True for run in runs)


def _find_content_start(paragraphs, study_year: int | None) -> int | None:
    texts = [_clean(paragraph.text) for paragraph in paragraphs]
    if study_year is not None:
        specific = next(
            (
                index
                for index, text in enumerate(texts)
                if re.search(
                    rf"содержание\s+программы\s+{study_year}-го\s+года\s+обучения",
                    text,
                    re.I,
                )
            ),
            None,
        )
        if specific is not None:
            return specific
    numbered = next(
        (
            index
            for index, text in enumerate(texts)
            if re.search(
                r"содержание\s+программы\s+\d+-го\s+года\s+обучения",
                text,
                re.I,
            )
        ),
        None,
    )
    if numbered is not None:
        return numbered
    return next(
        (
            index
            for index, text in enumerate(texts)
            if re.search(r"содержание\s+программы", text, re.I)
        ),
        None,
    )


def _year_number_from_text(text: str) -> int | None:
    low = text.casefold()
    for token, number in (
        ("перв", 1),
        ("втор", 2),
        ("трет", 3),
        ("четв", 4),
    ):
        if token in low:
            return number
    found = re.search(r"\b(\d+)\s*[-–—]?\s*го?\s*года", low)
    if found:
        return int(found.group(1))
    found = re.search(r"\b(\d+)\s+года", low)
    if found:
        return int(found.group(1))
    return None


def _split_outcome_items(text: str) -> list[str]:
    """Разбить абзац исхода на отдельные формулировки без выдумок."""

    cleaned = _clean(re.sub(r"^[•\-–—]\s*", "", text)).strip(" ;.")
    if not cleaned:
        return []
    if re.match(r"(?i)^(?:знать|уметь|знания|умения(?:\s+и\s+навыки)?)\s*:?\s*$", cleaned):
        return []
    parts = re.split(r"[;]\s*", cleaned)
    items: list[str] = []
    for part in parts:
        item = _clean(part).strip(" ;.")
        if not item:
            continue
        if re.match(r"(?i)^\d+[\).]\s*", item) and len(item) < 12:
            continue
        items.append(item)
    return items


def _year_end_outcomes(
    paragraphs: list[str],
    study_year: int | None,
) -> tuple[tuple[str, ...], tuple[str, ...]]:
    """Извлечь «должны знать / должны уметь» для выбранного года обучения."""

    knowledge: list[str] = []
    skills: list[str] = []
    mode: str | None = None
    locked_year: int | None = study_year

    year_heading = re.compile(
        r"(?i)по\s+окончани[юя].*года\s+обучен\w*.*должн",
    )
    know_inline = re.compile(r"(?i)должны?\s+знать\s*:?\s*(.*)$")
    skill_inline = re.compile(r"(?i)должны?\s+уметь\s*:?\s*(.*)$")
    know_label = re.compile(r"(?i)^(?:знать|знания)\s*:?\s*$")
    skill_label = re.compile(r"(?i)^(?:уметь|умения(?:\s+и\s+навыки)?)\s*:?\s*$")

    def year_allowed(year: int | None) -> bool:
        if locked_year is None:
            return True
        return year is None or year == locked_year

    for paragraph in paragraphs:
        text = _clean(paragraph)
        if not text:
            continue

        if year_heading.search(text):
            heading_year = _year_number_from_text(text)
            if locked_year is None and heading_year is not None:
                locked_year = heading_year
            if not year_allowed(heading_year):
                mode = None
                continue
            know_match = know_inline.search(text)
            skill_match = skill_inline.search(text)
            if know_match:
                mode = "knowledge"
                knowledge.extend(_split_outcome_items(know_match.group(1)))
            elif skill_match:
                mode = "skills"
                skills.extend(_split_outcome_items(skill_match.group(1)))
            else:
                mode = "await_label"
            continue

        if mode is None:
            continue

        if re.search(
            r"(?i)^(?:учебно[- ]тематическ|содержание\s+программы|тесты\s+для|"
            r"список\s+литератур|(?:ожидаемые|планируемые)\s+результаты)",
            text,
        ):
            mode = None
            continue

        if know_label.match(text):
            mode = "knowledge"
            continue
        if skill_label.match(text):
            mode = "skills"
            continue

        if mode == "await_label":
            continue
        if mode == "knowledge":
            knowledge.extend(_split_outcome_items(text))
        elif mode == "skills":
            skills.extend(_split_outcome_items(text))

    return tuple(dict.fromkeys(knowledge)), tuple(dict.fromkeys(skills))


def _structured_expected_outcomes(
    paragraphs: list[str],
) -> tuple[tuple[str, ...], tuple[str, ...]]:
    """Подразделы знаний/умений внутри «Ожидаемые результаты…», без тестов."""

    knowledge: list[str] = []
    skills: list[str] = []
    mode: str | None = None
    active = False
    for paragraph in paragraphs:
        text = _clean(paragraph)
        if not text:
            continue
        if re.search(r"(?i)^(?:ожидаемые|планируемые)\s+результаты", text):
            active = True
            mode = None
            continue
        if not active:
            continue
        if re.search(
            r"(?i)^(?:тесты\s+для|учебно[- ]тематическ|содержание\s+программы|"
            r"по окончани[юя]|список\s+литератур)",
            text,
        ):
            break
        if re.match(r"(?i)^(?:знания|знать)\s*:?\s*$", text):
            mode = "knowledge"
            continue
        if re.match(r"(?i)^(?:умения(?:\s+и\s+навыки)?|уметь)\s*:?\s*$", text):
            mode = "skills"
            continue
        if re.match(r"(?i)^\d+[\).]\s+", text):
            # Тестовые варианты/вопросы.
            mode = None
            continue
        if mode == "knowledge":
            knowledge.extend(_split_outcome_items(text))
        elif mode == "skills":
            skills.extend(_split_outcome_items(text))
    return tuple(dict.fromkeys(knowledge)), tuple(dict.fromkeys(skills))


def _content_items(document, study_year: int | None = None) -> tuple[ProgramContentItem, ...]:
    paragraphs = document.paragraphs
    start = _find_content_start(paragraphs, study_year)
    if start is None:
        return ()
    items: list[ProgramContentItem] = []
    current_number: str | None = None
    current_title: str | None = None
    current_parent: str | None = None
    content: list[str] = []

    def flush() -> None:
        if current_title:
            items.append(
                ProgramContentItem(
                    current_number,
                    current_title,
                    "\n".join(content).strip(),
                    current_parent,
                )
            )

    for paragraph in paragraphs[start + 1 :]:
        text = _clean(paragraph.text)
        if not text:
            continue
        if re.search(r"^\(\d+-й\s+год\s+обучения\)", text, re.I):
            continue
        if re.search(r"по окончани[юя].*года обучения", text, re.I):
            break
        if re.search(r"^(?:ожидаемые|планируемые)\s+результаты", text, re.I):
            break
        number, title = _number_title(text)
        if _is_bold_heading(paragraph):
            flush()
            current_number, current_title, content = number, title, []
            if number and "." not in number:
                current_parent = title
            elif number is None and paragraph.alignment == 1:
                current_parent = title
            continue
        if current_title is None:
            continue
        content.append(text)
    flush()
    return tuple(items)


def parse_program_docx(data: bytes, study_year: int | None = None) -> ProgramData:
    """Извлечь только явно присутствующие формулировки из DOCX."""
    document = Document(BytesIO(data))
    paragraphs = [_clean(p.text) for p in document.paragraphs if _clean(p.text)]
    text = "\n".join(paragraphs)
    stop = (
        r"цель|задач|формы?\s+организац|методы?\s+обучен|"
        r"ожидаем|планируем|содержание"
    )
    tasks = _collect_after(paragraphs, r"задачи(?: программы)?", r"хочу заметить|срок реализации|" + stop)
    forms = tuple(_clean(match.group(1)) for paragraph in paragraphs if (match := re.match(r"формы?\s+организации\s+(?:образовательного\s+процесса|занятий)\s*[–—:-]\s*(.+)", paragraph, re.I)))
    methods = tuple(_clean(match.group(1)) for paragraph in paragraphs if (match := re.match(r"методы?\s+обучения\s*[:.-]\s*(.+)", paragraph, re.I)))
    results = tuple(paragraph for paragraph in paragraphs if re.match(r"за период освоения.*ожидается", paragraph, re.I))
    if not results:
        results_stop = (
            r"цель|задач|формы?\s+организац|методы?\s+обучен|содержание|"
            r"учебно[- ]тематическ|тесты\s+для|по окончани[юя]|список\s+литератур"
        )
        results = _collect_after(
            paragraphs,
            r"(?:ожидаемые|планируемые)\s+результаты(?:\s+освоения(?:\s+программы)?)?",
            results_stop,
        )
        results = tuple(
            item
            for item in results
            if item
            and not re.match(r"(?i)^\d+[\).]\s+", item)
            and "тест" not in item.casefold()
        )
    year_knowledge, year_skills = _year_end_outcomes(paragraphs, study_year)
    expected_knowledge, expected_skills = _structured_expected_outcomes(paragraphs)
    knowledge = tuple(dict.fromkeys([*year_knowledge, *expected_knowledge]))
    skills = tuple(dict.fromkeys([*year_skills, *expected_skills]))
    academic_year, academic_year_mentions = _program_academic_year(text)
    duration_text = _first_match(
        text,
        (
            r"срок реализации(?: программы)?\s*[:.-]\s*([^\n]+)",
            r"программа рассчитана на\s*([^\n]+)",
        ),
    )
    student_age = _first_match(
        text,
        (
            r"возраст (?:обучающихся|учащихся|детей)\s*[:.-]\s*([^\n]+)",
        ),
    )
    age_min, age_max = parse_age_range(student_age)
    return ProgramData(
        title=_first_match(
            text,
            (
                r"(?:программа|направленность)\s*[«\"]([^»\"]+)[»\"]",
                r"название программы\s*[:.-]\s*([^\n]+)",
                r"(?m)^[«\"]([^»\"]+)[»\"]$",
            ),
        ),
        duration=duration_text,
        student_age=student_age,
        goal=_first_match(
            text,
            (r"цель(?: программы)?\s*[:.-]\s*([^\n]+)",),
        ),
        tasks=tasks,
        lesson_forms=forms,
        teaching_methods=methods,
        expected_results=results,
        knowledge_outcomes=knowledge,
        skill_outcomes=skills,
        content_items=_content_items(document, study_year),
        duration_years=parse_duration_years(duration_text),
        age_min=age_min,
        age_max=age_max,
        attestation_statements=extract_attestation_statements(text),
        academic_year=academic_year,
        academic_year_mentions=academic_year_mentions,
    )


_ATTESTATION_PHRASE = re.compile(
    r"(?i)(?:обязательн\w{0,12}\s+)?(?:промежуточн\w+|итогова\w+)\s+аттестац\w+"
)


def extract_attestation_statements(text: str) -> tuple[str, ...]:
    """Вернуть только явные упоминания промежуточной или итоговой аттестации."""

    if not text:
        return ()
    found: list[str] = []
    for paragraph in text.splitlines():
        cleaned = _clean(paragraph)
        if cleaned and _ATTESTATION_PHRASE.search(cleaned):
            found.append(cleaned)
    return tuple(dict.fromkeys(found))


def parse_duration_years(value: str | None) -> int | None:
    """Число лет только из явного «N год/года/лет» в начале строки."""

    if not value:
        return None
    found = re.match(
        r"\s*(\d+)\s+(год(?:а|ов)?|лет)\b",
        value,
        flags=re.IGNORECASE,
    )
    if found is None:
        return None
    years = int(found.group(1))
    return years if years > 0 else None


def parse_age_range(value: str | None) -> tuple[int | None, int | None]:
    """Диапазон возраста только из явного «N–M лет» или «от N до M лет»."""

    if not value:
        return None, None
    text = value.casefold()
    found = re.search(
        r"(\d+)\s*[-–—]\s*(\d+)\s*(?:лет|года|год)?\b",
        text,
    )
    if found is None:
        found = re.search(
            r"от\s+(\d+)\s+до\s+(\d+)\s*(?:лет|года|год)?\b",
            text,
        )
    if found is None:
        return None, None
    low = int(found.group(1))
    high = int(found.group(2))
    if low <= 0 or high <= 0 or low > high or high > 80:
        return None, None
    return low, high


def infer_study_year_number(value: str | None) -> int | None:
    """Преобразовать «второй», «2 год» и т.п. в номер года обучения."""
    if not value:
        return None
    text = value.casefold()
    for token, number in (
        ("перв", 1),
        ("втор", 2),
        ("трет", 3),
        ("четв", 4),
    ):
        if token in text:
            return number
    found = re.search(r"\d+", value)
    return int(found.group()) if found else None


def study_year_label(
    metadata_year: str | None,
    *hints: str | None,
) -> str | None:
    """Тот же год обучения, что подставляет pipeline в DOCX, или None."""

    for raw in (metadata_year, *hints):
        number = infer_study_year_number(raw)
        if number is not None:
            return f"{number} год обучения"
    raw = (metadata_year or "").strip()
    if raw:
        return raw if "год" in raw.casefold() else f"{raw} год обучения"
    return None


def _program_academic_year(text: str) -> tuple[str | None, tuple]:
    from calendar_pedagoga.academic_year import (
        extract_academic_year_mentions,
        unique_academic_years,
    )

    mentions = extract_academic_year_mentions(text)
    years = unique_academic_years(mentions)
    return (years[0] if len(years) == 1 else None), mentions


def parse_program(
    data: bytes,
    filename: str,
    study_year: int | None = None,
) -> ProgramData:
    """Разобрать загруженную программу с безопасной обработкой расширения."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".doc":
        return parse_program_docx(convert_legacy_doc(data), study_year=study_year)
    if suffix == ".docx":
        return parse_program_docx(data, study_year=study_year)
    raise ValueError("Поддерживаются только файлы DOC и DOCX.")
