"""Генерация итогового календарного плана DOCX на основе шаблона и данных УТП."""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from math import ceil
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from calendar_pedagoga.lesson_resolution import ResolvedLessonRow
from calendar_pedagoga.lesson_display import (
    format_practice_cell,
    format_theory_cell,
    practice_clause_for_repeated_topic,
)
from calendar_pedagoga.organization_template import CalendarTemplateSelection, CalendarTemplateSource
from calendar_pedagoga.parsing import UtpParseResult
from calendar_pedagoga.content_generation import WeekTopicPart
from calendar_pedagoga.program_parsing import infer_study_year_number


_PROJECT_ROOT = Path(__file__).resolve().parents[2]
STANDARD_TEMPLATE_PATH = _PROJECT_ROOT / "references" / "Календарный план Образец.docx"
STANDARD_TABLE_FONT_FAMILY = "Times New Roman"
STANDARD_GROUP_SPACE_AFTER_PT = 8
_MONTH_CELL_MARGIN_DXA = 40
PRINT_TOP_MARGIN_CM = 1.0


@dataclass(frozen=True)
class _TableColumns:
    month: int
    week: int
    theory: int
    theory_mark: int
    practice: int
    lesson_type: int
    planned_result: int
    assessment: int
    lesson_type_mirror: int | None = None
    planned_result_mirror: int | None = None


def _columns_for_table(table) -> _TableColumns:
    count = len(table.columns)
    if count >= 10:
        return _TableColumns(0, 1, 2, 3, 4, 5, 7, 9, 6, 8)
    return _TableColumns(0, 1, 2, 3, 4, 5, 6, 7)


def _prevent_row_split(row) -> None:
    """Не разрывать строку таблицы между страницами."""

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _allow_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for marker in list(tr_pr.findall(qn("w:cantSplit"))):
        tr_pr.remove(marker)
    for marker in row._tr.xpath(".//w:pPr/w:pageBreakBefore"):
        marker.getparent().remove(marker)


def _protect_vertical_cell_height(table) -> None:
    """Give rotated identifiers an intrinsic minimum, never an exact height.

    Word's automatic row sizing can ignore the inline extent of rotated text.
    Measure explicit lines without wrapping: after rotation their width is the
    required row height. Paragraph/line spacing lies on the perpendicular axis;
    retaining an unwrapped line prevents Word from adding a clipped third line.
    """
    import pymupdf

    table_styles = []
    style = table.style
    while style is not None:
        table_styles.append(style.element)
        style = style.base_style
    defaults = table.part.document.styles.element

    def value(elements, path, attribute, default):
        for element in elements:
            if element is not None:
                matches = element.xpath(path)
                if matches and matches[0].get(qn('w:' + attribute)) is not None:
                    return matches[0].get(qn('w:' + attribute))
        return default

    fonts = {}
    for row in table.rows[2:]:
        minimum = 0.0
        for cell in row.cells:
            direction = value([cell._tc], './w:tcPr/w:textDirection', 'val', '')
            if direction not in ('btLr', 'tbRl') or not cell.text.strip():
                continue
            margins = {}
            for side in ('left', 'right', 'top', 'bottom'):
                inherited = value([table._tbl] + table_styles,
                                  './w:tblPr/w:tblCellMar/w:' + side, 'w', '0')
                margins[side] = float(value([cell._tc],
                    './w:tcPr/w:tcMar/w:' + side, 'w', inherited)) / 20
            for paragraph in cell.paragraphs:
                styles = []
                style = paragraph.style
                while style is not None:
                    styles.append(style.element)
                    style = style.base_style
                inherited = styles + table_styles
                line_width = widest = largest = 0.0
                for run in paragraph.runs:
                    run_styles = []
                    style = run.style
                    while style is not None:
                        run_styles.append(style.element)
                        style = style.base_style
                    sources = [run._r] + run_styles + inherited
                    default_size = value([defaults], './w:docDefaults/w:rPrDefault/w:rPr/w:sz', 'val', '24')
                    size = float(value(sources, './w:rPr/w:sz', 'val', default_size)) / 2
                    family = value(sources, './w:rPr/w:rFonts', 'ascii', 'Times New Roman').lower()
                    bold = value(sources, './w:rPr/w:b', 'val', '0') not in ('0', 'false', 'off')
                    # MuPDF supplies portable serif/sans metrics and Unicode
                    # fallback glyphs; no installed desktop font is required.
                    face = ('hebo' if bold else 'helv') if any(
                        name in family for name in ('arial', 'calibri', 'sans')) else ('tibo' if bold else 'tiro')
                    if face not in fonts:
                        fonts[face] = pymupdf.Font(face)
                    largest = max(largest, size)
                    for index, text in enumerate(run.text.replace('\r', '\n').split('\n')):
                        if index:
                            widest = max(widest, line_width)
                            line_width = 0
                        line_width += fonts[face].text_length(text, fontsize=size)
                widest = max(widest, line_width)
                # Include indents in the rotated inline axis. Line/paragraph
                # spacing is not added to this axis; it is preserved unchanged.
                indents = sum(float(value([paragraph._p] + inherited,
                    './w:pPr/w:ind', side, '0')) / 20 for side in ('left', 'right'))
                # A small font-relative allowance covers glyph overhang and
                # renderer rounding, in addition to the inherited cell padding.
                minimum = max(minimum, widest + indents + margins['left'] +
                              margins['right'] + margins['top'] + margins['bottom'] + largest / 4)
        if minimum:
            properties = row._tr.get_or_add_trPr()
            height = properties.find(qn('w:trHeight'))
            if height is None:
                height = OxmlElement('w:trHeight')
                properties.append(height)
            height.set(qn('w:val'), str(max(ceil(minimum * 20), int(height.get(qn('w:val'), '0')))))
            height.set(qn('w:hRule'), 'atLeast')


def _repeat_table_header_rows(table, header_row_count: int = 2) -> None:
    """Повторять шапку таблицы на каждой странице Word через w:tblHeader."""

    header_tag = qn("w:tblHeader")
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        existing = tr_pr.find(header_tag)
        if index < header_row_count:
            if existing is None:
                existing = OxmlElement("w:tblHeader")
                tr_pr.append(existing)
            for attr in list(existing.attrib):
                del existing.attrib[attr]
        elif existing is not None:
            tr_pr.remove(existing)


def _topic_display_numbers(utp: UtpParseResult) -> dict[tuple[str | None, str, str], str]:
    numbers: dict[tuple[str | None, str, str], str] = {}
    for section in utp.sections:
        child_index = 0
        for topic in utp.topics:
            section_name = topic.parent_section or topic.title
            if section_name != section.title:
                continue
            key = (topic.number, topic.title, section_name)
            if topic.number:
                numbers[key] = topic.number.rstrip(".")
            elif topic.is_standalone_section:
                numbers[key] = (section.number or "?").rstrip(".")
            else:
                child_index += 1
                section_number = (section.number or "?").rstrip(".")
                numbers[key] = f"{section_number}.{child_index}"
    return numbers


def _hours_word(hours: int) -> str:
    if hours == 1:
        return "час"
    if hours in {2, 3, 4}:
        return "часа"
    return "часов"


def _study_year_label(
    utp: UtpParseResult,
    study_year_hints: tuple[str | None, ...] = (),
) -> str:
    for raw in (utp.metadata.study_year, *study_year_hints):
        number = infer_study_year_number(raw)
        if number is not None:
            return f"{number} год обучения"
    raw = (utp.metadata.study_year or "").strip()
    if raw:
        return raw if "год" in raw.casefold() else f"{raw} год обучения"
    return "____ год обучения"


def _resolve_header_from_rows(
    rows: tuple[ResolvedLessonRow, ...],
    *,
    program_title: str | None,
    study_year_hints: tuple[str | None, ...],
) -> tuple[str | None, tuple[str, ...]]:
    """Название и год обучения брать из строк занятия, не только из метаданных УТП."""

    names: list[str] = []
    hints: list[str] = []
    for raw in study_year_hints:
        cleaned = (raw or "").strip()
        if cleaned and cleaned not in hints:
            hints.append(cleaned)
    for row in rows:
        content = row.source.source
        name = (content.source_program_name or "").strip()
        if name and name not in names:
            names.append(name)
        utp_name = (content.source_utp_name or "").strip()
        if utp_name and utp_name not in hints:
            hints.append(utp_name)
    return program_title or (names[0] if names else None), tuple(hints)


def _program_header_line(
    utp: UtpParseResult,
    *,
    program_title: str | None = None,
    study_year_hints: tuple[str | None, ...] = (),
) -> str:
    program = (program_title or utp.metadata.program_name or "").strip()
    program = program.strip("«»\"'") or "название программы"
    weekly = utp.metadata.hours_per_week
    weekly_part = (
        f"({weekly} {_hours_word(weekly)} в неделю)"
        if weekly
        else "(количество часов в неделю)"
    )
    return f"«{program}» — {_study_year_label(utp, study_year_hints)} {weekly_part}"


_GROUP_CLASS_RE = re.compile(
    r"Группа\s*№\s*(?:_{2,}|\S+)\s*\(\s*Класс\s+(?:_{2,}|\S+)\s*\)",
    flags=re.IGNORECASE,
)


def _group_class_line(
    group_number: str | None = None,
    class_name: str | None = None,
) -> str:
    group = (group_number or "").strip() or "___________"
    klass = (class_name or "").strip() or "_________"
    return f"Группа № {group} (Класс {klass})"


def _replace_group_class_prefix(
    text: str,
    group_number: str | None,
    class_name: str | None,
) -> str:
    if "группа" not in text.casefold() or not (group_number or class_name):
        return text
    if not _GROUP_CLASS_RE.search(text):
        return text
    return _GROUP_CLASS_RE.sub(_group_class_line(group_number, class_name), text, count=1)


def _append_teacher_name_to_group_line(text: str, teacher_name: str | None) -> str:
    if "группа" not in text.casefold():
        return text
    match = _GROUP_CLASS_RE.search(text.replace("\t", ""))
    if match is None:
        match = _GROUP_CLASS_RE.search(text)
    if match is None:
        return text
    group = match.group(0)
    name = (teacher_name or "").strip()
    if name:
        return f"\t{group}\t{name}"
    return f"\t{group}"


def _content_width_twips(document) -> int:
    section = document.sections[0]
    return int(
        section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    )


def _apply_group_teacher_tabs(paragraph, document) -> None:
    """Группа/класс по центру, ФИО справа: один абзац, center tab + right tab."""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    width = _content_width_twips(document)
    properties = paragraph._p.get_or_add_pPr()
    tabs = properties.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        properties.append(tabs)
    for old in list(tabs.findall(qn("w:tab"))):
        tabs.remove(old)
    center = OxmlElement("w:tab")
    center.set(qn("w:val"), "center")
    center.set(qn("w:pos"), str(width // 2))
    tabs.append(center)
    right = OxmlElement("w:tab")
    right.set(qn("w:val"), "right")
    right.set(qn("w:pos"), str(width))
    tabs.append(right)


def _enable_cell_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is not None:
        tc_pr.remove(no_wrap)


def _first_run_properties(paragraph):
    for run in paragraph.runs:
        properties = run._r.find(qn("w:rPr"))
        if properties is not None:
            return deepcopy(properties)
    paragraph_properties = paragraph._p.find(qn("w:pPr"))
    if paragraph_properties is not None:
        properties = paragraph_properties.find(qn("w:rPr"))
        if properties is not None:
            return deepcopy(properties)
    return None


def _strip_header_sample_paragraph_layout(paragraph_properties) -> None:
    """Убрать свойства пустой/заголовочной строки, которые растягивают данные.

    Шапка таблицы не вызывается отсюда. Размер и шрифт остаются.
    firstLine/tabs сужают колонку; jc=center — выравнивание шапки, не текста данных.
    Без явного spacing действует docDefaults шаблона (after=200, line=276).
    """

    if paragraph_properties is None:
        return
    for tag in ("w:ind", "w:tabs", "w:jc"):
        node = paragraph_properties.find(qn(tag))
        if node is not None:
            paragraph_properties.remove(node)
    spacing = paragraph_properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        paragraph_properties.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    snap = paragraph_properties.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        paragraph_properties.append(snap)
    snap.set(qn("w:val"), "0")


def _strip_cloned_sample_row_layout(table_row) -> None:
    """Строка данных не наследует фиксированную высоту образца шапки/пустышки."""

    row_properties = table_row.find(qn("w:trPr"))
    if row_properties is not None:
        height = row_properties.find(qn("w:trHeight"))
        if height is not None:
            row_properties.remove(height)
        leaked_header = row_properties.find(qn("w:tblHeader"))
        if leaked_header is not None:
            row_properties.remove(leaked_header)
    for cell in table_row.findall(qn("w:tc")):
        for paragraph in cell.findall(qn("w:p")):
            paragraph_properties = paragraph.find(qn("w:pPr"))
            if paragraph_properties is None:
                paragraph_properties = OxmlElement("w:pPr")
                paragraph.insert(0, paragraph_properties)
            _strip_header_sample_paragraph_layout(paragraph_properties)


def _without_bold(properties):
    """Свойства шрифта без начертания: данные календаря должны быть regular."""

    if properties is None:
        return None
    cleaned = deepcopy(properties)
    for tag in ("w:b", "w:bCs"):
        node = cleaned.find(qn(tag))
        if node is not None:
            cleaned.remove(node)
    bold = OxmlElement("w:b")
    bold.set(qn("w:val"), "0")
    bold_cs = OxmlElement("w:bCs")
    bold_cs.set(qn("w:val"), "0")
    cleaned.insert(0, bold_cs)
    cleaned.insert(0, bold)
    return cleaned


def _set_run_font_family(run, family: str) -> None:
    """Явно закрепить одно семейство шрифта для всех наборов символов OOXML."""

    run.font.name = family
    run_properties = run._r.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), family)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        qualified = qn(f"w:{attribute}")
        if qualified in fonts.attrib:
            del fonts.attrib[qualified]


def _apply_standard_table_font(table) -> None:
    """Унифицировать шрифт стандартной таблицы, сохранив размер и начертание."""

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font_family(run, STANDARD_TABLE_FONT_FAMILY)


def _apply_standard_header_font(document) -> None:
    """Закрепить шрифт создаваемой шапки standard-template без смены начертания."""

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font_family(run, STANDARD_TABLE_FONT_FAMILY)


def _set_paragraph_text_keep_format(paragraph, text: str, run_properties=None) -> None:
    """Заменить текст абзаца, сохранив pPr и rPr шаблона."""

    seed = run_properties if run_properties is not None else _first_run_properties(paragraph)
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if seed is not None:
        run._r.insert(0, deepcopy(seed))


def _set_cell_text(cell, value: str) -> None:
    """Подставить текст ячейки, сохранив шрифт и размер.

    Переносы строк — мягкие w:br в одном абзаце, как при cell.text,
    а не отдельные w:p: лишние абзацы наследуют межстрочный зазор шаблона
    и растягивают строку таблицы.
    """

    _enable_cell_wrap(cell)
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph()
        paragraphs = list(cell.paragraphs)
    seed_paragraph = paragraphs[0]
    seed_rpr = _without_bold(_first_run_properties(seed_paragraph))
    seed_ppr = seed_paragraph._p.get_or_add_pPr()
    _strip_header_sample_paragraph_layout(seed_ppr)
    mark = seed_ppr.find(qn("w:rPr"))
    if mark is not None:
        for tag in ("w:b", "w:bCs"):
            node = mark.find(qn(tag))
            if node is not None:
                mark.remove(node)
    for extra in list(cell.paragraphs)[1:]:
        cell._tc.remove(extra._p)
    _set_paragraph_text_keep_format(seed_paragraph, value, seed_rpr)


def _center_cell_text(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _ensure_text_direction_bt_lr(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    direction = tc_pr.find(qn("w:textDirection"))
    if direction is None:
        direction = OxmlElement("w:textDirection")
        tc_pr.append(direction)
    direction.set(qn("w:val"), "btLr")


def _set_cell_vertical_align(cell, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    valign = tc_pr.find(qn("w:vAlign"))
    if valign is None:
        valign = OxmlElement("w:vAlign")
        tc_pr.append(valign)
    valign.set(qn("w:val"), value)


def _set_month_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag in ("top", "left", "bottom", "right"):
        side = tc_mar.find(qn(f"w:{tag}"))
        if side is None:
            side = OxmlElement(f"w:{tag}")
            tc_mar.append(side)
        side.set(qn("w:w"), str(_MONTH_CELL_MARGIN_DXA))
        side.set(qn("w:type"), "dxa")


def _set_month_paragraph_format(cell) -> None:
    for paragraph in cell.paragraphs:
        # Для вертикального текста центрирование абзаца даёт визуальный центр
        # внутри объединённой ячейки.
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_properties = paragraph._p.get_or_add_pPr()
        text_alignment = paragraph_properties.find(qn("w:textAlignment"))
        if text_alignment is None:
            text_alignment = OxmlElement("w:textAlignment")
            paragraph_properties.append(text_alignment)
        text_alignment.set(qn("w:val"), "center")
        spacing = paragraph_properties.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            paragraph_properties.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")


def _remove_vmerge(cell) -> None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return
    merge = tc_pr.find(qn("w:vMerge"))
    if merge is not None:
        tc_pr.remove(merge)


def _set_vmerge_restart(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    merge = tc_pr.find(qn("w:vMerge"))
    if merge is None:
        merge = OxmlElement("w:vMerge")
        tc_pr.insert(0, merge)
    merge.set(qn("w:val"), "restart")


def _set_vmerge_continue(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    merge = tc_pr.find(qn("w:vMerge"))
    if merge is None:
        merge = OxmlElement("w:vMerge")
        tc_pr.insert(0, merge)
    # Пустой w:vMerge без val = continue (стандарт OOXML).
    if merge.get(qn("w:val")) is not None:
        del merge.attrib[qn("w:val")]


def _prepare_month_cell(cell, month: str) -> None:
    _set_cell_text(cell, month)
    _ensure_text_direction_bt_lr(cell)
    _set_cell_vertical_align(cell, "center")
    _set_month_cell_margins(cell)
    _set_month_paragraph_format(cell)


def _format_month_cell(cell, month: str) -> None:
    """Подпись месяца по центру объединённого блока на странице."""

    _prepare_month_cell(cell, month)


def _clear_month_continuation_cell(cell) -> None:
    _set_cell_text(cell, "")
    _ensure_text_direction_bt_lr(cell)
    _set_cell_vertical_align(cell, "center")


def _write_row(
    row,
    columns: _TableColumns,
    *,
    month: str,
    week_number: int,
    date_range: str,
    theory_text: str,
    practice_text: str,
    lesson_type: str,
    planned_result: str,
    assessment_method: str,
) -> None:
    cells = row.cells
    _format_month_cell(cells[columns.month], month)
    _set_cell_text(cells[columns.week], f"{week_number}\n{date_range}")
    _center_cell_text(cells[columns.week])
    _set_cell_text(cells[columns.theory], theory_text)
    _set_cell_text(cells[columns.theory_mark], "")
    _set_cell_text(cells[columns.practice], practice_text)
    _set_cell_text(cells[columns.lesson_type], lesson_type)
    _set_cell_text(cells[columns.planned_result], planned_result)
    _set_cell_text(cells[columns.assessment], assessment_method)
    if columns.lesson_type_mirror is not None:
        _set_cell_text(cells[columns.lesson_type_mirror], lesson_type)
    if columns.planned_result_mirror is not None:
        _set_cell_text(cells[columns.planned_result_mirror], planned_result)


def _week_topic_parts(source_row) -> tuple[WeekTopicPart, ...]:
    if source_row.week_parts:
        return source_row.week_parts
    return (
        WeekTopicPart(
            topic_number=source_row.topic_number,
            topic_title=source_row.topic_title,
            section=source_row.section,
            theory_hours=source_row.theory_hours,
            practice_hours=source_row.practice_hours,
            match_status=source_row.match_status,
            program_section=source_row.program_section,
            program_topic=source_row.program_topic,
            program_content_full=source_row.program_content_full,
            warnings=source_row.warnings,
        ),
    )


def _topic_part_key(part: WeekTopicPart) -> tuple[str | None, str, str]:
    return (part.topic_number, part.topic_title, part.section)


def _practice_appearance_counts(
    rows: tuple[ResolvedLessonRow, ...],
) -> dict[tuple[str | None, str, str], int]:
    counts: dict[tuple[str | None, str, str], int] = {}
    for lesson in rows:
        for part in _week_topic_parts(lesson.source.source):
            if part.practice_hours <= 0:
                continue
            key = _topic_part_key(part)
            counts[key] = counts.get(key, 0) + 1
    return counts


def _topic_cells_for_lesson(
    lesson: ResolvedLessonRow,
    display_numbers: dict[tuple[str | None, str, str], str],
    *,
    topic_counts: dict[tuple[str | None, str, str], int],
    topic_occurrences: dict[tuple[str | None, str, str], int],
) -> tuple[str, str]:
    source_row = lesson.source.source
    theory_lines: list[str] = []
    practice_lines: list[str] = []
    for part in _week_topic_parts(source_row):
        key = _topic_part_key(part)
        occurrence_index = 0
        if part.practice_hours:
            occurrence_index = topic_occurrences.get(key, 0)
            topic_occurrences[key] = occurrence_index + 1
        display_number = display_numbers.get(key, part.topic_number or "?")
        theory_cell = format_theory_cell(
            display_number,
            part.topic_title,
            part.program_content_full,
            part.theory_hours,
        )
        selected_clause = ""
        appearance_count = topic_counts.get(key, 0)
        if appearance_count > 1 and part.practice_hours:
            selected_clause = practice_clause_for_repeated_topic(
                topic_title=part.topic_title,
                content=part.program_content_full,
                theory_hours=part.theory_hours,
                practice_hours=part.practice_hours,
                occurrence_index=occurrence_index,
                planned_result=lesson.planned_result,
                appearance_count=appearance_count,
            )
        practice_cell = format_practice_cell(
            display_number,
            part.topic_title,
            part.program_content_full,
            part.practice_hours,
            selected_clause,
        )
        if theory_cell:
            theory_lines.append(theory_cell)
        if practice_cell:
            practice_lines.append(practice_cell)
    return "\n".join(theory_lines), "\n".join(practice_lines)


def _apply_print_safe_margins(document) -> None:
    """Верхнее поле 1 см — заголовок и шапка таблицы в безопасной зоне печати."""

    target = Cm(PRINT_TOP_MARGIN_CM)
    for section in document.sections:
        section.top_margin = target


def _load_template(template: CalendarTemplateSelection) -> Document:
    if template.source is CalendarTemplateSource.STANDARD:
        if not STANDARD_TEMPLATE_PATH.is_file():
            raise FileNotFoundError(
                f"Стандартный шаблон не найден: {STANDARD_TEMPLATE_PATH}"
            )
        document = Document(str(STANDARD_TEMPLATE_PATH))
    else:
        assert template.content is not None
        document = Document(BytesIO(template.content))
    _apply_print_safe_margins(document)
    return document


def _ensure_data_rows(table, expected_rows: int) -> None:
    source_index = 2 if len(table.rows) > 2 else len(table.rows) - 1
    prototype = deepcopy(table.rows[source_index]._tr)
    _strip_cloned_sample_row_layout(prototype)
    while len(table.rows) < 2 + expected_rows:
        table._tbl.append(deepcopy(prototype))
    while len(table.rows) > 2 + expected_rows:
        table._tbl.remove(table.rows[-1]._tr)
    for row in table.rows[2:]:
        _strip_cloned_sample_row_layout(row._tr)


def _month_cell_at(row, column_index: int):
    """Ячейка месяца по физическому tc, без remap python-docx после vMerge."""

    from docx.table import _Cell

    return _Cell(row._tr.tc_lst[column_index], row)


def _merge_month_cell_group(
    group_rows,
    columns: _TableColumns,
    month: str,
) -> None:
    """Объединить месяц внутри одного сегмента страницы через w:vMerge.

    Сегменты включают только целые строки одной измеренной страницы;
    после merge границы проверяются повторным рендером. Поэтому:
    - одинаковые месяцы сливаются в один блок на странице;
    - переходная неделя («Сентябрь / Октябрь») остаётся отдельной строкой;
    - при продолжении месяца на новом листе подпись появляется снова;
    - текст (btLr + vAlign=center) оказывается строго по центру блока,
      начинаясь с первой строки сегмента (без пустой колонки вверху страницы).
    """
    if not group_rows:
        return

    # Важно брать tc до установки vMerge: иначе row.cells[i] начинает
    # возвращать ячейку restart и затирает подпись при очистке continue.
    cells = [_month_cell_at(row, columns.month) for row in group_rows]
    first_cell = cells[0]
    _format_month_cell(first_cell, month)
    if len(cells) == 1:
        _remove_vmerge(first_cell)
        return

    _set_vmerge_restart(first_cell)
    for cell in cells[1:]:
        _clear_month_continuation_cell(cell)
        _set_vmerge_continue(cell)


def _merge_month_cells(table, columns: _TableColumns, months: tuple[str, ...]) -> None:
    if not months:
        return

    group_start = 0
    for index in range(1, len(months) + 1):
        if index < len(months) and months[index] == months[group_start]:
            continue

        group_rows = table.rows[2 + group_start : 2 + index]
        _merge_month_cell_group(group_rows, columns, months[group_start])
        group_start = index


def _set_page_break_before_row(row, columns: _TableColumns) -> None:
    """Жёстко начать строку с новой страницы в Microsoft Word.

    Границы страниц считаются по пагинации Word. Явный pageBreakBefore
    закрепляет найденные старты страниц в итоговом DOCX, чтобы merge месяца
    не «уездал» при повторном открытии.
    """

    cell = row.cells[columns.theory]
    paragraph = cell.paragraphs[0]
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = p_pr.find(qn("w:pageBreakBefore"))
    if page_break is None:
        p_pr.append(OxmlElement("w:pageBreakBefore"))


def _clear_page_break_before_row(row, columns: _TableColumns) -> None:
    cell = row.cells[columns.theory]
    for paragraph in cell.paragraphs:
        p_pr = paragraph._p.pPr
        if p_pr is None:
            continue
        page_break = p_pr.find(qn("w:pageBreakBefore"))
        if page_break is not None:
            p_pr.remove(page_break)


def _apply_explicit_page_breaks(
    table,
    columns: _TableColumns,
    rows_by_page: tuple[tuple[int, ...], ...],
) -> None:
    """Закрепить начало каждой рассчитанной страницы в самом DOCX."""

    for row in table.rows[2:]:
        _clear_page_break_before_row(row, columns)

    for page_rows in rows_by_page[1:]:
        if not page_rows:
            continue
        _set_page_break_before_row(table.rows[2 + page_rows[0]], columns)


def _merge_month_cells_by_page_segments(
    table,
    columns: _TableColumns,
    months: tuple[str, ...],
    rows_by_page: tuple[tuple[int, ...], ...],
) -> None:
    if not months:
        return

    for page_rows in rows_by_page:
        if not page_rows:
            continue

        group_start = page_rows[0]
        for index in range(1, len(page_rows)):
            row_index = page_rows[index]
            if (months[row_index] != months[group_start]
                    or row_index != page_rows[index - 1] + 1):
                group_rows = table.rows[2 + group_start : 2 + page_rows[index - 1] + 1]
                _merge_month_cell_group(group_rows, columns, months[group_start])
                group_start = row_index

        group_rows = table.rows[2 + group_start : 2 + page_rows[-1] + 1]
        _merge_month_cell_group(group_rows, columns, months[group_start])


def _quoted_program_name(program_title: str | None, utp: UtpParseResult) -> str:
    program = (program_title or utp.metadata.program_name or "").strip().strip("«»\"'")
    return program


def _study_year_number(utp: UtpParseResult, study_year_hints: tuple[str | None, ...]) -> int | None:
    for raw in (utp.metadata.study_year, *study_year_hints):
        number = infer_study_year_number(raw)
        if number is not None:
            return number
    return None


def _fill_hours_paren(paren: str, *, weekly: int | None, yearly: int | None) -> str:
    weekly_slot = "недел" in paren.casefold()
    hours = weekly if weekly_slot and weekly else (yearly or weekly)
    if not hours:
        return paren
    if re.search(r"\d+", paren):
        return re.sub(r"\d+", str(hours), paren, count=1)
    if weekly_slot:
        return f"({hours} {_hours_word(hours)} в неделю)"
    return f"({hours})"


def _fill_organization_header_paragraph(
    text: str,
    utp: UtpParseResult,
    *,
    program_title: str | None,
    study_year_hints: tuple[str | None, ...],
    group_number: str | None,
    class_name: str | None,
    teacher_name: str | None = None,
) -> str:
    updated = text
    program = _quoted_program_name(program_title, utp)
    if program and "«" in updated and "»" in updated:
        updated = re.sub(r"«[^»]*»", f"«{program}»", updated, count=1)
    elif program and "название программы" in updated.casefold():
        updated = re.sub(r"название программы", program, updated, count=1, flags=re.IGNORECASE)

    year_number = _study_year_number(utp, study_year_hints)
    if year_number is not None:
        if re.search(r"(?:_{2,}|\d+)\s+год обучения", updated, flags=re.IGNORECASE):
            updated = re.sub(
                r"(?:_{2,}|\d+)\s+год обучения",
                f"{year_number} год обучения",
                updated,
                count=1,
                flags=re.IGNORECASE,
            )
        elif re.search(r"(?:_{2,}|\d+)\s*г\.об", updated, flags=re.IGNORECASE):
            updated = re.sub(
                r"(?:_{2,}|\d+)\s*г\.об\.?",
                f"{year_number} г.об.",
                updated,
                count=1,
                flags=re.IGNORECASE,
            )

    if "(" in updated and ")" in updated and (
        "час" in updated.casefold() or "чса" in updated.casefold() or "недел" in updated.casefold()
    ):
        updated = re.sub(
            r"\([^()]*\)",
            lambda match: _fill_hours_paren(
                match.group(0),
                weekly=utp.metadata.hours_per_week,
                yearly=utp.metadata.hours_per_year,
            ),
            updated,
            count=1,
        )

    updated = _replace_group_class_prefix(updated, group_number, class_name)
    return _append_teacher_name_to_group_line(updated, teacher_name)


def _fill_organization_header(
    document,
    utp: UtpParseResult,
    *,
    program_title: str | None = None,
    study_year_hints: tuple[str | None, ...] = (),
    group_number: str | None = None,
    class_name: str | None = None,
    teacher_name: str | None = None,
) -> None:
    """Подставить значения в шапку шаблона организации, не меняя её состав и стили."""

    for paragraph in document.paragraphs:
        original = paragraph.text
        if not original.strip():
            continue
        updated = _fill_organization_header_paragraph(
            original,
            utp,
            program_title=program_title,
            study_year_hints=study_year_hints,
            group_number=group_number,
            class_name=class_name,
            teacher_name=teacher_name,
        )
        if updated != original:
            _set_paragraph_text_keep_format(paragraph, updated)
        if "группа" in paragraph.text.casefold() and _GROUP_CLASS_RE.search(
            paragraph.text.replace("\t", "")
        ):
            _apply_group_teacher_tabs(paragraph, document)


def _write_document_header(
    document,
    utp: UtpParseResult,
    *,
    academic_year: str,
    program_title: str | None = None,
    study_year_hints: tuple[str | None, ...] = (),
    group_number: str | None = None,
    class_name: str | None = None,
    teacher_name: str | None = None,
    uses_organization_template: bool = False,
) -> None:
    if uses_organization_template:
        _fill_organization_header(
            document,
            utp,
            program_title=program_title,
            study_year_hints=study_year_hints,
            group_number=group_number,
            class_name=class_name,
            teacher_name=teacher_name,
        )
        return

    if document.paragraphs:
        _set_paragraph_text_keep_format(document.paragraphs[0], "Календарный план")
    program_line = _program_header_line(
        utp,
        program_title=program_title,
        study_year_hints=study_year_hints,
    )
    year_line = f"{academic_year} учебный год"
    group_line = _group_class_line(group_number, class_name)
    teacher = (teacher_name or "").strip()
    group_teacher_line = f"\t{group_line}\t{teacher}" if teacher else group_line
    group_paragraph = None
    if len(document.paragraphs) > 3:
        _set_paragraph_text_keep_format(document.paragraphs[1], program_line)
        _set_paragraph_text_keep_format(document.paragraphs[2], year_line)
        group_paragraph = document.paragraphs[3]
        _set_paragraph_text_keep_format(group_paragraph, group_teacher_line)
    elif len(document.paragraphs) > 2:
        _set_paragraph_text_keep_format(
            document.paragraphs[1], f"{program_line}\n{year_line}"
        )
        group_paragraph = document.paragraphs[2]
        _set_paragraph_text_keep_format(group_paragraph, group_teacher_line)
    elif len(document.paragraphs) > 1:
        _set_paragraph_text_keep_format(
            document.paragraphs[1],
            f"{program_line}\n{year_line}\n{group_teacher_line}",
        )
        group_paragraph = document.paragraphs[1]
    if teacher and group_paragraph is not None:
        _apply_group_teacher_tabs(group_paragraph, document)
    if group_paragraph is not None:
        group_paragraph.paragraph_format.space_after = Pt(
            STANDARD_GROUP_SPACE_AFTER_PT
        )
    for paragraph in document.paragraphs:
        blob = paragraph.text.casefold()
        if "название программы" in blob or "г.об" in blob:
            _set_paragraph_text_keep_format(paragraph, program_line)
    _apply_standard_header_font(document)


def _populate_calendar_table(
    document,
    utp: UtpParseResult,
    rows: tuple[ResolvedLessonRow, ...],
    *,
    academic_year: str,
    program_title: str | None = None,
    study_year_hints: tuple[str | None, ...] = (),
    group_number: str | None = None,
    class_name: str | None = None,
    teacher_name: str | None = None,
    uses_organization_template: bool = False,
) -> tuple:
    """Заполнить таблицу календаря строками данных (без объединения месяцев)."""

    resolved_title, resolved_hints = _resolve_header_from_rows(
        rows,
        program_title=program_title,
        study_year_hints=study_year_hints,
    )
    _write_document_header(
        document,
        utp,
        academic_year=academic_year,
        program_title=resolved_title,
        study_year_hints=resolved_hints,
        group_number=group_number,
        class_name=class_name,
        teacher_name=teacher_name,
        uses_organization_template=uses_organization_template,
    )

    table = document.tables[0]
    columns = _columns_for_table(table)
    display_numbers = _topic_display_numbers(utp)
    _ensure_data_rows(table, len(rows))
    _repeat_table_header_rows(table)
    topic_counts = _practice_appearance_counts(rows)
    topic_occurrences: dict[tuple[str | None, str, str], int] = {}

    for index, lesson in enumerate(rows):
        theory_cell, practice_cell = _topic_cells_for_lesson(
            lesson,
            display_numbers,
            topic_counts=topic_counts,
            topic_occurrences=topic_occurrences,
        )
        source = lesson.source.source
        _write_row(
            table.rows[index + 2],
            columns,
            month=source.month,
            week_number=source.week_number,
            date_range=source.date_range,
            theory_text=theory_cell,
            practice_text=practice_cell,
            lesson_type=lesson.lesson_type,
            planned_result=lesson.planned_result,
            assessment_method=lesson.assessment_method,
        )
        _allow_row_split(table.rows[index + 2])

    if not uses_organization_template:
        _apply_standard_table_font(table)

    _protect_vertical_cell_height(table)

    months = tuple(lesson.source.source.month for lesson in rows)
    return table, columns, months


def _save_document(document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _merge_month_cells_for_pages(
    document,
    utp: UtpParseResult,
    rows: tuple[ResolvedLessonRow, ...],
    rows_by_page: tuple[tuple[int, ...], ...],
    *,
    keep_together: frozenset[int] = frozenset(),
    academic_year: str,
    program_title: str | None = None,
    study_year_hints: tuple[str | None, ...] = (),
    group_number: str | None = None,
    class_name: str | None = None,
    teacher_name: str | None = None,
    uses_organization_template: bool = False,
) -> bytes:
    """Собрать DOCX с объединением месяцев по сегментам страниц."""

    table, columns, months = _populate_calendar_table(
        document,
        utp,
        rows,
        academic_year=academic_year,
        program_title=program_title,
        study_year_hints=study_year_hints,
        group_number=group_number,
        class_name=class_name,
        teacher_name=teacher_name,
        uses_organization_template=uses_organization_template,
    )
    for index in keep_together:
        _prevent_row_split(table.rows[index + 2])
    for page_rows in rows_by_page:
        for index in page_rows:
            _prevent_row_split(table.rows[index + 2])
    _merge_month_cells_by_page_segments(table, columns, months, rows_by_page)
    return _save_document(document)


def build_output_filename(utp: UtpParseResult, academic_year: str) -> str:
    program = utp.metadata.program_name or "программа"
    safe = re.sub(r'[<>:"/\\|?*\s]+', "_", program).strip("_") or "календарь"
    year = academic_year.replace("–", "-")
    return f"Календарный_план_{safe}_{year}.docx"


def generate_calendar_docx(
    utp: UtpParseResult,
    rows: tuple[ResolvedLessonRow, ...],
    template: CalendarTemplateSelection,
    academic_year: str,
    *,
    program_title: str | None = None,
    study_year_hints: tuple[str | None, ...] = (),
    group_number: str | None = None,
    class_name: str | None = None,
    teacher_name: str | None = None,
) -> bytes:
    """Сформировать DOCX: месяц совпадает с датой и merge не пересекает страницы."""

    header = {
        "academic_year": academic_year,
        "program_title": program_title,
        "study_year_hints": study_year_hints,
        "group_number": group_number,
        "class_name": class_name,
        "teacher_name": teacher_name,
        "uses_organization_template": template.uses_organization_template,
    }

    # Первый проход: без merge. Он нужен, чтобы получить фактическую пагинацию
    # Microsoft Word для текущего шаблона и объёма текста.
    preview_document = _load_template(template)
    preview_table, _, _ = _populate_calendar_table(preview_document, utp, rows, **header)
    preview = _save_document(preview_document)

    def conservative_fallback():
        # No unverified split, merge or forced break when measurement fails.
        for row in preview_table.rows[2:]:
            _prevent_row_split(row)
        return _save_document(preview_document)

    from calendar_pedagoga.docx_qa import detect_data_row_page_spans

    spans = detect_data_row_page_spans(preview, total_rows=len(rows))
    if not spans:
        # Без надёжной пагинации безопаснее оставить месяц в каждой строке,
        # чем объединить его через границу страницы и получить пустой столбец.
        return conservative_fallback()

    # Merge only complete, contiguous rows on the same measured page. Never
    # turn a preview page boundary into a forced break in the final document.
    keep_together = set()
    for _ in range(len(rows) + 2):
        keep_together.update(index for index, span in enumerate(spans)
                             if span.start_page != span.end_page and not span.split_safe)
        pages = {}
        for index, span in enumerate(spans):
            if span.start_page == span.end_page:
                pages.setdefault(span.start_page, []).append(index)
        rows_by_page = tuple(tuple(indices) for indices in pages.values())
        document = _load_template(template)
        merged = _merge_month_cells_for_pages(
            document, utp, rows, rows_by_page, keep_together=frozenset(keep_together), **header
        )
        detected = detect_data_row_page_spans(merged, total_rows=len(rows))
        if detected == spans:
            if any(span.start_page != span.end_page and not span.split_safe for span in detected):
                return conservative_fallback()
            return merged
        if not detected:
            return conservative_fallback()
        spans = detected

    # An unstable merge must not impose an unverified pagination constraint.
    return conservative_fallback()
