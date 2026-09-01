"""Локальный разбор учебно-тематических планов в формате DOCX."""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import re
from typing import BinaryIO

from docx import Document


@dataclass(frozen=True)
class Hours:
    total: int
    theory: int
    practice: int


@dataclass(frozen=True)
class Topic:
    number: str | None
    title: str
    hours: Hours
    parent_section: str | None = None
    is_standalone_section: bool = False


@dataclass(frozen=True)
class Section:
    number: str | None
    title: str
    hours: Hours
    is_standalone_position: bool = False


@dataclass(frozen=True)
class UtpMetadata:
    program_name: str | None = None
    academic_year: str | None = None
    study_year: str | None = None
    student_age: str | None = None
    hours_per_week: int | None = None
    hours_per_year: int | None = None
    study_weeks: int | None = None
    teacher_name: str | None = None
    stated_schedule_hours: int | None = None
    workload_provenance: str | None = None


@dataclass(frozen=True)
class UtpParseResult:
    metadata: UtpMetadata
    sections: tuple[Section, ...]
    topics: tuple[Topic, ...]
    table_totals: Hours | None
    warnings: tuple[str, ...] = field(default_factory=tuple)


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _match(text: str, pattern: str) -> str | None:
    found = re.search(pattern, text, re.IGNORECASE)
    return _clean(found.group(1)) if found else None


def _integer(value: str | None) -> int:
    cleaned = _clean(value or "")
    if not cleaned or cleaned in {"-", "–", "—"}:
        return 0
    found = re.search(r"\d+", cleaned)
    return int(found.group()) if found else 0


def _number_and_title(value: str) -> tuple[str | None, str]:
    value = _clean(value)
    found = re.match(r"^(\d+(?:\.\d+)*)\.?\s*(.*)$", value)
    if not found:
        return None, value
    return found.group(1), _clean(found.group(2))


def _metadata(paragraphs: list[str]) -> UtpMetadata:
    text = "\n".join(_clean(p) for p in paragraphs if _clean(p))
    academic_year = _match(text, r"план\s+на\s+(\d{4}\s*[-–]\s*\d{4})\s+учебн")
    if academic_year:
        academic_year = re.sub(r"\s*[-–]\s*", "–", academic_year)
    weekly = _match(text, r"Количество часов в неделю:\s*(\d+)")
    yearly = _match(text, r"Общее количество часов в год:\s*(\d+)")
    weeks = _match(text, r"(\d+)\s+учебн\w*\s+недел")
    schedule_hours = _match(text, r"\d+\s+учебн\w*\s+недел\w*\s+на\s+(\d+)\s+час")
    return UtpMetadata(
        program_name=_match(text, r"программ(?:ой|е)\s+[«\"]([^»\"]+)[»\"]"),
        academic_year=academic_year,
        study_year=_match(text, r"Год обучения:\s*([^\n]+)"),
        student_age=_match(text, r"Возраст обучающихся:\s*([^\n]+)"),
        hours_per_week=int(weekly) if weekly else None,
        hours_per_year=int(yearly) if yearly else None,
        study_weeks=int(weeks) if weeks else None,
        teacher_name=_match(text, r"Педагог дополнительного образования:\s*([^\n]+)"),
        stated_schedule_hours=int(schedule_hours) if schedule_hours else None,
    )


_CALENDAR_TABLE_MARKERS = (
    "месяц",
    "неделя",
    "теоретические занятия",
    "практические занятия",
    "планируемый результат",
    "вид контроля",
)


def _header_blob(table: object, rows: int = 3) -> str:
    parts: list[str] = []
    for row in table.rows[: min(rows, len(table.rows))]:
        parts.extend(cell.text for cell in row.cells)
    return " ".join(parts).casefold()


def _is_calendar_table(table: object) -> bool:
    blob = _header_blob(table)
    return sum(marker in blob for marker in _CALENDAR_TABLE_MARKERS) >= 3


def _utp_table_score(table: object) -> int:
    if _is_calendar_table(table):
        return -1
    blob = _header_blob(table)
    score = 0
    if "час" in blob:
        score += 3
    if "тем" in blob or "раздел" in blob:
        score += 2
    if re.search(r"теор|лекц", blob):
        score += 2
    if "практик" in blob:
        score += 2
    if "всего" in blob:
        score += 1
    joined_rows = [" ".join(cell.text for cell in row.cells) for row in table.rows]
    if any(re.search(r"\bитого\b", text, re.IGNORECASE) for text in joined_rows):
        score += 5
    columns = len(table.columns)
    if columns in {4, 5, 6}:
        score += 2
    if columns >= 8:
        score -= 4
    return score


def _parse_numbered_hours_rows(
    rows: list[list[str]],
    *,
    label_idx: int,
    title_idx: int,
    hour_idxs: tuple[int, int, int],
) -> tuple[list[Section], list[Topic], Hours | None]:
    sections: list[Section] = []
    topics: list[Topic] = []
    totals: Hours | None = None
    current_section: str | None = None
    started = False
    for raw_cells in rows:
        cells = [_clean(cell) for cell in raw_cells]
        if not any(cells):
            continue
        label = cells[label_idx] if label_idx < len(cells) else ""
        title = cells[title_idx] if title_idx < len(cells) else ""
        blob = " ".join(cells)
        if re.search(r"\bитого\b", blob, re.IGNORECASE):
            totals = Hours(*(_integer(cells[i]) if i < len(cells) else 0 for i in hour_idxs))
            continue
        number, label_title = _number_and_title(label)
        if number is None:
            if started:
                continue
            continue
        started = True
        hours = Hours(*(_integer(cells[i]) if i < len(cells) else 0 for i in hour_idxs))
        if "." not in number:
            section_title = label_title or _number_and_title(title)[1]
            current_section = section_title
            sections.append(Section(number, section_title, hours))
        else:
            topic_title = title
            if topic_title == label or topic_title == f"{number}.":
                fallback = cells[1] if len(cells) > 1 else title
                topic_title = fallback
            topics.append(
                Topic(
                    number,
                    _number_and_title(topic_title)[1],
                    hours,
                    parent_section=current_section,
                )
            )
    return sections, topics, totals


def _parse_six_column_table(
    rows: list[list[str]],
) -> tuple[list[Section], list[Topic], Hours | None]:
    sections: list[Section] = []
    topics: list[Topic] = []
    totals: Hours | None = None
    current_section: str | None = None
    for cells in rows[2:]:
        cells = [_clean(cell) for cell in cells]
        if not any(cells):
            continue
        label, title = cells[0], cells[2] or cells[1]
        if re.match(r"^итого\b", title, re.IGNORECASE):
            totals = Hours(*(_integer(cells[i]) for i in (3, 4, 5)))
            continue
        number, label_title = _number_and_title(label)
        if number is None:
            continue
        hours = Hours(*(_integer(cells[i]) for i in (3, 4, 5)))
        if "." not in number:
            section_title = label_title or _number_and_title(title)[1]
            current_section = section_title
            sections.append(Section(number, section_title, hours))
        else:
            topic_title = title
            if topic_title == label or topic_title == f"{number}.":
                topic_title = cells[1]
            topics.append(
                Topic(
                    number,
                    _number_and_title(topic_title)[1],
                    hours,
                    parent_section=current_section,
                )
            )
    return sections, topics, totals


def _parse_compact_table(
    table: object,
) -> tuple[list[Section], list[Topic], Hours | None]:
    sections: list[Section] = []
    topics: list[Topic] = []
    totals: Hours | None = None
    section_order = 0
    for row in table.rows[2:]:
        cells = [cell.text for cell in row.cells]
        raw_title = cells[0].strip()
        if "всего часов" in _clean(raw_title).lower():
            totals = Hours(*(_integer(cells[i]) for i in (1, 2, 3)))
            continue
        title_paragraphs = [
            paragraph for paragraph in row.cells[0].paragraphs if _clean(paragraph.text)
        ]
        titles = [_clean(paragraph.text) for paragraph in title_paragraphs]
        columns = [
            [_clean(line) for line in cells[i].splitlines() if _clean(line)]
            for i in (1, 2, 3)
        ]
        if not titles:
            continue
        section_order += 1
        section_number, section_title = _number_and_title(titles[0])
        if section_number is None:
            num_pr = title_paragraphs[0]._p.pPr.numPr
            if num_pr is not None and int(num_pr.ilvl.val) == 0:
                section_number = str(section_order)
        section_hours = Hours(*(_integer(col[0]) if col else 0 for col in columns))
        sections.append(
            Section(
                section_number,
                section_title,
                section_hours,
                is_standalone_position=len(titles) == 1,
            )
        )
        if len(titles) == 1:
            continue
        for index, raw_topic in enumerate(titles[1:], start=1):
            number, title = _number_and_title(raw_topic)
            values = [
                _integer(column[index]) if index < len(column) else 0
                for column in columns
            ]
            topics.append(
                Topic(
                    number,
                    title,
                    Hours(*values),
                    parent_section=section_title,
                )
            )
    return sections, topics, totals


def _looks_like_valid_utp(
    sections: list[Section],
    topics: list[Topic],
    table_totals: Hours | None,
) -> bool:
    if not sections or not topics:
        return False
    titled = [
        topic
        for topic in topics
        if topic.title and topic.title.casefold() not in {"тема", "№ п/п", "№ пп"}
    ]
    if len(titled) < 3:
        return False
    if table_totals is not None and table_totals.total > 0:
        return True
    return any(topic.hours.total > 0 for topic in titled)


def _parse_table_structure(
    table: object,
) -> tuple[list[Section], list[Topic], Hours | None]:
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    if not rows:
        raise ValueError("Пустая таблица УТП.")
    width = len(rows[0])
    if width >= 6:
        return _parse_six_column_table(rows)
    if width == 5:
        return _parse_numbered_hours_rows(
            rows,
            label_idx=0,
            title_idx=1,
            hour_idxs=(2, 3, 4),
        )
    if width >= 4:
        return _parse_compact_table(table)
    raise ValueError("Не удалось распознать структуру таблицы УТП.")


def _finalize_utp_parse(
    document,
    sections: list[Section],
    topics: list[Topic],
    table_totals: Hours | None,
) -> UtpParseResult:
    for section in sections:
        if not any(topic.parent_section == section.title for topic in topics):
            topics.append(
                Topic(
                    section.number,
                    section.title,
                    section.hours,
                    parent_section=section.title,
                    is_standalone_section=True,
                )
            )
    result = UtpParseResult(
        metadata=_metadata([paragraph.text for paragraph in document.paragraphs]),
        sections=tuple(sections),
        topics=tuple(topics),
        table_totals=table_totals,
    )
    from calendar_pedagoga.validation import validate_utp

    return UtpParseResult(
        metadata=result.metadata,
        sections=result.sections,
        topics=result.topics,
        table_totals=result.table_totals,
        warnings=tuple(validate_utp(result)),
    )


def parse_utp(source: str | Path | bytes | BinaryIO) -> UtpParseResult:
    """Разобрать УТП DOCX, не изменяя исходный файл."""
    document = Document(BytesIO(source) if isinstance(source, bytes) else source)
    if not document.tables:
        raise ValueError("В УТП не найдена таблица с темами.")
    ranked = sorted(
        document.tables,
        key=_utp_table_score,
        reverse=True,
    )
    last_error: Exception | None = None
    for table in ranked:
        if _utp_table_score(table) < 0:
            continue
        try:
            sections, topics, table_totals = _parse_table_structure(table)
        except (ValueError, IndexError) as error:
            last_error = error
            continue
        if _looks_like_valid_utp(sections, topics, table_totals):
            return _finalize_utp_parse(document, sections, topics, table_totals)
    if last_error is not None:
        raise ValueError("Не удалось распознать структуру таблицы УТП.") from last_error
    raise ValueError("В документе не найдена таблица УТП с темами и часами.")
